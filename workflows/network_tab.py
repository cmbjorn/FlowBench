"""
Harp manifold network solver — Streamlit UI.

Solves Z-manifold and U-manifold (reverse-return) harp configurations with up
to 500 parallel channels, optionally two harps in series.  Uses the Global
Gradient Algorithm (GGA / Newton-Raphson nodal analysis) for fast, robust
convergence including flow reversal.

Call render_harp_network_tab() from app.py.
"""
from __future__ import annotations

import hashlib
import json
import math

import plotly.graph_objects as go
import streamlit as st

import multiphase_engine as engine
from pipe_network import (
    HarpTopology, SolveResult, StarvationReport,
    build_harp, build_series_harp,
    check_starvation, solve_network, solve_network_hardy_cross,
)
from pipe_network.topology import build_center_fed_harp, build_biinlet_harp
from standards.piping import MATERIAL_ROUGHNESS, PIPE_DATABASE

# ── Constants ─────────────────────────────────────────────────────────────────

_DN_OPTIONS = list(PIPE_DATABASE.keys())
_PN_OPTIONS = ["PN20", "PN25", "PN40"]
_LIQUID_OPTIONS = list(engine.LIQUID_COOLPROP_ID.keys())

_PRESETS = {
    "H₂ / Water (electrolyser)": dict(
        phase="two-phase", gas_species="H₂", gas_kgh=8.0,
        liquid_type="Water", liquid_kgh=2400.0, P_bara=30.0, T_C=60.0,
    ),
    "CO₂ / Water": dict(
        phase="two-phase", gas_species="CO2", gas_kgh=20.0,
        liquid_type="Water", liquid_kgh=1500.0, P_bara=10.0, T_C=40.0,
    ),
    "Air / Water": dict(
        phase="two-phase", gas_species="Air", gas_kgh=15.0,
        liquid_type="Water", liquid_kgh=800.0, P_bara=5.0, T_C=25.0,
    ),
    "Water (single-phase)": dict(
        phase="single-phase", gas_species="H₂", gas_kgh=0.0,
        liquid_type="Water", liquid_kgh=2000.0, P_bara=5.0, T_C=25.0,
    ),
    "Custom": dict(
        phase="two-phase", gas_species="H₂", gas_kgh=5.0,
        liquid_type="Water", liquid_kgh=500.0, P_bara=10.0, T_C=50.0,
    ),
}

_GAS_COOLPROP = {
    "H₂":  "H2", "CO2": "CarbonDioxide", "Air": "Air",
    "N₂":  "Nitrogen", "O₂": "O2", "CH₄": "Methane", "H₂S": "H2S",
}

_OK_COLOR     = "#22c55e"
_WARN_COLOR   = "#f59e0b"
_STAR_COLOR   = "#ef4444"
_HEADER_COLOR = "#1e3a5f"
_CONN_COLOR   = "#64748b"


def _k(name: str) -> str:
    return f"hn_{name}"


# ── Pipe section widget ────────────────────────────────────────────────────────

def _pipe_section(
    title: str, kp: str,
    def_dn: str = "DN50", def_pn: str = "PN40",
    def_len_m: float = 0.15,
    def_id_mm: float = 50.0, def_rough_mm: float = 0.046,
    show_angle: bool = False,
    expanded: bool = True,
    length_label: str = "Length",
    length_help: str = "",
) -> tuple:
    """
    Render an expandable pipe section.
    Returns (D_inner_m, roughness_m, length_m [, angle_rad]).
    Supports both standard DN/PN lookup and custom inner diameter in mm.
    Length can be entered in m or mm.
    """
    with st.expander(title, expanded=expanded):
        use_custom = st.toggle(
            "Custom internal diameter",
            value=st.session_state.get(f"{kp}_custom", False),
            key=f"{kp}_custom",
        )
        if not use_custom:
            sc1, sc2 = st.columns(2)
            dn = sc1.selectbox("DN", _DN_OPTIONS,
                               index=_DN_OPTIONS.index(st.session_state.get(f"{kp}_dn", def_dn)),
                               key=f"{kp}_dn")
            pn = sc2.selectbox("PN", _PN_OPTIONS,
                               index=_PN_OPTIONS.index(st.session_state.get(f"{kp}_pn", def_pn)),
                               key=f"{kp}_pn")
            D_m   = PIPE_DATABASE[dn][pn]
            eps_m = MATERIAL_ROUGHNESS["SS316L"]
            st.caption(f"ID = {D_m*1000:.1f} mm")
        else:
            ic1, ic2 = st.columns(2)
            id_mm = ic1.number_input(
                "ID (mm)", min_value=0.1, max_value=2000.0,
                value=float(st.session_state.get(f"{kp}_id_mm", def_id_mm)),
                step=0.5, format="%.2f", key=f"{kp}_id_mm",
            )
            rough_mm = ic2.number_input(
                "Roughness (mm)", min_value=0.0, max_value=5.0,
                value=float(st.session_state.get(f"{kp}_rough_mm", def_rough_mm)),
                step=0.001, format="%.4f", key=f"{kp}_rough_mm",
                help="SS316L ≈ 0.046 mm",
            )
            D_m   = id_mm / 1000.0
            eps_m = rough_mm / 1000.0

        lc1, lc2 = st.columns([3, 1])
        len_unit = lc2.selectbox("Unit", ["m", "mm"], key=f"{kp}_len_unit",
                                  label_visibility="hidden")
        if len_unit == "mm":
            raw = lc1.number_input(
                f"{length_label} (mm)", min_value=0.0, max_value=100_000.0,
                value=float(st.session_state.get(f"{kp}_len_mm_val", def_len_m * 1000)),
                step=1.0, format="%.1f", key=f"{kp}_len_mm_val",
                help=length_help or None,
            )
            length_m = raw / 1000.0
        else:
            raw = lc1.number_input(
                f"{length_label} (m)", min_value=0.0, max_value=1_000.0,
                value=float(st.session_state.get(f"{kp}_len_m_val", def_len_m)),
                step=0.01, format="%.4f", key=f"{kp}_len_m_val",
                help=length_help or None,
            )
            length_m = raw

        if show_angle:
            _angle_opts = {
                "Horizontal (0°)":      0.0,
                "Vertical up (+90°)":   math.pi / 2,
                "Vertical down (−90°)": -math.pi / 2,
            }
            angle_lbl = st.selectbox("Orientation", list(_angle_opts.keys()),
                                     key=f"{kp}_angle_lbl")
            return D_m, eps_m, length_m, _angle_opts[angle_lbl]

        return D_m, eps_m, length_m


# ── Plotly diagrams ────────────────────────────────────────────────────────────

def _ch_color(r, Vsl_threshold: float, single_phase: bool,
              mean_m: float) -> str:
    if single_phase:
        frac = r.m_kgs / max(mean_m, 1e-12)
        if frac < 0.7: return _STAR_COLOR
        if frac < 0.9: return _WARN_COLOR
        return _OK_COLOR
    if r.starved:       return _STAR_COLOR
    if r.Vsl < 2 * Vsl_threshold: return _WARN_COLOR
    return _OK_COLOR


def _make_harp_diagram(
    N: int,
    solve_res: SolveResult,
    topo1: HarpTopology,
    topo2: HarpTopology | None,
    rep1: StarvationReport,
    rep2: StarvationReport | None,
    Vsl_threshold: float,
    single_phase: bool,
) -> go.Figure:
    from plotly.subplots import make_subplots

    net    = solve_res.net
    series = topo2 is not None and rep2 is not None

    if series:
        fig = make_subplots(rows=2, cols=1, vertical_spacing=0.08)
    else:
        fig = go.Figure()

    def _draw(topo: HarpTopology, report: StarvationReport,
              x_offset: float, y_top: float = 1.0, y_bot: float = 0.0,
              hide_inlet: bool = False, hide_outlet: bool = False,
              row: int = 1) -> None:
        n     = len(topo.channel_edge_ids)
        is_u  = (topo.harp_type == "U")
        is_i  = (topo.harp_type == "I")
        is_di = (topo.harp_type == "DI")

        a_pres = [net.node(nid).P_pa / 1e5 for nid in topo.header_A_node_ids]
        b_pres = [net.node(nid).P_pa / 1e5 for nid in topo.header_B_node_ids]

        # Axis references change for subplot row 2
        xax = f"x{'' if row == 1 else row}"
        yax = f"y{'' if row == 1 else row}"

        def _at(trace):
            if series:
                fig.add_trace(trace, row=row, col=1)
            else:
                fig.add_trace(trace)

        def _aa(**kw):
            if series:
                fig.add_annotation(xref=xax, yref=yax, **kw)
            else:
                fig.add_annotation(**kw)

        def _as(**kw):
            if series:
                fig.add_shape(xref=xax, yref=yax, **kw)
            else:
                fig.add_shape(**kw)

        # ── Headers ──────────────────────────────────────────────────────────
        _at(go.Scatter(
            x=[x_offset + i * n / max(n, 1) for i in range(n + 1)],
            y=[y_top] * (n + 1),
            mode="lines+markers",
            line=dict(color=_HEADER_COLOR, width=4),
            marker=dict(size=5, color=_HEADER_COLOR),
            hovertemplate=[f"<b>A{i}</b><br>P = {p:.4f} bara<extra></extra>"
                           for i, p in enumerate(a_pres)],
            showlegend=False,
        ))
        b_color = "#0f4c75" if is_u else _HEADER_COLOR
        _at(go.Scatter(
            x=[x_offset + i * n / max(n, 1) for i in range(n + 1)],
            y=[y_bot] * (n + 1),
            mode="lines+markers",
            line=dict(color=b_color, width=4, dash="dot" if is_u else "solid"),
            marker=dict(size=5, color=b_color),
            hovertemplate=[f"<b>B{i}</b><br>P = {p:.4f} bara<extra></extra>"
                           for i, p in enumerate(b_pres)],
            showlegend=False,
        ))

        # ── Channels ─────────────────────────────────────────────────────────
        # Group by colour → 3 traces max regardless of N (was 1 trace per channel).
        mean_m = report.mean_m_kgs
        if n <= 500:
            _seg_x:  dict[str, list] = {}
            _seg_y:  dict[str, list] = {}
            _seg_ht: dict[str, list] = {}
            _seg_w:  dict[str, list] = {}
            for idx, r in enumerate(report.channel_results):
                xc    = x_offset + idx * n / max(n, 1)
                color = _ch_color(r, Vsl_threshold, single_phase, mean_m)
                w     = max(1, min(8, int(r.m_kgs / max(mean_m, 1e-12) * 4)))
                ht    = (
                    f"<b>ch{idx}</b><br>"
                    f"m = {r.m_kgs*1000:.2f} g/s<br>"
                    f"Vsl = {r.Vsl:.4f} m/s<br>"
                    + (f"Vsg = {r.Vsg:.4f} m/s<br>x = {r.x_gas:.4f}<br>"
                       if not single_phase else "")
                    + f"{'⚠ LOW' if r.starved else '✓ OK'}"
                )
                if color not in _seg_x:
                    _seg_x[color]  = []
                    _seg_y[color]  = []
                    _seg_ht[color] = []
                    _seg_w[color]  = []
                _seg_x[color].extend([xc, xc, None])
                _seg_y[color].extend([y_bot, y_top, None])
                _seg_ht[color].extend([ht, ht, None])
                _seg_w[color].append(w)
            for color in _seg_x:
                avg_w = max(1, int(sum(_seg_w[color]) / len(_seg_w[color])))
                _at(go.Scatter(
                    x=_seg_x[color], y=_seg_y[color], mode="lines",
                    line=dict(color=color, width=avg_w),
                    text=_seg_ht[color],
                    hovertemplate="%{text}<extra></extra>",
                    showlegend=False,
                ))

        # ── Visual constants ──────────────────────────────────────────────────
        P_in  = net.node(topo.inlet_node_id).P_pa / 1e5
        P_out = net.node(topo.outlet_node_id).P_pa / 1e5
        stub  = max(n * 0.04, 0.5)
        cap   = (y_top - y_bot) * 0.18
        mid_x = x_offset + n * 0.5

        _IN_COLOR  = "#16a34a"
        _OUT_COLOR = "#dc2626"
        _CAP_COLOR = "#475569"

        def _inlet_stub(x_attach, stub_dir, y, label_y_shift=0):
            x_tip = x_attach + stub_dir * stub
            _at(go.Scatter(
                x=[x_attach, x_tip], y=[y, y], mode="lines",
                line=dict(color=_IN_COLOR, width=5),
                hoverinfo="skip", showlegend=False,
            ))
            _at(go.Scatter(
                x=[x_tip], y=[y], mode="markers",
                marker=dict(size=16, color=_IN_COLOR, symbol="arrow-right",
                            angle=0 if stub_dir < 0 else 180,
                            line=dict(color="white", width=1)),
                hoverinfo="skip", showlegend=False,
            ))
            _aa(x=x_tip, y=y + label_y_shift,
                text=f"<b>INLET</b><br><span style='font-size:10px'>{P_in:.3f} bara</span>",
                showarrow=False,
                xanchor="right" if stub_dir < 0 else "left",
                xshift=-6 if stub_dir < 0 else 6,
                font=dict(size=12, color=_IN_COLOR),
                bgcolor="rgba(255,255,255,0.85)", borderpad=3)

        def _outlet_stub(x_attach, stub_dir, y, label_y_shift=0):
            x_tip = x_attach + stub_dir * stub
            _at(go.Scatter(
                x=[x_attach, x_tip], y=[y, y], mode="lines",
                line=dict(color=_OUT_COLOR, width=5),
                hoverinfo="skip", showlegend=False,
            ))
            _at(go.Scatter(
                x=[x_tip], y=[y], mode="markers",
                marker=dict(size=16, color=_OUT_COLOR, symbol="arrow-right",
                            angle=0 if stub_dir > 0 else 180,
                            line=dict(color="white", width=1)),
                hoverinfo="skip", showlegend=False,
            ))
            _aa(x=x_tip, y=y + label_y_shift,
                text=f"<b>OUTLET</b><br><span style='font-size:10px'>{P_out:.3f} bara</span>",
                showarrow=False,
                xanchor="left" if stub_dir > 0 else "right",
                xshift=6 if stub_dir > 0 else -6,
                font=dict(size=12, color=_OUT_COLOR),
                bgcolor="rgba(255,255,255,0.85)", borderpad=3)

        def _cap(x_pos, y, label_below=True):
            _at(go.Scatter(
                x=[x_pos, x_pos], y=[y - cap, y + cap], mode="lines",
                line=dict(color=_CAP_COLOR, width=4),
                hoverinfo="skip", showlegend=False,
            ))
            yshift = -cap - 0.04 if label_below else cap + 0.04
            _aa(x=x_pos, y=y + yshift, text="<i>cap</i>",
                showarrow=False, font=dict(size=8, color=_CAP_COLOR))

        # ── Per-type layout ───────────────────────────────────────────────────
        if is_i:
            cx = x_offset + (n // 2)
            _at(go.Scatter(x=[cx, cx], y=[y_top, y_top + stub],
                           mode="lines", line=dict(color=_IN_COLOR, width=5),
                           hoverinfo="skip", showlegend=False))
            _at(go.Scatter(x=[cx], y=[y_top + stub], mode="markers",
                           marker=dict(size=16, color=_IN_COLOR, symbol="arrow-down",
                                       line=dict(color="white", width=1)),
                           hoverinfo="skip", showlegend=False))
            _aa(x=cx, y=y_top + stub + 0.03,
                text=f"<b>INLET</b><br><span style='font-size:10px'>{P_in:.3f} bara</span>",
                showarrow=False, xanchor="center",
                font=dict(size=12, color=_IN_COLOR),
                bgcolor="rgba(255,255,255,0.85)", borderpad=3)
            _at(go.Scatter(x=[cx, cx], y=[y_bot, y_bot - stub],
                           mode="lines", line=dict(color=_OUT_COLOR, width=5),
                           hoverinfo="skip", showlegend=False))
            _at(go.Scatter(x=[cx], y=[y_bot - stub], mode="markers",
                           marker=dict(size=16, color=_OUT_COLOR, symbol="arrow-down",
                                       line=dict(color="white", width=1)),
                           hoverinfo="skip", showlegend=False))
            _aa(x=cx, y=y_bot - stub - 0.03,
                text=f"<b>OUTLET</b><br><span style='font-size:10px'>{P_out:.3f} bara</span>",
                showarrow=False, xanchor="center",
                font=dict(size=12, color=_OUT_COLOR),
                bgcolor="rgba(255,255,255,0.85)", borderpad=3)
            _cap(x_offset + 0, y_top, label_below=False)
            _cap(x_offset + n, y_top, label_below=False)
            _cap(x_offset + 0, y_bot, label_below=True)
            _cap(x_offset + n, y_bot, label_below=True)
            _aa(x=x_offset + n * 0.25, y=y_top, text="←", showarrow=False, yshift=14,
                font=dict(size=14, color=_HEADER_COLOR))
            _aa(x=x_offset + n * 0.75, y=y_top, text="→", showarrow=False, yshift=14,
                font=dict(size=14, color=_HEADER_COLOR))
            _aa(x=x_offset + n * 0.25, y=y_bot, text="→", showarrow=False, yshift=-16,
                font=dict(size=14, color=b_color))
            _aa(x=x_offset + n * 0.75, y=y_bot, text="←", showarrow=False, yshift=-16,
                font=dict(size=14, color=b_color))

        elif is_di:
            _inlet_stub(x_offset + 0, -1, y_top)
            _at(go.Scatter(x=[x_offset + n, x_offset + n + stub], y=[y_top, y_top],
                           mode="lines", line=dict(color=_IN_COLOR, width=5),
                           hoverinfo="skip", showlegend=False))
            _at(go.Scatter(x=[x_offset + n + stub], y=[y_top], mode="markers",
                           marker=dict(size=16, color=_IN_COLOR, symbol="arrow-right",
                                       angle=180, line=dict(color="white", width=1)),
                           hoverinfo="skip", showlegend=False))
            _aa(x=x_offset + n + stub, y=y_top, text="<b>INLET</b>",
                showarrow=False, xanchor="left", xshift=6,
                font=dict(size=12, color=_IN_COLOR),
                bgcolor="rgba(255,255,255,0.85)", borderpad=3)
            _outlet_stub(x_offset + n, +1, y_bot)
            _cap(x_offset + 0, y_bot, label_below=True)
            _aa(x=x_offset + n * 0.25, y=y_top, text="→", showarrow=False, yshift=14,
                font=dict(size=14, color=_HEADER_COLOR))
            _aa(x=x_offset + n * 0.75, y=y_top, text="←", showarrow=False, yshift=14,
                font=dict(size=14, color=_HEADER_COLOR))
            _aa(x=mid_x, y=y_bot, text="→", showarrow=False, yshift=-16,
                font=dict(size=14, color=b_color))

        else:
            # Z or U
            if not hide_inlet:
                _inlet_stub(x_offset, -1, y_top)
            out_x   = x_offset + (0 if is_u else n)
            out_dir = -1 if is_u else +1
            if not hide_outlet:
                _outlet_stub(out_x, out_dir, y_bot)
            _cap(x_offset + n, y_top, label_below=False)
            _cap(x_offset + (n if is_u else 0), y_bot, label_below=True)
            _aa(x=mid_x, y=y_top, text="→", showarrow=False, yshift=14,
                font=dict(size=14, color=_HEADER_COLOR))
            b_arrow = "←" if is_u else "→"
            _aa(x=mid_x, y=y_bot, text=b_arrow, showarrow=False, yshift=-16,
                font=dict(size=14, color=b_color))

        # ── Badge ─────────────────────────────────────────────────────────────
        badges = {"Z": "Z-manifold", "U": "U-manifold",
                  "I": "I-manifold (center-fed)", "DI": "Double-inlet Z"}
        badge = badges.get(topo.harp_type, topo.harp_type)
        _aa(x=x_offset + n / 2, y=y_top + 0.14,
            text=f"<b>{topo.harp_id}</b>  {badge}",
            showarrow=False, font=dict(size=10, color="#334155"))

    # ── Draw harps ────────────────────────────────────────────────────────────
    if series:
        _draw(topo1, rep1, x_offset=0.0, hide_outlet=True,  row=1)
        _draw(topo2, rep2, x_offset=0.0, hide_inlet=True,   row=2)
        # Connector label floats in the spacing band between subplots
        fig.add_annotation(
            x=0.5, y=0.5, xref="paper", yref="paper",
            text=f"↕  connector",
            showarrow=False,
            font=dict(size=10, color=_CONN_COLOR),
            bgcolor="rgba(255,255,255,0.85)", borderpad=3,
        )
    else:
        _draw(topo1, rep1, x_offset=0.0)

    # ── Legend ────────────────────────────────────────────────────────────────
    ok_lbl   = "OK" if not single_phase else "OK (≥90% mean)"
    warn_lbl = "Low V_sl" if not single_phase else "70–90% mean"
    star_lbl = "Starved" if not single_phase else "<70% mean"
    for label, color in [(ok_lbl, _OK_COLOR), (warn_lbl, _WARN_COLOR), (star_lbl, _STAR_COLOR)]:
        fig.add_trace(go.Scatter(x=[None], y=[None], mode="lines",
                                  line=dict(color=color, width=4),
                                  name=label, showlegend=True))

    _ax_kw = dict(showgrid=False, zeroline=False, showticklabels=False)
    _y_kw  = dict(**_ax_kw, range=[-0.55, 1.65])
    layout_kw: dict = dict(
        template="plotly_white",
        height=640 if series else 380,
        margin=dict(l=60, r=30, t=35, b=55),
        yaxis=_y_kw, xaxis=_ax_kw,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        hoverlabel=dict(bgcolor="white", font_size=11),
    )
    if series:
        layout_kw["yaxis2"] = _y_kw
        layout_kw["xaxis2"] = _ax_kw
    fig.update_layout(**layout_kw)
    return fig


def _make_flow_bar_chart(
    reports: list[tuple[StarvationReport, str]],
    Vsl_threshold: float,
    single_phase: bool,
    max_bars: int = 300,
) -> go.Figure:
    fig = go.Figure()
    for report, label in reports:
        mean_m = report.mean_m_kgs
        rs = report.channel_results
        # Downsample if too many channels
        if len(rs) > max_bars:
            step = max(1, len(rs) // max_bars)
            rs   = rs[::step]
        colors = [_ch_color(r, Vsl_threshold, single_phase, mean_m) for r in rs]
        fig.add_trace(go.Bar(
            x=[f"ch{r.channel_index}" for r in rs],
            y=[r.m_kgs * 1000 for r in rs],
            name=label, marker_color=colors,
            hovertemplate="%{x}<br>%{y:.2f} g/s<extra></extra>",
        ))
        fig.add_hline(y=mean_m * 1000, line_dash="dash",
                      line_color="#64748b", line_width=1,
                      annotation_text=f"{label} mean {mean_m*1000:.1f} g/s",
                      annotation_position="top right", annotation_font_size=10)
    fig.update_layout(
        template="plotly_white", height=260,
        margin=dict(l=50, r=20, t=30, b=40),
        yaxis_title="Flow (g/s)", xaxis_title="Channel",
        barmode="group",
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
    )
    return fig


def _make_pressure_profile(
    solve_res: SolveResult,
    topo1: HarpTopology,
    topo2: HarpTopology | None,
) -> go.Figure:
    net = solve_res.net
    fig = go.Figure()

    def _add(topo: HarpTopology, lbl: str, cA: str, cB: str, x_off: int = 0) -> None:
        n   = len(topo.header_A_node_ids)
        p_A = [net.node(nid).P_pa / 1e5 for nid in topo.header_A_node_ids]
        p_B = [net.node(nid).P_pa / 1e5 for nid in topo.header_B_node_ids]
        xs  = [x_off + i for i in range(n)]
        fig.add_trace(go.Scatter(x=xs, y=p_A, mode="lines+markers",
                                  name=f"{lbl} A", line=dict(color=cA, width=2),
                                  marker=dict(size=4)))
        fig.add_trace(go.Scatter(x=xs, y=p_B, mode="lines+markers",
                                  name=f"{lbl} B", line=dict(color=cB, width=2, dash="dot"),
                                  marker=dict(size=4)))

    _add(topo1, "H1", "#2563EB", "#93c5fd")
    if topo2:
        _add(topo2, "H2", "#d97706", "#fcd34d",
             x_off=len(topo1.header_A_node_ids) + 2)

    fig.update_layout(
        template="plotly_white", height=230,
        margin=dict(l=50, r=20, t=30, b=40),
        yaxis_title="Pressure (bara)", xaxis_title="Node position",
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
    )
    return fig


# ── Engineering report ────────────────────────────────────────────────────────

def _generate_harp_report(
    solve_res: SolveResult,
    topo1: HarpTopology,
    topo2: HarpTopology | None,
    rep1: StarvationReport,
    rep2: StarvationReport | None,
    inputs: dict,
    fig_diagram: "go.Figure",
    fig_bar: "go.Figure",
    fig_pressure: "go.Figure",
) -> bytes:
    """
    Generate a Word (.docx) engineering report for the harp solve result.
    Returns bytes suitable for st.download_button.
    """
    import re
    from io import BytesIO
    from datetime import datetime

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ── Derived values ─────────────────────────────────────────────────────────
    single_phase = inputs.get("single_phase", False)
    k            = inputs.get("channels_per_branch", 1)
    N            = inputs.get("N", 0)
    c_D          = inputs.get("c_D", 0.025)
    A_ch         = math.pi * c_D ** 2 / 4.0
    solver_name  = inputs.get("solver_method", "GGA (Newton-Raphson)")
    solve_mode   = inputs.get("solve_mode", "Specify flow")
    found_m_kgs  = inputs.get("found_m_kgs", inputs.get("m_total", 0))
    P_in_result  = solve_res.net.node(topo1.inlet_node_id).P_pa / 1e5
    P_out        = inputs.get("P_outlet", 0.0)
    dP_bar       = P_in_result - P_out
    dP_str       = f"{dP_bar*1000:.1f} mbar" if dP_bar < 1 else f"{dP_bar:.4f} bar"

    # ── Helper: write Markdown subset into docx ────────────────────────────────
    def _bold_runs(para, text: str) -> None:
        """Add runs to *para*, rendering **bold** and *italic* markers."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                para.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                para.add_run(part[1:-1]).italic = True
            elif part:
                para.add_run(part)

    def _md_section(md_text: str) -> None:
        """Render the expert-review Markdown into the open document."""
        for line in md_text.split('\n'):
            s = line.strip()
            if not s:
                continue
            if s.startswith('#### '):
                doc.add_heading(s[5:], level=3)
            elif s == '---':
                doc.add_paragraph('─' * 72).style = 'Normal'
            elif re.match(r'^\d+\.\s', s):
                p = doc.add_paragraph(style='List Number')
                _bold_runs(p, re.sub(r'^\d+\.\s+', '', s))
            elif s.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                _bold_runs(p, s[2:])
            else:
                p = doc.add_paragraph()
                _bold_runs(p, s)

    # ── Document skeleton ──────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Harp Manifold — Network Solver Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d  %H:%M')}   |   "
        f"Solver: {solver_name}   |   "
        f"Converged: {'Yes' if solve_res.converged else 'No'}  "
        f"({solve_res.iterations} iterations, residual = {solve_res.residual:.2e})"
    ).runs[0].font.size = Pt(9)

    # ── 1  Configuration ───────────────────────────────────────────────────────
    doc.add_heading("1  Configuration", level=1)
    cfg_table = doc.add_table(rows=1, cols=2)
    cfg_table.style = "Table Grid"
    hdr = cfg_table.rows[0].cells
    hdr[0].text = "Parameter"; hdr[1].text = "Value"

    _type_str = inputs.get("harp_type_str", "Z")
    _type_labels = {"Z": "Z-manifold", "U": "U-manifold (reverse-return)",
                    "I": "I-manifold (center-fed)", "DI": "Double-inlet Z"}

    _total_ch = N * k * (2 if topo2 else 1)
    _ch_desc  = (f"{N} branches × {k} ch/branch = {N*k} physical channels per harp"
                 if k > 1 else f"{N} channels per harp")

    cfg_rows = [
        ("Manifold type",      _type_labels.get(_type_str, _type_str)),
        ("Harps in series",    "Yes" if inputs.get("series_mode") else "No"),
        ("Channels per harp",  _ch_desc),
        ("Header ID",          f"{inputs.get('h_D', 0)*1000:.1f} mm"),
        ("Channel ID (physical)", f"{c_D*1000:.1f} mm"),
        ("Channel length",     f"{inputs.get('c_len', 0):.3f} m"),
        ("Header seg length",  f"{inputs.get('h_len', 0):.3f} m"),
        ("Liquid",             inputs.get("liq_type", "—")),
        ("Gas species",        inputs.get("gas_lbl", "—")),
        ("Solve mode",         solve_mode),
        ("Solver method",      solver_name),
        ("Total flow",         f"{found_m_kgs*3600:.1f} kg/h"),
        ("x_gas (inlet)",      f"{inputs.get('x_inlet', 0):.4f}"),
        ("P_outlet (pinned)",  f"{P_out:.2f} bara"),
        ("P_inlet (result)",   f"{P_in_result:.4f} bara"),
        ("ΔP (total)",         dP_str),
        ("T",                  f"{inputs.get('T_C', 0):.0f} °C"),
        ("Starvation threshold", f"{inputs.get('Vsl_threshold', 0.05):.4f} m/s"),
    ]
    for param, val in cfg_rows:
        r = cfg_table.add_row().cells
        r[0].text = param; r[1].text = str(val)

    # ── 2  Network diagram ─────────────────────────────────────────────────────
    doc.add_heading("2  Network diagram", level=1)
    try:
        import plotly.io as pio
        _diag_h = int(getattr(fig_diagram.layout, "height", None) or (640 if topo2 else 380))
        buf_diag = BytesIO(pio.to_image(fig_diagram, format="png",
                                         width=900, height=_diag_h, scale=2))
        doc.add_picture(buf_diag, width=Inches(6.5))
    except Exception:
        doc.add_paragraph("[Diagram could not be embedded — kaleido may not be installed]")

    # ── 3  Flow distribution ───────────────────────────────────────────────────
    doc.add_heading("3  Flow distribution", level=1)
    try:
        buf_bar = BytesIO(pio.to_image(fig_bar, format="png", width=900, height=280, scale=2))
        doc.add_picture(buf_bar, width=Inches(6.5))
    except Exception:
        doc.add_paragraph("[Chart could not be embedded]")

    # ── 4  Header pressure profile ─────────────────────────────────────────────
    doc.add_heading("4  Header pressure profile", level=1)
    try:
        buf_pres = BytesIO(pio.to_image(fig_pressure, format="png", width=700, height=250, scale=2))
        doc.add_picture(buf_pres, width=Inches(6.0))
    except Exception:
        doc.add_paragraph("[Chart could not be embedded]")

    # ── 5  Summary ─────────────────────────────────────────────────────────────
    doc.add_heading("5  Summary", level=1)
    for rep, label in [(rep1, "Harp 1")] + ([(rep2, "Harp 2")] if rep2 else []):
        doc.add_heading(label, level=2)
        mdi    = rep.maldistribution_index
        grade  = _er_grade(mdi).replace("✅", "").replace("🟡", "").replace(
                     "🟠", "").replace("🔴", "").strip()
        mean_ch_gs = rep.mean_m_kgs * 1000 / k
        p = doc.add_paragraph()
        p.add_run(f"Grade: {grade}   |   MDI = {mdi:.4f}   |   "
                  f"Low-flow branches: {rep.n_channels_starved}/{rep.n_channels_total}   |   "
                  f"Mean branch flow: {rep.mean_m_kgs*1000:.2f} g/s")
        if k > 1:
            p.add_run(f"   |   Mean per physical channel: {mean_ch_gs:.2f} g/s")

    # ── 6  Per-channel results ─────────────────────────────────────────────────
    doc.add_heading("6  Per-branch / per-channel results", level=1)
    if k > 1:
        doc.add_paragraph(
            f"Values shown per physical channel (branch flow ÷ {k}). "
            f"V_sl/ch computed with physical channel ID = {c_D*1000:.1f} mm."
        ).runs[0].font.size = Pt(9)

    def _add_ch_table(rep: StarvationReport, label: str) -> None:
        doc.add_heading(label, level=2)
        _m_hdr  = "m/ch (g/s)" if k > 1 else "m (g/s)"
        _v_hdr  = "V_sl/ch (m/s)" if k > 1 else "V_sl (m/s)"
        headers = ["Branch" if k > 1 else "Ch", _m_hdr, _v_hdr]
        if not single_phase:
            headers += ["V_sg (m/s)", "x_gas", "Regime"]
        headers += ["Status"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
        for r in rep.channel_results:
            m_ch  = r.m_kgs / k
            vsl_ch = m_ch / max(1000.0 * A_ch, 1e-12)
            row = tbl.add_row().cells
            row[0].text = str(r.channel_index)
            row[1].text = f"{m_ch*1000:.2f}"
            row[2].text = f"{vsl_ch:.4f}"
            col = 3
            if not single_phase:
                row[col].text = f"{r.Vsg:.4f}"; col += 1
                row[col].text = f"{r.x_gas:.4f}"; col += 1
                row[col].text = r.regime; col += 1
            status = "LOW" if r.starved else "OK"
            row[col].text = status
            if r.starved:
                row[col].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    _add_ch_table(rep1, "Harp 1")
    if rep2:
        _add_ch_table(rep2, "Harp 2")

    # ── 7  Expert review ───────────────────────────────────────────────────────
    doc.add_heading("7  Expert Review", level=1)
    _review_md = _expert_review(
        solve_res, topo1, topo2, rep1, rep2,
        N=N, h_D=inputs.get("h_D", 0.05),
        c_D=c_D, c_len=inputs.get("c_len", 0.5),
        h_len=inputs.get("h_len", 0.15),
        P_outlet=P_out, single_phase=single_phase,
        channels_per_branch=k,
    )
    _md_section(_review_md)

    # ── 8  Method ──────────────────────────────────────────────────────────────
    doc.add_heading("8  Method", level=1)
    _is_hc = "Hardy" in solver_name
    if _is_hc:
        doc.add_paragraph(
            "Solver: Hardy-Cross (1936) loop-balancing method. "
            "Sequential flow correction applied to each independent loop until "
            "max|ΔQ|/m_total converges. "
            f"Converged in {solve_res.iterations} iterations "
            f"(residual = {solve_res.residual:.2e})."
        )
    else:
        doc.add_paragraph(
            "Solver: Global Gradient Algorithm (GGA), Todini & Pilati (1988). "
            "Newton-Raphson simultaneous solution for nodal pressures and pipe flows. "
            f"Converged in {solve_res.iterations} iterations "
            f"(residual = {solve_res.residual:.2e})."
        )
    doc.add_paragraph(
        f"Two-phase correlation: {inputs.get('corr', 'Beggs-Brill')}. "
        f"Void-fraction model: {inputs.get('void', 'Homogeneous')}. "
        "T-junction K-factors: K_fwd = 0.5 (dividing T), K_rev = 0.3 (combining T, Idelchik §7-23). "
        "Phase split: homogeneous (equal quality at every junction). "
        "Momentum recovery in header: not modelled (conservative)."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Expert review helpers ──────────────────────────────────────────────────────

def _er_grade(mdi: float) -> str:
    if mdi < 0.05:  return "✅ Excellent"
    if mdi < 0.20:  return "🟡 Acceptable"
    if mdi < 0.50:  return "🟠 Poor"
    return "🔴 Critical"


def _er_pattern(
    channel_results: list,
    harp_type: str,
) -> tuple:
    """Return (label, pearson_r, def_first, def_last, def_centre)."""
    flows = [r.m_kgs for r in channel_results]
    n = len(flows)
    if n < 2:
        return ("uniform", 0.0, 0.0, 0.0, 0.0)

    mean_f = sum(flows) / n
    if mean_f < 1e-12:
        return ("uniform", 0.0, 0.0, 0.0, 0.0)

    std_f = math.sqrt(sum((f - mean_f) ** 2 for f in flows) / n)
    if std_f < 1e-12 * mean_f:
        return ("uniform", 0.0, 0.0, 0.0, 0.0)

    third = max(1, n // 3)
    first_mean  = sum(flows[:third]) / third
    last_mean   = sum(flows[n - third:]) / third
    mid_slice   = flows[third: n - third]
    centre_mean = sum(mid_slice) / max(len(mid_slice), 1) if mid_slice else mean_f

    def_first  = (mean_f - first_mean)  / mean_f
    def_last   = (mean_f - last_mean)   / mean_f
    def_centre = (mean_f - centre_mean) / mean_f

    xs = list(range(n))
    mx = (n - 1) / 2.0
    std_x = math.sqrt(sum((x - mx) ** 2 for x in xs) / n)
    cov   = sum((xs[i] - mx) * (flows[i] - mean_f) for i in range(n)) / n
    r     = cov / max(std_x * std_f, 1e-12)

    if harp_type == "I":
        # Centre-fed: centre should be high, ends low → "end-starved" is normal
        if max(abs(def_first), abs(def_last)) > 0.10:
            return ("end-starved", r, def_first, def_last, def_centre)
        return ("uniform", r, def_first, def_last, def_centre)

    if abs(r) < 0.3 and max(abs(def_first), abs(def_last), abs(def_centre)) < 0.10:
        return ("uniform", r, def_first, def_last, def_centre)

    if harp_type in ("Z", "DI"):
        if r > 0.4:    return ("inlet-starved",  r, def_first, def_last, def_centre)
        if r < -0.4:   return ("outlet-starved", r, def_first, def_last, def_centre)
    elif harp_type == "U":
        if def_first > 0.10 and def_last > 0.10:
            return ("end-starved",    r, def_first, def_last, def_centre)
        if def_centre > 0.10:
            return ("centre-starved", r, def_first, def_last, def_centre)

    return ("mixed", r, def_first, def_last, def_centre)


def _er_header_ratio(net, topo) -> tuple:
    """Return (R, dP_header_Pa, dP_ch_mean_Pa)."""
    a_ids = topo.header_A_node_ids
    if topo.harp_type in ("I", "DI"):
        centre = len(a_ids) // 2
        dP_hdr = abs(net.node(a_ids[centre]).P_pa - net.node(a_ids[0]).P_pa)
    else:
        dP_hdr = abs(net.node(a_ids[0]).P_pa - net.node(a_ids[-1]).P_pa)

    ch_dps     = [abs(net.edge(eid).dP_Pa) for eid in topo.channel_edge_ids]
    dP_ch_mean = sum(ch_dps) / max(len(ch_dps), 1)
    R          = dP_hdr / max(dP_ch_mean, 1.0)
    return R, dP_hdr, dP_ch_mean


def _er_type_recs(harp_type: str, mdi: float, N: int) -> list:
    recs = []
    if harp_type == "Z":
        if mdi > 0.20 and N < 10:
            recs.append("Switch to **Double-inlet Z**: feeds both header ends simultaneously, "
                        "halving the header pressure gradient.")
        elif mdi > 0.20 and N >= 10:
            recs.append("Switch to **U-manifold** (reverse-return): header pressure gradients "
                        "partially cancel, typically halving MDI for this geometry.")
        if mdi > 0.30 and N >= 20:
            recs.append("Or consider **I-manifold** (center-fed): inlet at header midpoint "
                        "makes distribution symmetric — best uniformity for N > 20.")
    elif harp_type == "U":
        if mdi > 0.30:
            recs.append("Consider **I-manifold** (center-fed) or **Double-inlet Z** "
                        "for better uniformity at this channel count.")
    elif harp_type in ("I", "DI"):
        if mdi > 0.30:
            recs.append("Distribution topology is already near-optimal; "
                        "focus on increasing the header ID (see above).")
    return recs


def _er_channel_count_rec(N: int, R: float, mdi: float):
    if mdi > 0.20 and R > 0.20:
        N_crit = int(N * math.sqrt(0.20 / R))
        if N_crit < int(0.70 * N) and N_crit >= 2:
            return N_crit
    return None


def _er_two_phase(channel_results: list, net, ch_edge_ids: list) -> list:
    notes = []
    x_vals = [r.x_gas for r in channel_results]
    mean_x = sum(x_vals) / max(len(x_vals), 1)

    if mean_x > 1e-4:
        std_x = math.sqrt(sum((x - mean_x) ** 2 for x in x_vals) / max(len(x_vals), 1))
        cv    = std_x / max(mean_x, 1e-12)
        if cv > 0.15:
            notes.append(
                f"Gas quality varies across channels (CV = {cv:.0%}) — "
                "phase maldistribution: gas-rich channels may trend toward dry-out "
                "or flow reversal at higher total flow."
            )

    n_rev = sum(1 for eid in ch_edge_ids if net.edge(eid).m_kgs < -1e-8)
    if n_rev:
        notes.append(
            f"⚠ **{n_rev} channel(s) have reversed flow** — liquid is being pushed "
            "backward by the header pressure gradient. "
            "Increase total flow or widen the header."
        )

    from collections import Counter
    regimes = [r.regime for r in channel_results if r.regime]
    if regimes:
        counts = Counter(regimes)
        majority, maj_count = counts.most_common(1)[0]
        n_mixed = len(regimes) - maj_count
        if n_mixed > 0:
            minority = [rg for rg, _ in counts.most_common() if rg != majority]
            notes.append(
                f"Mixed flow regimes: {n_mixed} channel(s) in "
                f"**{', '.join(minority)}**, majority in **{majority}** — "
                "channels are near a regime-transition boundary."
            )

    return notes


def _er_series_note(solve_res, topo1, topo2, rep1, rep2) -> list:
    lines = []
    P_h1_out = solve_res.net.node(topo1.outlet_node_id).P_pa / 1e5
    P_h2_in  = solve_res.net.node(topo2.inlet_node_id).P_pa  / 1e5
    dP_conn  = (P_h1_out - P_h2_in) * 1e5
    lines.append(
        f"Connector ΔP = **{dP_conn:.0f} Pa** "
        f"({P_h1_out:.4f} → {P_h2_in:.4f} bara)"
    )
    mdi1, mdi2 = rep1.maldistribution_index, rep2.maldistribution_index
    if mdi2 > mdi1 * 1.5 and mdi1 > 0.02:
        lines.append(
            f"Second harp shows significantly worse distribution "
            f"(MDI H2 = {mdi2:.3f} vs H1 = {mdi1:.3f}) — "
            "the connector pipe pressure drop may be insufficient to equalise "
            "inlet conditions. Consider increasing connector pipe ID."
        )
    return lines


def _expert_review(
    solve_res: SolveResult,
    topo1: HarpTopology,
    topo2: HarpTopology | None,
    rep1: StarvationReport,
    rep2: StarvationReport | None,
    N: int,
    h_D: float,
    c_D: float,
    c_len: float,
    h_len: float,
    P_outlet: float,
    single_phase: bool,
    channels_per_branch: int = 1,
) -> str:
    net = solve_res.net

    if not solve_res.converged:
        return "*Solver did not converge — expert review unavailable.*"
    if rep1.n_channels_total < 1:
        return "*No channel data available.*"

    k = channels_per_branch

    def _harp_section(topo: HarpTopology, rep: StarvationReport, label: str) -> str:
        lines = []
        mdi   = rep.maldistribution_index
        grade = _er_grade(mdi)
        n_ch  = rep.n_channels_total
        mean_per_ch = rep.mean_m_kgs / k

        ch_desc = (f"{n_ch} branches × {k} ch = {n_ch*k} physical channels"
                   if k > 1 else f"{n_ch} channels")
        lines.append(f"#### {label}")
        lines.append(
            f"**Grade:** {grade} &nbsp;·&nbsp; MDI = {mdi:.3f} &nbsp;·&nbsp; "
            f"{ch_desc} &nbsp;·&nbsp; "
            f"mean {'branch' if k>1 else 'channel'} flow = {rep.mean_m_kgs*1000:.1f} g/s"
            + (f" ({mean_per_ch*1000:.1f} g/s per physical channel)" if k > 1 else "")
        )

        # ── Pattern ──────────────────────────────────────────────────────────
        if n_ch >= 2:
            pattern, r, def_first, def_last, def_centre = _er_pattern(
                rep.channel_results, topo.harp_type)
            third = max(1, n_ch // 3)

            if pattern == "uniform":
                lines.append("**Pattern:** Uniform distribution — no significant positional bias.")
            elif pattern == "inlet-starved":
                lines.append(
                    f"**Pattern:** Near-inlet channels (ch0–ch{third-1}) receive "
                    f"**{def_first:.0%} less** flow than far-end channels — "
                    "classic Z-manifold header-momentum pattern."
                )
            elif pattern == "outlet-starved":
                lines.append(
                    f"**Pattern:** Far-end channels (ch{n_ch-third}–ch{n_ch-1}) receive "
                    f"**{def_last:.0%} less** flow than near-inlet channels."
                )
            elif pattern == "end-starved":
                lines.append(
                    f"**Pattern:** End channels receive less flow than central channels "
                    f"(first-third deficit {def_first:.0%}, last-third deficit {def_last:.0%}) — "
                    "consistent with U-manifold reverse-return behaviour."
                )
            elif pattern == "centre-starved":
                lines.append(
                    f"**Pattern:** Central channels receive **{def_centre:.0%} less** flow "
                    "than end channels."
                )
            elif pattern == "end-starved" and topo.harp_type == "I":
                lines.append(
                    f"**Pattern:** End channels receive less flow than the centre-fed "
                    f"region (end deficit {max(def_first, def_last):.0%}) — "
                    "expected for I-manifold; reduce only if unacceptable."
                )
            else:
                lines.append(
                    f"**Pattern:** Irregular distribution (Pearson r = {r:.2f}) — "
                    "no single dominant positional trend."
                )

        # ── Root cause ────────────────────────────────────────────────────────
        R, dP_hdr, dP_ch = _er_header_ratio(net, topo)
        dP_hdr_mbar = dP_hdr / 100.0
        dP_ch_mbar  = dP_ch  / 100.0
        if R > 0.50:
            cause_label = "Header-dominated"
        elif R > 0.20:
            cause_label = "Moderate header effect"
        else:
            cause_label = "Channel-dominated (uniform by design)"

        lines.append(
            f"**Root cause:** Supply header ΔP = {dP_hdr_mbar:.1f} mbar = "
            f"**{R:.2f}×** channel ΔP ({dP_ch_mbar:.1f} mbar) — {cause_label}."
        )

        # ── Recommendations ───────────────────────────────────────────────────
        recs = []

        if R > 0.20:
            D_rec = h_D * (R / 0.10) ** 0.25
            recs.append(
                f"🔧 Increase header ID from **{h_D*1000:.0f} mm** to "
                f"**{D_rec*1000:.0f} mm** — reduces header/channel ΔP ratio "
                f"from {R:.2f} to 0.10."
            )

        for tr in _er_type_recs(topo.harp_type, mdi, n_ch):
            emoji = "🔄" if "U-manifold" in tr or "Double-inlet" in tr else "⚙️"
            recs.append(f"{emoji} {tr}")

        N_crit = _er_channel_count_rec(n_ch, R, mdi)
        if N_crit is not None:
            recs.append(
                f"⚙️ Reduce channel count from **{n_ch}** to **{N_crit}** "
                "to lower per-channel header velocity and reduce header ΔP."
            )

        if recs:
            lines.append("**Recommendations:**")
            for i, rec in enumerate(recs, 1):
                lines.append(f"{i}. {rec}")
        elif mdi < 0.05:
            lines.append("No geometry changes needed — distribution is within ±5% of mean.")

        # ── Two-phase notes ───────────────────────────────────────────────────
        if not single_phase:
            tp = _er_two_phase(rep.channel_results, net, topo.channel_edge_ids)
            if tp:
                lines.append("\n*Two-phase notes:*")
                for note in tp:
                    lines.append(f"- {note}")

        return "\n\n".join(lines)

    parts = [_harp_section(topo1, rep1, "Harp 1")]

    if topo2 is not None and rep2 is not None:
        parts.append("---")
        parts.append(_harp_section(topo2, rep2, "Harp 2"))
        series_lines = _er_series_note(solve_res, topo1, topo2, rep1, rep2)
        if series_lines:
            parts.append("---")
            parts.append("#### Series connector\n\n" + "\n\n".join(series_lines))

    return "\n\n".join(parts)


# ── Main render ────────────────────────────────────────────────────────────────

def render_harp_network_tab() -> None:
    # Apply any pending flow value from the per-channel calculator (must happen
    # before the liq_kgh widget is instantiated — Streamlit forbids setting a
    # widget key after the widget has been rendered).
    st.markdown(
        "Solve **flow distribution** in harp manifolds. "
        "Detects channels starved of liquid due to header pressure maldistribution."
    )

    col_in, col_out = st.columns([1, 2], gap="large")

    with col_in:
        st.markdown("#### Network geometry")

        gc1, gc2 = st.columns(2)
        series_mode = gc1.toggle("Two harps in series", key=_k("series"),
                                  value=st.session_state.get(_k("series"), False))
        harp_type   = gc2.radio(
            "Manifold type",
            ["Z-manifold", "U-manifold", "I-manifold", "Double-inlet Z"],
            key=_k("harp_type"), horizontal=False,
            help=(
                "**Z-manifold**: inlet left, outlet right (opposite ends).\n\n"
                "**U-manifold**: inlet and outlet on the **same** end "
                "(reverse-return).\n\n"
                "**I-manifold**: inlet at centre of supply header, outlet at "
                "centre of collection header — symmetric distribution, "
                "best uniformity for large N.\n\n"
                "**Double-inlet Z**: feed enters at **both** ends of the supply "
                "header simultaneously — halves header pressure gradient, "
                "much more uniform than single-inlet Z."
            ),
            index=0,
        )
        _type_map = {
            "Z-manifold": "Z", "U-manifold": "U",
            "I-manifold": "I", "Double-inlet Z": "DI",
        }
        harp_type_str = _type_map[harp_type]

        # Series mode only makes sense for Z and U (I/DI can be added later)
        if harp_type_str in ("I", "DI") and series_mode:
            st.caption("ℹ Series mode is not yet supported for I / Double-inlet. "
                       "Solving as single harp.")

        N = st.slider("Channels per harp", 2, 500,
                      value=st.session_state.get(_k("N"), 8), key=_k("N"))

        h_D, h_eps, h_len = _pipe_section(
            "Header pipe", "hn_h",
            def_dn="DN50", def_pn="PN40", def_len_m=0.15,
            def_id_mm=50.0, def_rough_mm=0.046, expanded=True,
            length_label="Spacing between branches",
            length_help=(
                "Distance between consecutive branch tap points along the header. "
                "Total header length = N × this value. "
                "Example: 10 branches × 150 mm = 1.5 m total header."
            ),
        )
        c_D, c_eps, c_len, c_angle = _pipe_section(
            "Channel pipe", "hn_c",
            def_dn="DN25", def_pn="PN40", def_len_m=0.50,
            def_id_mm=25.0, def_rough_mm=0.046, show_angle=True, expanded=True,
            length_label="Channel length",
        )
        # Rectangular channel override — uses hydraulic diameter Dh = 2WH/(W+H)
        with st.expander("Rectangular channel (optional)", expanded=False):
            use_rect = st.toggle("Use rectangular cross-section",
                                  key=_k("use_rect"),
                                  value=st.session_state.get(_k("use_rect"), False),
                                  help="Overrides the channel diameter with Dh = 2WH/(W+H).")
            if use_rect:
                rc1, rc2 = st.columns(2)
                rect_w = rc1.number_input("Width W (mm)", min_value=0.1, max_value=500.0,
                                           value=float(st.session_state.get(_k("rect_w"), 10.0)),
                                           step=0.5, format="%.2f", key=_k("rect_w"))
                rect_h = rc2.number_input("Height H (mm)", min_value=0.1, max_value=500.0,
                                           value=float(st.session_state.get(_k("rect_h"), 5.0)),
                                           step=0.5, format="%.2f", key=_k("rect_h"))
                W_m, H_m = rect_w / 1000.0, rect_h / 1000.0
                Dh = 2 * W_m * H_m / (W_m + H_m)   # hydraulic diameter
                A_rect = W_m * H_m
                st.caption(f"D_h = {Dh*1000:.2f} mm  |  A = {A_rect*1e6:.2f} mm²"
                           f"  |  Aspect ratio = {max(rect_w,rect_h)/min(rect_w,rect_h):.2f}")
                c_D = Dh  # override diameter used in builder
            else:
                rect_w = rect_h = 0.0

        # Parallel channels per branch
        channels_per_branch = int(st.number_input(
            "Channels per branch", min_value=1, max_value=8,
            value=int(st.session_state.get(_k("cpb"), 1)), step=1,
            key=_k("cpb"),
            help="Number of parallel, identical channels at each header tap. "
                 "2 = two channels side-by-side at every branch point. "
                 "The solver uses an equivalent diameter; results are shown "
                 "per physical channel.",
        ))
        # Equivalent diameter for k parallel channels (turbulent Darcy-Weisbach):
        # resistance R ∝ 1/D^5; k parallel pipes → R_eq = R/k²  → D_eq = D·k^0.4
        c_D_eff = c_D * (channels_per_branch ** 0.4) if channels_per_branch > 1 else c_D
        if channels_per_branch > 1:
            st.caption(
                f"{channels_per_branch} ch/branch → equivalent channel ID "
                f"{c_D_eff*1000:.1f} mm  (physical ID {c_D*1000:.1f} mm)"
            )
        if series_mode:
            k_D, k_eps, k_len = _pipe_section(
                "Connector pipe", "hn_k",
                def_dn="DN50", def_pn="PN40", def_len_m=0.30,
                def_id_mm=50.0, def_rough_mm=0.046, expanded=True,
            )
        else:
            k_D = k_eps = 0.05; k_len = 0.30

        # ── Fluid ─────────────────────────────────────────────────────────────
        st.markdown("#### Fluid")
        preset_name = st.selectbox("Preset", list(_PRESETS.keys()), key=_k("preset"))
        preset = _PRESETS[preset_name]

        single_phase = (preset["phase"] == "single-phase") or st.toggle(
            "Single-phase liquid only",
            value=st.session_state.get(_k("single_phase"), preset["phase"] == "single-phase"),
            key=_k("single_phase"),
        )

        fc1, fc2 = st.columns(2)
        T_C      = fc1.number_input("T (°C)", -50.0, 350.0,
                                     value=float(st.session_state.get(_k("T"), preset["T_C"])),
                                     step=1.0, format="%.0f", key=_k("T"))
        P_outlet = fc2.number_input("P_outlet (bara)", 0.01, 500.0,
                                     value=float(st.session_state.get(
                                         _k("P_outlet"), preset["P_bara"] * 0.95)),
                                     step=0.1, format="%.2f", key=_k("P_outlet"),
                                     help="Outlet pressure (pinned). P_inlet is a solver result.")

        if not single_phase:
            _gas_opts = list(_GAS_COOLPROP.keys())
            gc1, gc2 = st.columns(2)
            _gas_lbl = gc1.selectbox("Gas species", _gas_opts, key=_k("gas_species"),
                                      index=_gas_opts.index(preset["gas_species"])
                                      if preset["gas_species"] in _gas_opts else 0)
            gas_kgh  = gc2.number_input("Gas flow (kg/h)", 0.0, 10000.0,
                                         value=float(st.session_state.get(_k("gas_kgh"),
                                                                          preset["gas_kgh"])),
                                         step=0.5, format="%.1f", key=_k("gas_kgh"))
        else:
            _gas_lbl = "H₂"; gas_kgh = 0.0

        liq_type = st.selectbox("Liquid", _LIQUID_OPTIONS, key=_k("liq_type"),
                                 index=_LIQUID_OPTIONS.index(preset["liquid_type"])
                                 if preset["liquid_type"] in _LIQUID_OPTIONS else 0)

        # ── Solve mode ────────────────────────────────────────────────────────
        _MODE_FLOW   = "Specify flow"
        _MODE_DP     = "Specify ΔP"
        _MODE_PERCHAN= "Specify flow/channel"
        st.markdown("#### Solve mode")
        solve_mode = st.radio(
            "",
            [_MODE_FLOW, _MODE_DP, _MODE_PERCHAN],
            index=[_MODE_FLOW, _MODE_DP, _MODE_PERCHAN].index(
                st.session_state.get(_k("solve_mode"), _MODE_FLOW)),
            key=_k("solve_mode"),
            horizontal=True,
            label_visibility="hidden",
            help=(
                "**Specify flow** — enter total liquid flow → result: P_inlet and ΔP.\n\n"
                "**Specify ΔP** — enter desired pressure drop → solver finds the flow.\n\n"
                "**Specify flow/channel** — enter per-channel flow → result: total flow and ΔP."
            ),
        )

        _ch_A       = math.pi * c_D ** 2 / 4.0
        _rho_approx = 1000.0
        goal_seek      = False
        P_inlet_target = None

        if solve_mode == _MODE_FLOW:
            liq_kgh = st.number_input(
                f"{liq_type} flow (kg/h)", 0.0, 500_000.0,
                value=float(st.session_state.get(_k("liq_kgh"), preset["liquid_kgh"])),
                step=10.0, format="%.0f", key=_k("liq_kgh"),
            )
            P_bara = P_outlet + 1.0
            _ch_liq_Lh = (liq_kgh / 3600.0 / max(N, 1)) / _rho_approx * 1_000_000.0 * 3.6
            _ch_vsl    = (liq_kgh / 3600.0 / max(N, 1)) / (_rho_approx * _ch_A) if _ch_A > 0 else 0
            if single_phase:
                st.caption(
                    f"Total: **{liq_kgh:.0f} kg/h** — "
                    f"**{_ch_liq_Lh:.1f} L/h** per channel ({_ch_vsl:.3f} m/s)"
                )
            else:
                _ch_gas_lh = gas_kgh / max(N, 1)
                st.caption(
                    f"Total: **{gas_kgh+liq_kgh:.0f} kg/h** — "
                    f"x = **{gas_kgh/max(gas_kgh+liq_kgh,1e-12):.4f}** — "
                    f"liq **{_ch_liq_Lh:.1f} L/h/ch** + gas **{_ch_gas_lh:.2f} kg/h/ch**"
                )

        elif solve_mode == _MODE_DP:
            dp1, dp2 = st.columns([3, 1])
            _dp_unit = dp2.selectbox("", ["mbar", "bar"], key=_k("dp_unit"),
                                      label_visibility="hidden")
            _dp_val  = dp1.number_input(
                "Target ΔP",
                min_value=0.001, max_value=(500_000.0 if _dp_unit == "mbar" else 500.0),
                value=float(st.session_state.get(_k("dp_val"), 100.0)),
                step=(1.0 if _dp_unit == "mbar" else 0.01),
                format=("%.1f" if _dp_unit == "mbar" else "%.3f"),
                key=_k("dp_val"),
            )
            dP_bar         = _dp_val / 1000.0 if _dp_unit == "mbar" else _dp_val
            P_inlet_target = P_outlet + dP_bar
            P_bara         = P_inlet_target
            goal_seek      = True
            liq_kgh        = max(float(st.session_state.get(_k("liq_kgh"),
                                                             preset["liquid_kgh"])), 1.0)
            st.caption(
                f"Solver will find flow for **ΔP = {_dp_val:.1f} {_dp_unit}** "
                f"({dP_bar:.4f} bar) — P_outlet = {P_outlet:.2f} bara"
            )

        else:  # _MODE_PERCHAN
            pc1, pc2 = st.columns([3, 1])
            _ch_unit = pc2.selectbox("", ["L/h", "m/s"], key=_k("ch_unit"),
                                      label_visibility="hidden")
            _ch_label_in = "per physical channel" if channels_per_branch > 1 else "per channel"
            if _ch_unit == "L/h":
                _ch_target = pc1.number_input(
                    f"Liquid flow {_ch_label_in} (L/h)",
                    min_value=0.0, max_value=1_000_000.0,
                    value=float(st.session_state.get(_k("ch_target_lh"), 100.0)),
                    step=1.0, format="%.1f", key=_k("ch_target_lh"),
                )
                liq_kgh = _ch_target * N * channels_per_branch   # total = per-ch × branches × k
                _ch_vsl = (_ch_target / 3600.0 / 1000.0) / _ch_A if _ch_A > 0 else 0
            else:
                _ch_target = pc1.number_input(
                    f"V_sl {_ch_label_in} (m/s)",
                    min_value=0.0, max_value=20.0,
                    value=float(st.session_state.get(_k("ch_target_vsl"), 0.5)),
                    step=0.05, format="%.3f", key=_k("ch_target_vsl"),
                )
                liq_kgh = _ch_target * _ch_A * _rho_approx * N * channels_per_branch * 3600.0
                _ch_vsl = _ch_target
            P_bara = P_outlet + 1.0
            _total_ch_count = N * channels_per_branch
            st.caption(
                f"Total: **{liq_kgh:.0f} kg/h** = "
                f"**{_ch_target:.1f} {_ch_unit}** × {_total_ch_count} physical channels "
                f"(V_sl ≈ {_ch_vsl:.3f} m/s)"
            )

        m_total = (gas_kgh + liq_kgh) / 3600.0
        x_inlet = gas_kgh / max(gas_kgh + liq_kgh, 1e-12)

        # ── Starvation threshold ──────────────────────────────────────────────
        _ch_area   = math.pi * c_D ** 2 / 4.0
        _vsl_to_lh = _ch_area * 3_600_000.0
        _lh_to_vsl = 1.0 / _vsl_to_lh if _vsl_to_lh > 0 else 1.0

        tc1, tc2 = st.columns([3, 1])
        _thr_unit = tc2.selectbox("Unit", ["m/s", "L/h"], key=_k("thr_unit"),
                                   label_visibility="hidden")
        if _thr_unit == "L/h":
            _def_lh = float(st.session_state.get(_k("vsl_thr"), 0.05)) * _vsl_to_lh
            _raw_lh = tc1.number_input("Min liquid flow per channel",
                                        min_value=0.0, max_value=1_000_000.0,
                                        value=float(st.session_state.get(
                                            _k("thr_lh_val"), round(_def_lh, 3))),
                                        step=max(0.01, round(_def_lh / 10, 3)),
                                        format="%.3f", key=_k("thr_lh_val"))
            Vsl_threshold = _raw_lh * _lh_to_vsl
            st.caption(f"= {Vsl_threshold:.4f} m/s  (ID = {c_D*1000:.1f} mm)")
        else:
            Vsl_threshold = tc1.number_input("Min liquid velocity per channel",
                                              min_value=0.0, max_value=10.0,
                                              value=float(st.session_state.get(_k("vsl_thr"), 0.05)),
                                              step=0.005, format="%.4f", key=_k("vsl_thr"))
            st.caption(f"= {Vsl_threshold * _vsl_to_lh:.3f} L/h  (ID = {c_D*1000:.1f} mm)")

        # Effective threshold for equivalent pipe (V_eq = V_per_ch × k^0.2)
        Vsl_threshold_eff = Vsl_threshold * (channels_per_branch ** 0.2)

        # ── Solver settings ───────────────────────────────────────────────────
        with st.expander("Solver settings", expanded=False):
            solver_method = st.radio(
                "Solver method",
                ["GGA (Newton-Raphson)", "Hardy-Cross (classic)"],
                index=0 if st.session_state.get(_k("solver_method"), "GGA") == "GGA" else 1,
                horizontal=True,
                key=_k("solver_method"),
                help=(
                    "**GGA**: Global Gradient Algorithm — simultaneous Newton-Raphson "
                    "for all pressures and flows. Quadratic convergence, ~2–5 iterations. "
                    "Recommended for production use.\n\n"
                    "**Hardy-Cross**: Sequential loop-balancing (1936 method). "
                    "Linear convergence, ~20–100 iterations. "
                    "Useful for validation and comparison."
                ),
            )
            use_hardy_cross = solver_method.startswith("Hardy")
            corr = st.selectbox("Correlation", engine.TWO_PHASE_CORRELATIONS,
                                 index=0, key=_k("corr"))
            void = st.selectbox("Void-fraction model", engine.VOIDAGE_METHODS,
                                 index=0, key=_k("void"))
            sc1, sc2 = st.columns(2)
            _relax_default = 0.5 if use_hardy_cross else 1.0
            relax    = sc1.number_input("Relaxation", 0.1, 1.0,
                                         value=float(st.session_state.get(_k("relax"), _relax_default)),
                                         step=0.1, format="%.2f", key=_k("relax"))
            _iter_default = 300 if use_hardy_cross else 50
            max_iter = sc2.number_input("Max iterations", 10, 500,
                                         value=int(st.session_state.get(_k("max_iter"), _iter_default)),
                                         step=10, key=_k("max_iter"))

        solve_clicked = st.button("▶  Solve network", use_container_width=True, type="primary")

    # ── Run solver ────────────────────────────────────────────────────────────
    _cache_key = _k("result_cache")
    _hash_key  = _k("result_hash")

    _inputs_dict = dict(
        N=N, series_mode=series_mode, harp_type=harp_type_str,
        h_D=h_D, h_eps=h_eps, h_len=h_len,
        c_D=c_D, c_eps=c_eps, c_len=c_len, c_angle=c_angle,
        channels_per_branch=channels_per_branch,
        k_D=k_D, k_eps=k_eps, k_len=k_len,
        P_bara=P_bara, T_C=T_C, P_outlet=P_outlet,
        gas_lbl=_gas_lbl, gas_kgh=gas_kgh,
        liq_type=liq_type, liq_kgh=liq_kgh,
        single_phase=single_phase,
        use_rect=st.session_state.get(_k("use_rect"), False),
        rect_w=st.session_state.get(_k("rect_w"), 0),
        rect_h=st.session_state.get(_k("rect_h"), 0),
        corr=corr, void=void, relax=relax, max_iter=max_iter,
        solve_mode=solve_mode, goal_seek=goal_seek, P_inlet_target=P_inlet_target,
        solver_method=solver_method,
    )
    _inputs_hash = hashlib.md5(
        json.dumps(_inputs_dict, sort_keys=True).encode()
    ).hexdigest()

    need_solve = solve_clicked or (
        st.session_state.get(_hash_key) != _inputs_hash
        and _cache_key not in st.session_state
    )

    if need_solve:
        with col_out:
            with st.spinner("Solving…"):
                try:
                    gas_cp_id = _GAS_COOLPROP.get(_gas_lbl, _gas_lbl)
                    gas_flows = {} if single_phase else {gas_cp_id: gas_kgh}
                    liq_flows = {liq_type: liq_kgh}

                    _common = dict(
                        header_D_inner_m=h_D, header_roughness_m=h_eps,
                        channel_D_inner_m=c_D_eff, channel_roughness_m=c_eps,
                        header_segment_length=h_len,
                        channel_length=c_len, channel_angle_rad=c_angle,
                        correlation=corr, voidage_method=void,
                        P_inlet_pa=P_bara * 1e5, T_C=T_C, x_inlet=x_inlet,
                    )

                    effective_series = series_mode and harp_type_str in ("Z", "U")

                    if effective_series:
                        net, topo1, topo2 = build_series_harp(
                            N,
                            connector_D_inner_m=k_D, connector_roughness_m=k_eps,
                            connector_length=k_len,
                            harp_type=harp_type_str,
                            **_common,
                        )
                        inlet_id  = topo1.inlet_node_id
                        outlet_id = topo2.outlet_node_id
                    elif harp_type_str == "I":
                        net, topo1 = build_center_fed_harp(N, **_common)
                        topo2     = None
                        inlet_id  = topo1.inlet_node_id
                        outlet_id = topo1.outlet_node_id
                    elif harp_type_str == "DI":
                        net, topo1 = build_biinlet_harp(N, **_common)
                        topo2     = None
                        inlet_id  = topo1.inlet_node_id
                        outlet_id = topo1.outlet_node_id
                    else:  # Z or U, single harp
                        net, topo1 = build_harp(N, harp_type=harp_type_str, **_common)
                        topo2     = None
                        inlet_id  = topo1.inlet_node_id
                        outlet_id = topo1.outlet_node_id

                    _solve_kwargs = dict(
                        gas_flows_kgh=gas_flows,
                        liquid_type=liq_type,
                        liquid_flows_kgh=liq_flows,
                        inlet_node_id=inlet_id,
                        outlet_node_id=outlet_id,
                        P_outlet_pa=P_outlet * 1e5,
                        max_iter=int(max_iter),
                        relax=float(relax),
                    )

                    def _rebuild_net():
                        """Rebuild a fresh network (solver mutates in-place)."""
                        if effective_series:
                            _n, _t1, _t2 = build_series_harp(
                                N, connector_D_inner_m=k_D,
                                connector_roughness_m=k_eps,
                                connector_length=k_len,
                                harp_type=harp_type_str, **_common)
                            return _n, _t1, _t2, _t1.inlet_node_id, _t2.outlet_node_id
                        elif harp_type_str == "I":
                            _n, _t1 = build_center_fed_harp(N, **_common)
                            return _n, _t1, None, _t1.inlet_node_id, _t1.outlet_node_id
                        elif harp_type_str == "DI":
                            _n, _t1 = build_biinlet_harp(N, **_common)
                            return _n, _t1, None, _t1.inlet_node_id, _t1.outlet_node_id
                        else:
                            _n, _t1 = build_harp(N, harp_type=harp_type_str, **_common)
                            return _n, _t1, None, _t1.inlet_node_id, _t1.outlet_node_id

                    def _run_solver(_net, _m_kgs):
                        if use_hardy_cross:
                            return solve_network_hardy_cross(
                                _net, m_total_kgs=_m_kgs, **_solve_kwargs)
                        return solve_network(_net, m_total_kgs=_m_kgs, **_solve_kwargs)

                    if goal_seek and P_inlet_target is not None:
                        # Outer loop: adjust m_total so P_inlet_result ≈ P_inlet_target
                        # Uses dP ∝ m^n scaling (n≈2 turbulent) for fast convergence.
                        _m = m_total
                        _gs_warnings = []
                        for _gs_iter in range(6):
                            _gn, _gt1, _gt2, _gi, _go = _rebuild_net()
                            _sr = _run_solver(_gn, _m)
                            _P_in = _gn.node(_gi).P_pa / 1e5
                            _dP_result = _P_in - P_outlet
                            _dP_target = P_inlet_target - P_outlet
                            if _dP_result <= 0 or _dP_target <= 0:
                                _gs_warnings.append("ΔP ≤ 0 — check that P_target > P_outlet.")
                                break
                            _err = abs(_P_in - P_inlet_target) / max(abs(P_inlet_target), 1.0)
                            if _err < 1e-4:
                                break
                            # Scale: dP ∝ m^2 → m_new = m_old × √(dP_target/dP_result)
                            _m = _m * (_dP_target / _dP_result) ** 0.5
                            _m = max(_m, 1e-6)

                        solve_res = _sr
                        topo1, topo2 = _gt1, _gt2
                        net = _gn
                        _found_m_kgs = _m
                        # Only forward error warnings; the flow result is shown as a metric
                        solve_res.warnings = _gs_warnings + solve_res.warnings
                    else:
                        _net_b, topo1, topo2, inlet_id, outlet_id = _rebuild_net()
                        net = _net_b
                        solve_res = _run_solver(net, m_total)
                        _found_m_kgs = m_total

                    rep1 = check_starvation(solve_res, topo1.channel_edge_ids,
                                             Vsl_threshold=Vsl_threshold_eff)
                    rep2 = (check_starvation(solve_res, topo2.channel_edge_ids,
                                              Vsl_threshold=Vsl_threshold_eff)
                            if topo2 else None)

                    st.session_state[_cache_key] = (
                        solve_res, topo1, topo2, rep1, rep2, _found_m_kgs, solve_mode)
                    st.session_state[_hash_key]  = _inputs_hash

                except Exception as exc:
                    st.error(f"Solver error: {exc}")
                    st.session_state.pop(_cache_key, None)
                    return

    # ── Results ───────────────────────────────────────────────────────────────
    if _cache_key not in st.session_state:
        with col_out:
            st.info("Configure the manifold and click **▶ Solve network**.")
        return

    _cached = st.session_state[_cache_key]
    if len(_cached) != 7:
        st.session_state.pop(_cache_key, None)
        with col_out:
            st.info("Cache format updated — please click **▶ Solve network** to re-run.")
        return
    solve_res, topo1, topo2, rep1, rep2, found_m_kgs, cached_mode = _cached

    with col_out:
        total_starved = rep1.n_channels_starved + (rep2.n_channels_starved if rep2 else 0)
        total_ch      = rep1.n_channels_total   + (rep2.n_channels_total   if rep2 else 0)
        mdi_combined  = (rep1.maldistribution_index
                         + (rep2.maldistribution_index if rep2 else 0)) / (2 if rep2 else 1)

        P_in_result  = solve_res.net.node(topo1.inlet_node_id).P_pa / 1e5
        dP_bar       = P_in_result - P_outlet
        dP_str       = f"{dP_bar*1000:.1f} mbar" if dP_bar < 1 else f"{dP_bar:.4f} bar"
        found_kgh     = found_m_kgs * 3600.0
        n_total_harps = 2 if topo2 else 1
        # Per-physical-channel values (branch flow ÷ channels_per_branch)
        _n_branches   = N * n_total_harps
        found_ch_lh   = found_kgh / max(_n_branches * channels_per_branch, 1)
        found_ch_vsl  = (found_m_kgs / max(_n_branches * channels_per_branch, 1)) \
                        / (_rho_approx * _ch_A) if _ch_A > 0 else 0.0

        bc = st.columns(5)
        _solver_lbl = "Hardy-Cross" if use_hardy_cross else "GGA"
        bc[0].metric("Solver", f"{'✅' if solve_res.converged else '⚠️'} {_solver_lbl} · {solve_res.iterations} iters")

        if cached_mode == _MODE_DP:
            bc[1].metric("Total flow (result)", f"{found_kgh:.1f} kg/h")
            bc[2].metric("ΔP achieved", dP_str)
        elif cached_mode == _MODE_PERCHAN:
            bc[1].metric("Total flow", f"{found_kgh:.1f} kg/h")
            bc[2].metric("ΔP (total)", dP_str)
        else:  # _MODE_FLOW
            bc[1].metric("P_inlet (result)", f"{P_in_result:.4f} bara")
            bc[2].metric("ΔP (total)", dP_str)

        bc[3].metric("Low-flow channels", f"{total_starved} / {total_ch}")
        bc[4].metric("MDI", f"{mdi_combined:.4f}")

        # Per-channel summary line beneath metrics
        _ch_label = "per physical channel" if channels_per_branch > 1 else "per channel"
        st.caption(
            f"P_inlet = **{P_in_result:.4f} bara** — "
            f"total **{found_kgh:.1f} kg/h** — "
            f"**{found_ch_lh:.1f} L/h** {_ch_label} (V_sl ≈ {found_ch_vsl:.3f} m/s)"
        )

        if not solve_res.converged:
            st.warning("Solver did not converge — results are approximate. "
                       "Try reducing relaxation or increasing max iterations.")
        for w in solve_res.warnings:
            if "did not converge" not in w.lower():
                st.caption(f"ℹ {w}")

        st.markdown("##### Network diagram")
        _cpb = channels_per_branch
        st.caption(
            f"{N} branches"
            + (f" × {_cpb} ch/branch = {N*_cpb} physical channels" if _cpb > 1 else f" · {N} channels")
            + f" · {'single-phase' if single_phase else 'two-phase'}"
            + (" · 2 harps in series" if topo2 else "")
            + " — hover for detail"
        )
        fig_diag = _make_harp_diagram(N, solve_res, topo1, topo2, rep1, rep2,
                                       Vsl_threshold_eff, single_phase)
        st.plotly_chart(fig_diag, use_container_width=True,
                        config={"displayModeBar": False})

        dc1, dc2 = st.columns([3, 2])
        with dc1:
            st.markdown("##### Channel flow distribution")
            reports_bar = [(rep1, "H1")] + ([(rep2, "H2")] if rep2 else [])
            st.plotly_chart(_make_flow_bar_chart(reports_bar, Vsl_threshold_eff, single_phase),
                            use_container_width=True, config={"displayModeBar": False})
        with dc2:
            st.markdown("##### Header pressure profile")
            st.plotly_chart(_make_pressure_profile(solve_res, topo1, topo2),
                            use_container_width=True, config={"displayModeBar": False})

        # ── Per-channel table ──────────────────────────────────────────────────
        st.markdown("##### Channel detail")
        import pandas as pd

        def _build_table(report: StarvationReport, label: str) -> None:
            k   = channels_per_branch
            A_s = math.pi * c_D ** 2 / 4.0   # physical single-channel area
            rows = []
            for r in report.channel_results:
                # Per-physical-channel values (branch values ÷ k)
                m_ch  = r.m_kgs / k
                vsl_ch = m_ch / max(1000.0 * A_s, 1e-12)   # ρ≈1000 approximation
                if single_phase:
                    frac = r.m_kgs / max(report.mean_m_kgs, 1e-12)
                    status = ("🔴 LOW" if frac < 0.7 else
                              "🟡" if frac < 0.9 else "🟢 OK")
                else:
                    status = ("🔴 STARVED" if r.starved else
                              "🟡 Low" if vsl_ch < 2 * Vsl_threshold else "🟢 OK")
                col_m   = "m/ch (g/s)" if k > 1 else "m (g/s)"
                col_vsl = "V_sl/ch (m/s)" if k > 1 else "V_sl (m/s)"
                row = {"Branch": r.channel_index,
                       col_m:   round(m_ch * 1000, 2),
                       col_vsl: round(vsl_ch, 4)}
                if not single_phase:
                    row["V_sg (m/s)"] = round(r.Vsg / max(k ** 0.2, 1), 4)
                    row["x_gas"]      = round(r.x_gas, 4)
                    row["Regime"]     = r.regime
                row["Status"] = status
                rows.append(row)
            df = pd.DataFrame(rows)

            vsl_col = "V_sl/ch (m/s)" if k > 1 else "V_sl (m/s)"
            def _style(row):
                v = row.get(vsl_col, 1)
                if v < Vsl_threshold:
                    return ["background-color:#fee2e2;font-weight:bold"] * len(row)
                if v < 2 * Vsl_threshold:
                    return ["background-color:#fef3c7"] * len(row)
                return [""] * len(row)

            caption = (f"{label} — MDI = {report.maldistribution_index:.4f}"
                       + (f"  ⚠ {report.n_channels_starved} low-flow"
                          if report.n_channels_starved else "")
                       + (f"  · values per physical channel (branch ÷ {k})" if k > 1 else ""))
            st.caption(caption)
            h = min(420, 52 + len(rows) * 36)
            st.dataframe(df.style.apply(_style, axis=1),
                         use_container_width=True, hide_index=True, height=h)

        if series_mode and rep2 is not None:
            dt1, dt2 = st.tabs(["Harp 1", "Harp 2"])
            with dt1: _build_table(rep1, "H1")
            with dt2: _build_table(rep2, "H2")
        else:
            _build_table(rep1, "H1")

        # ── Expert Review ─────────────────────────────────────────────────
        with st.expander("Expert Review", expanded=True):
            st.markdown(_expert_review(
                solve_res, topo1, topo2, rep1, rep2,
                N=N, h_D=h_D, c_D=c_D, c_len=c_len, h_len=h_len,
                P_outlet=P_outlet, single_phase=single_phase,
                channels_per_branch=channels_per_branch,
            ))

        # ── Export report ──────────────────────────────────────────────────
        st.markdown("##### Export")
        with st.spinner("Generating report…") if False else st.empty():
            pass
        if st.button("📄  Generate Word report", key=_k("gen_report"),
                     use_container_width=False):
            with st.spinner("Building report…"):
                try:
                    _fig_diag  = _make_harp_diagram(N, solve_res, topo1, topo2,
                                                    rep1, rep2, Vsl_threshold_eff, single_phase)
                    _fig_bar   = _make_flow_bar_chart(
                        [(rep1, "H1")] + ([(rep2, "H2")] if rep2 else []),
                        Vsl_threshold, single_phase)
                    _fig_pres  = _make_pressure_profile(solve_res, topo1, topo2)
                    _report_bytes = _generate_harp_report(
                        solve_res, topo1, topo2, rep1, rep2,
                        dict(
                            harp_type_str=harp_type_str, series_mode=series_mode,
                            N=N, h_D=h_D, c_D=c_D, c_len=c_len, h_len=h_len,
                            channels_per_branch=channels_per_branch,
                            liq_type=liq_type, gas_lbl=_gas_lbl,
                            m_total=m_total, found_m_kgs=found_m_kgs,
                            x_inlet=x_inlet,
                            P_bara=P_bara, P_outlet=P_outlet, T_C=T_C,
                            Vsl_threshold=Vsl_threshold, single_phase=single_phase,
                            solve_mode=cached_mode, solver_method=solver_method,
                            corr=corr, void=void,
                        ),
                        _fig_diag, _fig_bar, _fig_pres,
                    )
                    st.session_state[_k("report_bytes")] = _report_bytes
                except Exception as exc:
                    st.error(f"Report error: {exc}")

        if _k("report_bytes") in st.session_state:
            import datetime
            fname = (f"harp_{harp_type_str}_{N}ch_"
                     f"{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx")
            st.download_button(
                "⬇  Download report (.docx)",
                data=st.session_state[_k("report_bytes")],
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=_k("dl_report"),
                use_container_width=False,
            )

        with st.expander("ℹ Method", expanded=False):
            st.markdown(
                f"""
**Global Gradient Algorithm** (Todini & Pilati, 1988) — same method as EPANET.

Newton-Raphson simultaneous solution for all nodal pressures and pipe flows.
Converges quadratically — typically **2–5 iterations** regardless of network size.

**Sign convention (§A):** `dP = r·m·|m|` — pressure drop always opposes flow.
Reversed-flow channels automatically get the correct opposing dP.

**T-junction K-factors (§B):** each channel carries K_fwd = 0.5 (dividing T)
or K_rev = 0.3 (combining T), switching every iteration based on flow direction.
Correct for the near-inlet T-junctions in a U-manifold where flow reverses.

**Boundary conditions:** outlet pressure is **pinned** at {P_outlet:.2f} bara.
The displayed P_inlet is a **result** — the pressure required at the inlet
to push the specified flow through the network.
                """
            )


# Alias for app.py references
render_network_solver_tab = render_harp_network_tab

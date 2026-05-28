# report_generator.py
"""
Generates a Word (.docx) calculation report from multiphase hydraulics calculation results.
"""
from io import BytesIO
from datetime import datetime
import atexit
import threading

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_HDR_BG = "2563EB"
_HDR_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ALT_BG = "F1F5F9"

_IMG_TIMEOUT = 60  # seconds per render attempt before declaring kaleido hung
_png_cache: dict = {}  # (id(fig), width, height, scale) → PNG bytes or None


def _kill_kaleido_proc():
    """Kill the current kaleido subprocess so the next call spawns a fresh one."""
    try:
        import plotly.io as pio
        scope = getattr(getattr(pio, "kaleido", None), "scope", None)
        if scope is not None:
            proc = getattr(scope, "_proc", None)
            if proc is not None and proc.poll() is None:
                proc.kill()
            scope._proc = None
    except Exception:
        pass


def _cleanup_kaleido():
    """Kill the kaleido subprocess on process exit.
    Prevents orphaned kaleido Node.js processes after Streamlit shuts down.
    """
    _kill_kaleido_proc()


atexit.register(_cleanup_kaleido)


def _fig_to_png(fig, width=900, height=400, scale=2):
    """Render a Plotly figure to PNG bytes with a hard timeout.

    Tries twice: if the first attempt times out (kaleido deadlock), the hung
    subprocess is killed and a fresh one is used for the retry.
    Returns bytes on success, None if both attempts fail.
    Uses daemon threads so a hung render never blocks process exit.
    """
    import plotly.io as pio
    _key = (id(fig), width, height, scale)
    if _key in _png_cache:
        return _png_cache[_key]

    for _attempt in range(2):
        _result = [None]
        def _render():
            try:
                _result[0] = pio.to_image(
                    fig, format="png", width=width, height=height, scale=scale)
            except Exception:
                pass
        _t = threading.Thread(target=_render, daemon=True)
        _t.start()
        _t.join(timeout=_IMG_TIMEOUT)
        if _result[0] is not None:
            break
        # Timed out — kill the hung kaleido process; retry will spawn a fresh one
        _kill_kaleido_proc()

    _png_cache[_key] = _result[0]
    return _result[0]


def prefetch_figures(specs):
    """Warm the PNG cache for a list of (fig, width, height, scale) tuples.

    Renders sequentially — kaleido uses a single subprocess internally so
    concurrent calls just queue up and earlier timeouts cut off later ones.
    Each render uses the retry logic in _fig_to_png.
    specs with None fig are silently skipped.
    """
    for fig, w, h, s in specs:
        if fig is not None:
            _fig_to_png(fig, w, h, s)


def clear_fig_cache():
    """Clear the PNG render cache (call between report sessions if needed)."""
    _png_cache.clear()


def _shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _style_header(row, font_size=9):
    for cell in row.cells:
        _shd(cell, _HDR_BG)
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            for run in para.runs:
                run.font.bold       = True
                run.font.color.rgb  = _HDR_FG
                run.font.size       = Pt(font_size)


def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_inches):
            cell.width = Inches(w)


def _fig_caption(doc, text):
    """Add a small grey italic caption paragraph below a figure."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    if p.runs:
        p.runs[0].font.size      = Pt(8)
        p.runs[0].font.italic    = True
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def _cell_font(cell, size_pt=9):
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        for run in para.runs:
            run.font.size = Pt(size_pt)


def _add_toc(doc):
    """Insert a Word TOC field (levels 1–2). Word updates it on open."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el

    toc_h = doc.add_heading("Table of Contents", level=1)
    toc_h.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    para = doc.add_paragraph()
    run  = para.add_run()

    def _fc(type_):
        fc = _el("w:fldChar"); fc.set(_qn("w:fldCharType"), type_); return fc

    run._r.append(_fc("begin"))
    instr = _el("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    run._r.append(_fc("separate"))
    placeholder = _el("w:r")
    t = _el("w:t"); t.text = "[Right-click → Update Field to populate]"
    placeholder.append(t)
    para._p.append(placeholder)
    run._r.append(_fc("end"))

    doc.add_page_break()


def _add_footer_page_numbers(doc):
    """Add a centred PAGE / NUMPAGES field in the first section's footer."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove any existing runs, keep pPr (alignment)
    p = para._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    def _field_runs(instr):
        """Return three w:r elements: fldChar begin, instrText, fldChar end."""
        r1 = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r1.append(fc)
        r2 = OxmlElement("w:r")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
        it.text = f" {instr} "; r2.append(it)
        r3 = OxmlElement("w:r")
        fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end"); r3.append(fc2)
        return r1, r2, r3

    for r in _field_runs("PAGE"):
        p.append(r)
    r_sep = OxmlElement("w:r")
    t_sep = OxmlElement("w:t"); t_sep.set(qn("xml:space"), "preserve"); t_sep.text = " / "
    r_sep.append(t_sep); p.append(r_sep)
    for r in _field_runs("NUMPAGES"):
        p.append(r)


def _kv_table(doc, rows_data, col_widths=(2.5, 3.5)):
    """Two-column key-value table with blue header."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=2)
    tbl.style = "Table Grid"
    _style_header(tbl.rows[0])
    tbl.rows[0].cells[0].text = "Parameter"
    tbl.rows[0].cells[1].text = "Value"
    for i, (label, value) in enumerate(rows_data, start=1):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        if i % 2 == 0:
            _shd(row.cells[0], _ALT_BG)
            _shd(row.cells[1], _ALT_BG)
        _cell_font(row.cells[0])
        _cell_font(row.cells[1])
    _set_col_widths(tbl, col_widths)
    return tbl


def generate_report(
    P_bara, T_C,
    gas_flows_kgh,        # dict {species: kg/h}
    liquid_type,          # str
    q_lye,
    props,
    grid_records,
    segments,
    total_dp_kpa,
    outlet_pressure_bara,
    pipe_length_m,
    cumulative_distance,
    fig_sch=None,
    fig_prof=None,
    fig_regime_h=None,
    fig_regime_v=None,
    case_label="Case",
    flow_mode=None,            # "liquid_only","gas_only","gas_liquid","vle"
    custom_liquid=None,        # dict (used for KOH concentration)
    stream_records=None,       # list of stream-balance dicts (VLE phase distribution)
    sensitivity_results=None,  # list of dicts from run_sensitivity()
    slug_records=None,         # list of dicts from slug_dynamics() per slug segment
):
    doc = Document()

    # A4 portrait, 20 mm margins
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    # Derive flow mode if not provided
    if flow_mode is None:
        _has_gas = bool(gas_flows_kgh)
        _has_liq = bool(liquid_type)
        if _has_gas and _has_liq:
            flow_mode = "gas_liquid"
        elif _has_gas:
            flow_mode = "gas_only"
        elif _has_liq:
            flow_mode = "liquid_only"
        else:
            flow_mode = "vle"

    _GAS_ONLY  = flow_mode == "gas_only"
    _LIQ_ONLY  = flow_mode == "liquid_only"
    _IS_VLE    = flow_mode == "vle"
    _TWO_PHASE = flow_mode in ("gas_liquid", "vle")

    _MODE_LABEL = {
        "liquid_only": "Single-Phase Liquid",
        "gas_only":    "Single-Phase Gas (Compressible)",
        "gas_liquid":  "Gas–Liquid Two-Phase",
        "vle":         "Saturated VLE / Two-Phase",
    }
    _mode_str = _MODE_LABEL.get(flow_mode, "Gas–Liquid Two-Phase")

    # Fluid description string
    _gas_label = " + ".join((gas_flows_kgh or {}).keys())
    if _IS_VLE:
        _vle_fl_name = props.get("vle_fluid") or (liquid_type or "—")
        _fluid_desc  = f"{_vle_fl_name}  (saturated VLE)"
    elif _GAS_ONLY:
        _fluid_desc = _gas_label or "—"
    elif _LIQ_ONLY:
        _liq_display = liquid_type or "—"
        if liquid_type == "KOH solution" and custom_liquid:
            _koh_c = custom_liquid.get("koh_conc_wt")
            if _koh_c is not None:
                _liq_display = f"KOH solution ({_koh_c:.0f} wt%)"
        _fluid_desc = _liq_display
    else:
        _liq_display = liquid_type or "—"
        if liquid_type == "KOH solution" and custom_liquid:
            _koh_c = custom_liquid.get("koh_conc_wt")
            if _koh_c is not None:
                _liq_display = f"KOH solution ({_koh_c:.0f} wt%)"
        _fluid_desc = f"{_gas_label} / {_liq_display}"

    _phase_str = "Single-phase" if (_GAS_ONLY or _LIQ_ONLY) else "Two-phase"

    # Section counter
    _sec = [0]
    def _h1(title):
        _sec[0] += 1
        doc.add_heading(f"{_sec[0]}. {title}", level=1)

    def _body(text):
        _p = doc.add_paragraph(text)
        _p.paragraph_format.space_after = Pt(4)
        if _p.runs:
            _p.runs[0].font.size = Pt(9)
        return _p

    # ── Title block ──────────────────────────────────────────────────────────
    h = doc.add_heading(f"Branch Line Hydraulic Calculation — {case_label}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"{_mode_str}  ·  {datetime.now().strftime('%d %B %Y  %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub.runs[0].font.size = Pt(10)

    desc = doc.add_paragraph(
        f"{_phase_str} pressure drop  ·  {_fluid_desc}  ·  Steady-state"
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if desc.runs:
        desc.runs[0].font.size = Pt(9)
        desc.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()

    # ── 1. Purpose ───────────────────────────────────────────────────────────
    _h1("Purpose")
    if _LIQ_ONLY:
        _purpose_text = (
            f"This calculation determines the single-phase liquid pressure drop along the "
            f"{case_label} branch pipeline, which carries {_fluid_desc}. "
            f"The Darcy-Weisbach equation with minor losses (Crane TP-410 equivalent-length "
            f"method) is used. The result — inlet pressure, outlet pressure, and total ΔP — "
            f"is used to size the pipe and set the required equipment outlet pressure."
        )
    elif _GAS_ONLY:
        _purpose_text = (
            f"This calculation determines the single-phase compressible gas pressure drop along "
            f"the {case_label} branch pipeline, which carries {_fluid_desc}. "
            f"The Darcy-Weisbach equation is applied with pressure marching — gas density is "
            f"re-evaluated at each segment inlet to capture compressibility. "
            f"Minor losses are accounted for using equivalent lengths (Crane TP-410). "
            f"The result — inlet pressure, outlet pressure, and total ΔP — is used to "
            f"size the pipe and set the required equipment outlet pressure."
        )
    elif _IS_VLE:
        _purpose_text = (
            f"This calculation determines the two-phase pressure drop along the {case_label} "
            f"branch pipeline, which carries {_fluid_desc}. "
            f"The inlet quality (vapour mass fraction) and total mass flow rate are specified; "
            f"the fluid maintains thermodynamic equilibrium at each segment, with quality "
            f"evolving isenthalpically as pressure falls. "
            f"The result — inlet pressure, outlet pressure, total ΔP, and phase distribution — "
            f"is used to size the pipe and assess the evolving two-phase character along the route."
        )
    else:
        _purpose_text = (
            f"This calculation determines the two-phase pressure drop along the {case_label} "
            f"branch pipeline, which carries {_fluid_desc} from a "
            f"process unit to a gas–liquid separator. "
            f"The result — inlet pressure, outlet pressure, and total ΔP — is used to "
            f"size the pipe and, in combination with the collecting header calculation, "
            f"to establish the required equipment outlet pressure."
        )
    _body(_purpose_text)
    doc.add_paragraph()

    # ── 2. Method ────────────────────────────────────────────────────────────
    _h1("Method")
    if _GAS_ONLY or _LIQ_ONLY:
        _method_paras = [
            (
                "Single-phase Darcy-Weisbach: ΔP_fric = f (L/D) ρ V² / 2. "
                "Friction factor from Churchill (1977), covering laminar, transition, and "
                "turbulent regimes without separate regime checks. "
                "For compressible gas, density is re-evaluated at each segment inlet using the "
                "ideal-gas law; gas viscosity is obtained from CoolProp. "
                "Minor losses: equivalent-length method, Crane TP-410. "
                "Erosion check: API RP 14E, C = 100 continuous service. "
                "Packages: fluids · CoolProp · python-docx."
            ),
        ]
    else:
        _method_paras = [
            (
                "Six two-phase ΔP correlations are available: Beggs & Brill (1973, default), "
                "Friedel, Lockhart-Martinelli, Müller-Steinhagen & Heck, Chisholm, and "
                "Kim-Mudawar. The selected correlation is shown in the Segment Analysis table. "
                "The pipeline is divided into user-defined segments with pressure marching — "
                "gas density is re-evaluated at each segment inlet to capture compressibility. "
                "Each segment's ΔP is decomposed into frictional, gravitational, and "
                "accelerational components."
            ),
            (
                "Void fraction uses either the homogeneous model "
                "(α = (x/ρ_g) / (x/ρ_g + (1−x)/ρ_l)) or the Rouhani-1 slip-flow model. "
                "Flow regime is classified automatically: Taitel-Dukler + Mandhane-Gregory-Aziz "
                "for horizontal segments (|θ| ≤ 15°); Wallis/Taitel (1980) annular-onset "
                "criterion with void-fraction thresholds for vertical segments (|θ| ≥ 75°). "
                "Gas properties use the ideal-gas law and CoolProp viscosities; water vapour "
                "is included via Dalton's Law for aqueous liquids."
            ),
            (
                "Minor losses: equivalent-length method, Crane TP-410. "
                "Erosion: API RP 14E, C = 100 continuous service. "
                "Packages: fluids · CoolProp · python-docx."
            ),
        ]
        if _IS_VLE:
            _method_paras.append(
                "VLE mode: thermodynamic equilibrium is maintained at each segment. "
                "Inlet enthalpy is computed from the specified quality and inlet pressure "
                "via CoolProp; at each subsequent segment the quality is re-derived from the "
                "updated pressure by isenthalpic flashing (constant total enthalpy). "
                "Phase properties (density, viscosity, surface tension) are evaluated with "
                "CoolProp at the local saturation state. "
                "Note: the current implementation assumes a single-component fluid; "
                "mixtures require an explicit flash calculation not yet implemented."
            )
    for _txt in _method_paras:
        _body(_txt)
    doc.add_paragraph()

    # ── 3. Process Conditions ────────────────────────────────────────────────
    _h1("Process Conditions")
    _T_display = (f"{T_C:.1f} °C" if T_C is not None
                  else f"{props.get('T_sat_C', props.get('T_C', 0.0)):.1f} °C  (T_sat)")
    _cond_rows = [
        ("Case",           case_label),
        ("Flow Mode",      _mode_str),
        ("Inlet Pressure", f"{P_bara:.2f} bara"),
        ("Temperature",    _T_display),
    ]
    if _IS_VLE:
        _vle_x  = props.get("x_gas", 0.0)
        _m_tot  = props.get("m_total_kgs", 0.0) * 3600.0
        _cond_rows += [
            ("Fluid (saturated VLE)", _vle_fl_name),
            ("Total Mass Flow",       f"{_m_tot:.2f} kg/h"),
            ("Inlet Quality x",       f"{_vle_x:.3f}"),
            ("T_sat at P_inlet",      _T_display),
        ]
    elif _GAS_ONLY:
        for _sp, _flow in (gas_flows_kgh or {}).items():
            _cond_rows.append((f"{_sp} Mass Flow", f"{_flow:.3f} kg/h"))
    elif _LIQ_ONLY:
        _cond_rows += [
            ("Liquid Type",        _fluid_desc),
            ("Liquid Volume Flow", f"{q_lye:.3f} m³/h"),
        ]
    else:
        # gas_liquid
        for _sp, _flow in (gas_flows_kgh or {}).items():
            _cond_rows.append((f"{_sp} Mass Flow", f"{_flow:.3f} kg/h"))
        _cond_rows += [
            ("Liquid Type",        _fluid_desc.split(" / ", 1)[-1] if " / " in _fluid_desc else _fluid_desc),
            ("Liquid Volume Flow", f"{q_lye:.3f} m³/h"),
        ]
    _cond_rows.append(("Number of Segments", str(len(segments))))
    _kv_table(doc, _cond_rows)
    doc.add_paragraph()

    # ── 4. Phase Thermodynamics ──────────────────────────────────────────────
    _h1("Phase Thermodynamics  (inlet)")
    _thermo_rows = []
    if not _LIQ_ONLY:
        _thermo_rows += [
            ("Gas / Vapour Density ρ_g",    f"{props.get('rho_g', 0.0):.4f} kg/m³"),
            ("Gas Mixture MW",              f"{props.get('MW_mix_gmol', 0.0):.3f} g/mol"),
            ("Gas Dynamic Viscosity μ_g",   f"{props.get('mu_g', 0.0)*1e6:.2f} µPa·s"),
        ]
    _thermo_rows.append(("Liquid Type", props.get("liquid_type", liquid_type or "—")))
    if not _GAS_ONLY:
        _thermo_rows += [
            ("Liquid Density ρ_l",          f"{props.get('rho_l', 0.0):.2f} kg/m³"),
            ("Liquid Dynamic Viscosity μ_l", f"{props.get('mu_l', 0.0)*1e3:.4f} mPa·s"),
            ("Surface Tension σ",           f"{props.get('sigma', 0.0)*1e3:.3f} mN/m"),
        ]
    if _TWO_PHASE:
        _thermo_rows += [
            ("Mass Quality x",              f"{props.get('x_gas', 0.0)*100:.4f} %"),
            ("Void Fraction α",             f"{props.get('alpha', 0.0)*100:.2f} %"),
        ]
    if props.get("P_sat_H2O_pa", 0) > 0:
        _thermo_rows += [
            ("H₂O Saturation Pressure", f"{props['P_sat_H2O_pa']/1e5:.4f} bara"),
            ("H₂O Vapour Flow",         f"{props['m_vapor_h2o_kgh']:.4f} kg/h"),
        ]
    _kv_table(doc, _thermo_rows)
    doc.add_paragraph()

    # ── 5. Segment Analysis ──────────────────────────────────────────────────
    _h1("Segment Analysis")
    if _GAS_ONLY or _LIQ_ONLY:
        _p5_text = (
            "The segment table shows the hydraulic result for each individual pipe section, "
            "marching from inlet to outlet. "
            "ΔP (kPa) shows the pressure drop consumed by each segment — any segment that "
            "contributes disproportionately is the first candidate for bore enlargement or "
            "fitting reduction. "
            "V_m/V_e is the ratio of flow velocity to the API RP 14E erosional velocity "
            "(C = 100 for continuous service); values above 1.0 flag an erosion concern."
        )
    else:
        _p5_text = (
            "The segment table shows the hydraulic result for each individual pipe section, "
            "marching from inlet to outlet. Three columns are of primary interest. "
            "Regime shows the predicted two-phase flow pattern — Bubbly, Slug, Annular, or Mist — "
            "which governs the dominant loss mechanism and drives vibration and pulsation loads. "
            "Slug flow on near-horizontal lines produces cyclic pressure surges at bends and "
            "supports; Annular and Mist flow indicate a gas-dominated, high-velocity service "
            "associated with higher erosion risk. A change in regime from one segment to the next "
            "signals a shift in flow character that the designer should be aware of. "
            "ΔP (kPa) shows how much of the total pressure budget is consumed by each segment — "
            "any segment that contributes disproportionately is the first candidate for bore "
            "enlargement or fitting reduction. "
            "V_m/V_e is the ratio of mixture velocity to the API RP 14E erosional velocity "
            "(C = 100 for continuous service); values above 1.0 flag an erosion concern."
        )
    _body(_p5_text)
    doc.add_paragraph()

    _COLS   = ["Seg", "Pipe", "ID (mm)", "Type", "L (m)", "L_eq (m)", "Fittings", "Regime",
               "V_m (m/s)", "V_m/V_e", "ΔP (kPa)", "P_out (bara)"]
    _WIDTHS = [0.25,  0.50,   0.42,     0.80,   0.38,   0.42,      0.70,      0.85,
               0.48,         0.44,      0.50,       0.53]

    if grid_records:
        tbl3 = doc.add_table(rows=len(grid_records) + 1, cols=len(_COLS))
        tbl3.style = "Table Grid"
        _style_header(tbl3.rows[0], font_size=8)
        for j, col in enumerate(_COLS):
            tbl3.rows[0].cells[j].text = col
        for i, rec in enumerate(grid_records, start=1):
            row = tbl3.rows[i]
            if i % 2 == 0:
                for cell in row.cells:
                    _shd(cell, _ALT_BG)
            for j, col in enumerate(_COLS):
                cell = row.cells[j]
                cell.text = str(rec.get(col, ""))
                _cell_font(cell, size_pt=8)
        _set_col_widths(tbl3, _WIDTHS)

    doc.add_paragraph()

    # ── 6. VLE Phase Distribution (VLE mode only) ────────────────────────────
    if _IS_VLE and stream_records:
        _h1("VLE Phase Distribution")
        _body(
            "The phase distribution table traces the vapour quality, phase flow rates, and "
            "saturation temperature at each stream boundary as pressure falls along the pipeline. "
            "Quality x increases as pressure drops and the fluid partially vaporises; the "
            "vapour fraction α increases correspondingly. "
            "A rapid rise in x or α in a short segment indicates a high-vaporisation zone "
            "that may require attention to flow regime and erosion velocity."
        )
        doc.add_paragraph()
        _vfl = props.get("vle_fluid") or ""
        _vap_key = f"{_vfl} vapour  kg/h"
        _liq_key = f"{_vfl} liquid  kg/h"
        _VD_COLS = ["Stream", "P (bara)", "T_sat (°C)", "x (−)", "Vapour (kg/h)", "Liquid (kg/h)", "α (−)"]
        _VD_W    = [1.10,     0.60,       0.65,         0.50,    0.75,            0.75,            0.52]
        tbl_vd = doc.add_table(rows=len(stream_records) + 1, cols=len(_VD_COLS))
        tbl_vd.style = "Table Grid"
        _style_header(tbl_vd.rows[0], font_size=8)
        for j, col in enumerate(_VD_COLS):
            tbl_vd.rows[0].cells[j].text = col
        for i, sr in enumerate(stream_records, start=1):
            row = tbl_vd.rows[i]
            if i % 2 == 0:
                for cell in row.cells:
                    _shd(cell, _ALT_BG)
            _vals = [
                sr.get("Stream", ""),
                f"{sr.get('P (bara)', 0.0):.3f}",
                f"{sr.get('T (°C)', 0.0):.2f}",
                f"{sr.get('x (−)', 0.0):.4f}",
                f"{sr.get(_vap_key, 0.0):.2f}",
                f"{sr.get(_liq_key, 0.0):.2f}",
                f"{sr.get('α (−)', 0.0):.4f}",
            ]
            for j, v in enumerate(_vals):
                row.cells[j].text = v
                _cell_font(row.cells[j], size_pt=8)
        _set_col_widths(tbl_vd, _VD_W)
        doc.add_paragraph()

    # ── 7. System Totals ─────────────────────────────────────────────────────
    _h1("System Totals")
    _body(
        "The system totals consolidate all segment results into the key hydraulic deliverable. "
        "Total ΔP is the pressure budget consumed by the branch pipeline; it must be "
        "subtracted from the process unit outlet pressure to obtain the separator operating "
        "pressure, or alternatively added to the separator target pressure to determine "
        "the required process unit outlet pressure. "
        "For a goal-seek run, the inlet pressure was found iteratively so that the pipeline "
        "outlet reaches the target separator pressure exactly. "
        "The effective length (including fitting equivalent lengths per Crane TP-410) "
        "confirms that all minor losses — valves, elbows, reducers, tees — have been "
        "captured in the friction model."
    )
    doc.add_paragraph()
    _kv_table(doc, [
        ("Case",                                 case_label),
        ("Inlet Pressure",                       f"{P_bara:.4f} bara"),
        ("Outlet Pressure",                      f"{outlet_pressure_bara:.4f} bara"),
        ("Total Pressure Drop ΔP",               f"{total_dp_kpa:.4f} kPa"),
        ("Total Pressure Drop ΔP",               f"{total_dp_kpa / 100:.6f} bar"),
        ("Pipe Length (physical segments)",      f"{pipe_length_m:.2f} m"),
        ("Effective Length (incl. fittings)",    f"{cumulative_distance:.2f} m"),
    ])

    # ── 8. Slug Flow Dynamics (if slug segments present) ────────────────────
    if slug_records:
        _h1("Slug Flow Dynamics")
        _body(
            "One or more pipeline segments were classified as slug or intermittent flow. "
            "The table below characterises each slug segment: arrival frequency, translational "
            "velocity, liquid holdup, indicative slug length, and the peak pressure pulse and "
            "force that a liquid slug imposes at a 90° elbow. "
            "Design values include a Dynamic Load Factor of 2.0 per ASME B31.3 occasional-load "
            "provisions. Use these results as first-pass inputs for pipe-support and structural "
            "assessment; ±30 % accuracy is typical for empirical slug correlations."
        )
        doc.add_paragraph()
        _SEV_COLOURS = {"Low": "D1FAE5", "Moderate": "FEF3C7",
                        "Severe": "FEE2E2", "High": "FEE2E2"}
        _SLUG_COLS = [
            "Seg", "DN", "Regime", "Severity",
            "f_slug (Hz)", "f_slug (slugs/min)",
            "V_slug (m/s)", "H_Ls", "L_slug (m)",
            "ΔP_pulse (kPa)", "ΔP_design (kPa)",
            "F_elbow (N)", "F_design (N)",
        ]
        _SLUG_W = [0.28, 0.38, 0.90, 0.68, 0.52, 0.72, 0.58, 0.38, 0.55, 0.68, 0.72, 0.58, 0.62]
        _sev_col_idx = _SLUG_COLS.index("Severity")
        tbl_slug = doc.add_table(rows=len(slug_records) + 1, cols=len(_SLUG_COLS))
        tbl_slug.style = "Table Grid"
        _style_header(tbl_slug.rows[0], font_size=8)
        for j, col in enumerate(_SLUG_COLS):
            tbl_slug.rows[0].cells[j].text = col
        for i, rec in enumerate(slug_records, start=1):
            row = tbl_slug.rows[i]
            if i % 2 == 0:
                for cell in row.cells:
                    _shd(cell, _ALT_BG)
            for j, col in enumerate(_SLUG_COLS):
                row.cells[j].text = str(rec.get(col, ""))
                _cell_font(row.cells[j], size_pt=8)
            sev_val = rec.get("Severity", "")
            if sev_val in _SEV_COLOURS:
                _shd(row.cells[_sev_col_idx], _SEV_COLOURS[sev_val])
        _set_col_widths(tbl_slug, _SLUG_W)
        doc.add_paragraph()
        # Summary row
        _max_freq   = max(r["f_slug (Hz)"]       for r in slug_records)
        _max_vel    = max(r["V_slug (m/s)"]      for r in slug_records)
        _max_pulse  = max(r["ΔP_pulse (kPa)"]    for r in slug_records)
        _max_design = max(r["ΔP_design (kPa)"]   for r in slug_records)
        _max_fe     = max(r["F_elbow (N)"]        for r in slug_records)
        _max_fd     = max(r["F_design (N)"]       for r in slug_records)
        _worst_sev  = next(
            (s for s in ("Severe", "High", "Moderate", "Low")
             if any(r.get("Severity") == s for r in slug_records)), "—"
        )
        _kv_table(doc, [
            ("Max slug frequency",               f"{_max_freq:.3f} Hz"),
            ("Max slug velocity",                f"{_max_vel:.2f} m/s"),
            ("Max ΔP pulse at 90° elbow",        f"{_max_pulse:.2f} kPa"),
            ("Max ΔP design (DLF 2.0)",          f"{_max_design:.2f} kPa"),
            ("Max elbow force",                  f"{_max_fe:.1f} N"),
            ("Max elbow force — design (DLF 2)", f"{_max_fd:.1f} N"),
            ("Overall severity",                 _worst_sev),
        ])
        doc.add_paragraph()
        # Severity criteria reference tables
        _body("Severity Criteria")
        _kv_n_table(doc,
            ["Criterion A — Momentum Flux  ρ_L × V_slug²  (NORSOK P-001)",
             "Threshold", "Action"],
            [("Low",      "< 50,000 kg/m/s²",         "Standard supports adequate"),
             ("Moderate", "50,000 – 150,000 kg/m/s²", "Dynamic support design recommended"),
             ("Severe",   "> 150,000 kg/m/s²",        "Structural analysis required")],
            col_widths=[1.0, 1.6, 3.87]
        )
        doc.add_paragraph()
        _kv_n_table(doc,
            ["Criterion B — ΔP Pulse / P_operating  (ASME B31.3)",
             "Threshold", "Action"],
            [("Low",      "< 5 %",    "Negligible pressure transient"),
             ("Moderate", "5 – 15 %", "Check flange and valve ratings"),
             ("Severe",   "> 15 %",   "Formal occasional-load check required")],
            col_widths=[1.0, 1.6, 3.87]
        )
        doc.add_paragraph()
        _kv_n_table(doc,
            ["Criterion C — Frequency vs Resonance  (structural dynamics)",
             "Threshold", "Action"],
            [("Low",      "< 0.5 Hz",   "Well below structural resonance range"),
             ("Moderate", "0.5 – 2 Hz", "Verify support spacing"),
             ("High",     "> 2 Hz",     "Resonance risk — structural assessment needed")],
            col_widths=[1.0, 1.6, 3.87]
        )
        doc.add_paragraph()
        slug_note = doc.add_paragraph(
            "References: Gregory-Scott (1969) slug frequency (horizontal empirical); "
            "Bendiksen (1984) slug translational velocity; Gregory et al. (1978) slug body "
            "liquid holdup; Brill & Mukherjee 30D slug length rule-of-thumb. "
            "Pressure pulse: momentum balance at 90° bend. "
            "Design values: ASME B31.3 DLF = 2.0 occasional-load provision. "
            "Momentum flux thresholds: NORSOK P-001. "
            "ΔP% thresholds: ASME B31.3 occasional-load framework. "
            "Overall severity = worst of criteria A, B, C."
        )
        if slug_note.runs:
            slug_note.runs[0].font.size      = Pt(8)
            slug_note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── 9. Method Sensitivity Analysis (if available) ────────────────────────
    if sensitivity_results:
        _ok = [r for r in sensitivity_results if r.get("ok")]
        if _ok:
            doc.add_page_break()
            _h1("Method Sensitivity Analysis")
            _body(
                "All 12 method combinations (6 ΔP correlations × 2 void-fraction models) "
                "were evaluated to quantify uncertainty in total ΔP due to correlation choice. "
                "The spread between the minimum and maximum values defines the uncertainty band "
                "for this service. Combinations that failed to converge are excluded."
            )
            doc.add_paragraph()
            _CORR_S = {"Beggs-Brill": "BB", "Friedel": "Friedel",
                       "Lockhart_Martinelli": "L-M", "Muller_Steinhagen_Heck": "MSH",
                       "Chisholm": "Chisholm", "Kim_Mudawar": "Kim-M"}
            _VOID_S = {"Homogeneous": "Homo", "Rouhani-1 (slip)": "Rouhani-1"}
            _all_dp = [r["total_dp_kpa"] for r in _ok]
            _kv_table(doc, [
                ("Selected method ΔP (kPa)", f"{total_dp_kpa:.3f}"),
                ("Minimum ΔP (kPa)",         f"{min(_all_dp):.3f}"),
                ("Maximum ΔP (kPa)",         f"{max(_all_dp):.3f}"),
                ("Spread (kPa)",             f"{max(_all_dp) - min(_all_dp):.3f}"),
                ("Spread (%)",               f"{(max(_all_dp)-min(_all_dp))/max(total_dp_kpa,1e-9)*100:.1f} %"),
            ])
            doc.add_paragraph()
            # Per-method detail table
            _sm_rows = []
            for r in sensitivity_results:
                _c = _CORR_S.get(r.get("correlation",""), r.get("correlation",""))
                _v = _VOID_S.get(r.get("voidage",""), r.get("voidage",""))
                _dp_str = (f"{r['total_dp_kpa']:.3f}" if r.get("ok")
                           else f"FAIL: {r.get('error','')}")
                _sm_rows.append((f"{_c} / {_v}", _dp_str))
            _kv_table(doc, _sm_rows, col_widths=(3.5, 2.97))
            doc.add_paragraph()

    # ── 9. Visualisations ────────────────────────────────────────────────────
    if fig_sch is not None or fig_prof is not None or fig_regime_h is not None:
        doc.add_page_break()
        _h1("Visualisations")

        if fig_sch is not None:
            doc.add_heading("Pipeline Schematic", level=2)
            _body(
                "The schematic plots each pipe segment as a coloured bar, where the colour "
                "indicates the predicted flow regime. Reading left to right follows the "
                "flow direction from process unit outlet to separator inlet. "
                "A colour change along the route signals a regime transition — for instance, "
                "from Slug to Annular flow as pressure drops and gas expands toward the outlet. "
                "The segment label shows the pipe specification (DN/PN) and the V_m/V_e ratio; "
                "segments where V_m/V_e > 1.0 are highlighted as potential erosion risk locations."
            )
            img = _fig_to_png(fig_sch, width=900, height=520, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out — export without kaleido)")
            doc.add_paragraph()

        if fig_prof is not None:
            doc.add_heading("Pressure Profile", level=2)
            _body(
                "The pressure profile shows absolute pressure (bara) as a function of "
                "cumulative pipe distance from the process unit outlet. "
                "Coloured markers or bands on the profile indicate the predicted flow regime "
                "in each segment, linking the pressure trend to the local flow character. "
                "A steep slope in any one segment identifies the location with the highest "
                "resistance per unit length — typically a segment with many fittings, "
                "a smaller bore, or a significant elevation change."
            )
            img = _fig_to_png(fig_prof, width=900, height=400, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out — export without kaleido)")

        if fig_regime_h is not None or fig_regime_v is not None:
            doc.add_heading("Flow Regime Maps", level=2)
            _body(
                "Each map plots superficial gas velocity (V_sg) against superficial liquid "
                "velocity (V_sl) on a log-log scale. Background zones show the predicted "
                "flow regime across the full operating envelope; coloured markers show each "
                "pipe segment's actual operating point. "
                "Left: horizontal map (Taitel-Dukler 1976 + Mandhane-Gregory-Aziz 1974). "
                "Right: vertical map (Wallis annular criterion + void-fraction thresholds). "
                "Regime boundaries are computed at inlet fluid conditions for the first pipe segment."
            )
            doc.add_paragraph()
            if fig_regime_h is not None:
                doc.add_heading("Horizontal Flow Regime Map", level=3)
                img_h = _fig_to_png(fig_regime_h, width=800, height=600, scale=2)
                if img_h:
                    doc.add_picture(BytesIO(img_h), width=Inches(5.5))
                else:
                    doc.add_paragraph("(chart rendering timed out — export without kaleido)")
                doc.add_paragraph()
            if fig_regime_v is not None:
                doc.add_heading("Vertical Flow Regime Map", level=3)
                img_v = _fig_to_png(fig_regime_v, width=800, height=600, scale=2)
                if img_v:
                    doc.add_picture(BytesIO(img_v), width=Inches(5.5))
                else:
                    doc.add_paragraph("(chart rendering timed out — export without kaleido)")

    # ── Disclaimer ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    _gas_str = " / ".join((gas_flows_kgh or {}).keys())
    if _TWO_PHASE:
        _note_text = (
            f"Engineering Note: The two-phase correlations used here were developed primarily "
            f"for oil/gas systems. Their application to this service ({_fluid_desc}) carries "
            f"an estimated uncertainty of ±20–30 %. Use the sensitivity analysis (Method "
            f"Sensitivity section or Compare tab) to bracket the ΔP range across all available "
            f"methods. Treat as a first-pass engineering estimate; validate against commissioning "
            f"data before use in safety-critical design."
        )
    else:
        _note_text = (
            f"Engineering Note: Single-phase Darcy-Weisbach calculations for this service "
            f"({_fluid_desc}) are well-established, with typical accuracy of ±10–15 % "
            f"depending on surface roughness assumptions and minor-loss data. "
            f"Treat as a first-pass engineering estimate; validate against commissioning data."
        )
    note = doc.add_paragraph(_note_text)
    if note.runs:
        note.runs[0].font.size      = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================================
# COMPARISON REPORT  (Case A vs Case B)
# ============================================================================

def _kv_n_table(doc, headers, rows_data, col_widths=None):
    """Generic n-column table. headers[0] = 'Parameter', rest = case labels."""
    n = len(headers)
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=n)
    tbl.style = "Table Grid"
    _style_header(tbl.rows[0])
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for i, row_vals in enumerate(rows_data, start=1):
        row = tbl.rows[i]
        for j, v in enumerate(row_vals):
            row.cells[j].text = str(v)
        if i % 2 == 0:
            for cell in row.cells:
                _shd(cell, _ALT_BG)
        for cell in row.cells:
            _cell_font(cell)
    if col_widths:
        _set_col_widths(tbl, col_widths)
    return tbl


def _kv3_table(doc, rows_data, col_widths=(2.2, 2.1, 2.1),
               label_a="Case A", label_b="Case B"):
    """Three-column comparison table: Parameter | <label_a> | <label_b>."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=3)
    tbl.style = "Table Grid"
    _style_header(tbl.rows[0])
    for j, label in enumerate(["Parameter", label_a, label_b]):
        tbl.rows[0].cells[j].text = label
    for i, (label, va, vb) in enumerate(rows_data, start=1):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(va)
        row.cells[2].text = str(vb)
        if i % 2 == 0:
            for cell in row.cells:
                _shd(cell, _ALT_BG)
        for cell in row.cells:
            _cell_font(cell)
    _set_col_widths(tbl, col_widths)
    return tbl


def generate_comparison_report(
    results_a, results_b,
    label_a="Case A", label_b="Case B",
    fig_cmp=None, fig_bar=None,
    sensitivity_data=None,
    stack_dp=None,
):
    """
    Generate a Word report comparing two cases side by side.

    results_a / results_b: dicts returned by run_case() in app.py.
    label_a / label_b: custom case names used throughout the report.
    sensitivity_data: optional dict {"sa": [...], "sb": [...], "fig": Figure}
                      from run_sensitivity(); adds a Method Sensitivity section.
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27);  sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9);   sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9);   sec.bottom_margin = Inches(0.9)

    # ── Title ────────────────────────────────────────────────────────────────
    h = doc.add_heading(f"Branch Line Comparison — {label_a}  vs.  {label_b}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Hydraulic Comparison  ·  {datetime.now().strftime('%d %B %Y  %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 1. Purpose ───────────────────────────────────────────────────────────
    doc.add_heading("1. Purpose", level=1)
    _p = doc.add_paragraph(
        f"This report compares the pressure drop along two branch pipelines: "
        f"{label_a} and {label_b}. Both carry flow from an upstream process unit "
        f"to a separator. The comparison supports pipe sizing — selecting the smallest "
        f"bore that keeps ΔP within budget and velocity below the erosion threshold — "
        f"and identifies which line governs the required equipment outlet pressure. "
        f"The sensitivity analysis (if run) quantifies the uncertainty in ΔP due to "
        f"correlation choice across all 12 method combinations."
    )
    _p.paragraph_format.space_after = Pt(4)
    if _p.runs:
        _p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 2. Method ────────────────────────────────────────────────────────────
    doc.add_heading("2. Method", level=1)
    for _txt in [
        ("Six two-phase ΔP correlations are available: Beggs & Brill (1973, default), "
         "Friedel, Lockhart-Martinelli, Müller-Steinhagen & Heck, Chisholm, and Kim-Mudawar. "
         f"{label_a} and {label_b} may use different correlations and void-fraction models. "
         "Gas density is re-evaluated at each segment inlet (pressure marching); each "
         "segment's ΔP is split into frictional, gravitational, and accelerational components."),
        ("Void fraction: homogeneous model or Rouhani-1 slip-flow model. "
         "Flow regime classified automatically — Taitel-Dukler + Mandhane-Gregory-Aziz "
         "for horizontal, Wallis/Taitel (1980) for vertical segments. "
         "Gas properties: ideal-gas law, CoolProp viscosities, Dalton's Law for water vapour. "
         "Minor losses: Crane TP-410. Erosion: API RP 14E, C = 100."),
        ("Where a sensitivity analysis was run, the final section shows total ΔP and flow "
         "regimes across all 12 method combinations (6 correlations × 2 void-fraction models). "
         "Packages: fluids · CoolProp · python-docx."),
    ]:
        _p = doc.add_paragraph(_txt)
        _p.paragraph_format.space_after = Pt(4)
        if _p.runs:
            _p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 2. Process Conditions ────────────────────────────────────────────────
    doc.add_heading("2. Process Conditions", level=1)

    # Collect all unique gas species across both cases (safe for VLE/liquid-only)
    _all_species = list(dict.fromkeys(
        list((results_a.get("gas_flows_kgh") or {}).keys()) +
        list((results_b.get("gas_flows_kgh") or {}).keys())
    ))
    def _tc_str(res):
        tc = res.get("T_C")
        if tc is not None:
            return f"{tc:.1f}"
        tsat = res.get("props", {}).get("T_sat_C", res.get("props", {}).get("T_C", 0.0))
        return f"{tsat:.1f}  (T_sat)"

    def _fluid_label(res):
        fm = res.get("flow_mode", "gas_liquid")
        if fm == "vle":
            fl = res.get("props", {}).get("vle_fluid") or res.get("vle_fluid") or "—"
            x  = res.get("props", {}).get("x_gas", 0.0)
            return f"{fl} (VLE, x={x:.3f})"
        liq = res.get("liquid_type") or "—"
        gas_keys = list((res.get("gas_flows_kgh") or {}).keys())
        if fm == "liquid_only" or not gas_keys:
            return liq
        if fm == "gas_only":
            return " + ".join(gas_keys)
        return f"{' + '.join(gas_keys)} / {liq}"

    _cond_rows = [
        ("Flow mode",
         (results_a.get("flow_mode") or "gas_liquid").replace("_", " ").title(),
         (results_b.get("flow_mode") or "gas_liquid").replace("_", " ").title()),
        ("Inlet Pressure (bara)",
         f"{results_a['P_bara']:.2f}",
         f"{results_b['P_bara']:.2f}"),
        ("Temperature (°C)",
         _tc_str(results_a),
         _tc_str(results_b)),
    ]
    for _sp in _all_species:
        _cond_rows.append((
            f"{_sp} mass flow (kg/h)",
            f"{(results_a.get('gas_flows_kgh') or {}).get(_sp, 0.0):.3f}",
            f"{(results_b.get('gas_flows_kgh') or {}).get(_sp, 0.0):.3f}",
        ))
    # VLE total mass flow
    for _res, _col in [(results_a, 1), (results_b, 2)]:
        if _res.get("flow_mode") == "vle":
            pass  # handled in fluid label row below
    _cond_rows += [
        ("Fluid / service",
         _fluid_label(results_a),
         _fluid_label(results_b)),
        ("Liquid volume flow (m³/h)",
         f"{results_a.get('q_lye', 0.0):.3f}",
         f"{results_b.get('q_lye', 0.0):.3f}"),
    ]
    _kv3_table(doc, _cond_rows, label_a=label_a, label_b=label_b)
    doc.add_paragraph()

    # ── 3. Phase Thermodynamics ──────────────────────────────────────────────
    doc.add_heading("3. Phase Thermodynamics  (inlet conditions)", level=1)
    pa, pb = results_a["props"], results_b["props"]
    _thermo_rows = [
        ("Gas density ρ_g (kg/m³)",          f"{pa['rho_g']:.4f}",             f"{pb['rho_g']:.4f}"),
        ("Gas mixture MW (g/mol)",            f"{pa['MW_mix_gmol']:.3f}",       f"{pb['MW_mix_gmol']:.3f}"),
        ("Liquid density ρ_l (kg/m³)",        f"{pa['rho_l']:.2f}",             f"{pb['rho_l']:.2f}"),
        ("Liquid viscosity μ_l (mPa·s)",      f"{pa['mu_l']*1e3:.4f}",          f"{pb['mu_l']*1e3:.4f}"),
        ("Gas viscosity μ_g (µPa·s)",         f"{pa['mu_g']*1e6:.2f}",          f"{pb['mu_g']*1e6:.2f}"),
        ("Surface tension σ (mN/m)",          f"{pa['sigma']*1e3:.3f}",         f"{pb['sigma']*1e3:.3f}"),
        ("Mass quality x (%)",                f"{pa['x_gas']*100:.4f}",         f"{pb['x_gas']*100:.4f}"),
        ("Void fraction α (%)",               f"{pa['alpha']*100:.2f}",         f"{pb['alpha']*100:.2f}"),
    ]
    if pa.get("P_sat_H2O_pa", 0) > 0 or pb.get("P_sat_H2O_pa", 0) > 0:
        _thermo_rows += [
            ("H₂O saturation pressure (bara)",
             f"{pa.get('P_sat_H2O_pa',0)/1e5:.4f}" if pa.get('P_sat_H2O_pa',0) > 0 else "—",
             f"{pb.get('P_sat_H2O_pa',0)/1e5:.4f}" if pb.get('P_sat_H2O_pa',0) > 0 else "—"),
            ("H₂O vapour flow (kg/h)",
             f"{pa.get('m_vapor_h2o_kgh',0):.4f}" if pa.get('P_sat_H2O_pa',0) > 0 else "—",
             f"{pb.get('m_vapor_h2o_kgh',0):.4f}" if pb.get('P_sat_H2O_pa',0) > 0 else "—"),
        ]
    _kv3_table(doc, _thermo_rows, label_a=label_a, label_b=label_b)
    doc.add_paragraph()

    # ── 4. System Totals ─────────────────────────────────────────────────────
    doc.add_heading("4. System Totals", level=1)
    _p4cmp = doc.add_paragraph(
        "This section compares the two branch lines at the system level. "
        "The governing branch — the one with the higher total ΔP — sets the required "
        "inlet pressure and therefore the minimum process unit outlet pressure. "
        "For an electrolyser or fuel cell where both branches carry gas from a common "
        "pressure boundary, the ΔP difference (last row) is the hydraulic imbalance "
        "that appears as a differential pressure across the internal membrane. "
        "Minimising this differential is the primary design objective alongside "
        "keeping both lines within the erosion velocity limit (V_m/V_e ≤ 1.0)."
    )
    _p4cmp.paragraph_format.space_after = Pt(4)
    if _p4cmp.runs:
        _p4cmp.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _dp_delta = results_b["total_dp_kpa"] - results_a["total_dp_kpa"]
    _max_ve_a = max((r["V_m/V_e"] for r in results_a["grid_records"]), default=0.0)
    _max_ve_b = max((r["V_m/V_e"] for r in results_b["grid_records"]), default=0.0)
    _tot_rows = [
        ("Outlet pressure (bara)",
         f"{results_a['outlet_pressure_bara']:.4f}",
         f"{results_b['outlet_pressure_bara']:.4f}"),
        ("Total ΔP (kPa)",
         f"{results_a['total_dp_kpa']:.4f}",
         f"{results_b['total_dp_kpa']:.4f}"),
        ("Total ΔP (bar)",
         f"{results_a['total_dp_kpa']/100:.6f}",
         f"{results_b['total_dp_kpa']/100:.6f}"),
        (f"ΔP difference {label_b} − {label_a} (kPa)",
         "—",
         f"{_dp_delta:+.4f}"),
        ("Pipe length (m)",
         f"{results_a['pipe_length_m']:.2f}",
         f"{results_b['pipe_length_m']:.2f}"),
        ("Effective length incl. fittings (m)",
         f"{results_a['cumulative_distance']:.2f}",
         f"{results_b['cumulative_distance']:.2f}"),
        ("Worst V_m/V_e (–)",
         f"{_max_ve_a:.3f}",
         f"{_max_ve_b:.3f}"),
    ]
    _kv3_table(doc, _tot_rows, label_a=label_a, label_b=label_b)
    doc.add_paragraph()

    # ── 5. Segment Analysis ──────────────────────────────────────────────────
    _COLS   = ["Seg","Pipe","ID (mm)","Type","L (m)","L_eq (m)","Fittings","Regime",
               "V_m (m/s)","V_m/V_e","ΔP (kPa)","P_out (bara)"]
    _WIDTHS = [0.25, 0.50, 0.42, 0.80, 0.38, 0.42, 0.70, 0.85, 0.48, 0.44, 0.50, 0.53]

    def _seg_table(doc, grid_records):
        if not grid_records:
            return
        tbl = doc.add_table(rows=len(grid_records)+1, cols=len(_COLS))
        tbl.style = "Table Grid"
        _style_header(tbl.rows[0], font_size=8)
        for j, col in enumerate(_COLS):
            tbl.rows[0].cells[j].text = col
        for i, rec in enumerate(grid_records, start=1):
            row = tbl.rows[i]
            if i % 2 == 0:
                for cell in row.cells:
                    _shd(cell, _ALT_BG)
            for j, col in enumerate(_COLS):
                row.cells[j].text = str(rec.get(col, ""))
                _cell_font(row.cells[j], size_pt=8)
        _set_col_widths(tbl, _WIDTHS)

    doc.add_heading(f"5. Segment Analysis — {label_a}", level=1)
    _pseg_a = doc.add_paragraph(
        f"Per-segment results for the {label_a} branch. "
        f"The Regime column shows the predicted two-phase flow pattern for each segment — "
        f"note any changes in regime along the pipe route as these indicate a shift in "
        f"the dominant loss mechanism. "
        f"V_m/V_e > 1.0 flags erosion risk (API RP 14E, C = 100)."
    )
    _pseg_a.paragraph_format.space_after = Pt(4)
    if _pseg_a.runs:
        _pseg_a.runs[0].font.size = Pt(9)
    _seg_table(doc, results_a["grid_records"])
    doc.add_paragraph()
    doc.add_heading(f"6. Segment Analysis — {label_b}", level=1)
    _pseg_b = doc.add_paragraph(
        f"Per-segment results for the {label_b} branch. "
        f"Compare the Regime and V_m/V_e columns with {label_a} above to identify "
        f"where the two branches differ in flow character and loss distribution."
    )
    _pseg_b.paragraph_format.space_after = Pt(4)
    if _pseg_b.runs:
        _pseg_b.runs[0].font.size = Pt(9)
    _seg_table(doc, results_b["grid_records"])
    doc.add_paragraph()

    # ── 7. Slug Flow Dynamics (if slug segments in either case) ──────────────
    _slug_a = results_a.get("slug_records") or []
    _slug_b = results_b.get("slug_records") or []
    if _slug_a or _slug_b:
        doc.add_heading("7. Slug Flow Dynamics", level=1)
        _pslug = doc.add_paragraph(
            "One or more pipeline segments were classified as slug or intermittent flow. "
            "Slug frequency (Gregory-Scott 1969), translational velocity (Bendiksen 1984), "
            "liquid holdup (Gregory et al. 1978), and 90° elbow pressure pulse / force "
            "(momentum balance, ASME B31.3 DLF = 2.0) are shown per case below."
        )
        _pslug.paragraph_format.space_after = Pt(4)
        if _pslug.runs:
            _pslug.runs[0].font.size = Pt(9)
        doc.add_paragraph()
        _SEV_COLOURS_CMP = {"Low": "D1FAE5", "Moderate": "FEF3C7",
                            "Severe": "FEE2E2", "High": "FEE2E2"}
        _SCOLS = ["Seg","DN","Regime","Severity","f_slug (Hz)","f_slug (slugs/min)",
                  "V_slug (m/s)","H_Ls","L_slug (m)",
                  "ΔP_pulse (kPa)","ΔP_design (kPa)","F_elbow (N)","F_design (N)"]
        _SW = [0.28,0.38,0.90,0.68,0.52,0.72,0.58,0.38,0.55,0.68,0.72,0.58,0.62]
        _sev_idx_cmp = _SCOLS.index("Severity")
        def _slug_tbl(doc, records):
            if not records:
                doc.add_paragraph("No slug segments in this case.")
                return
            tbl = doc.add_table(rows=len(records)+1, cols=len(_SCOLS))
            tbl.style = "Table Grid"
            _style_header(tbl.rows[0], font_size=8)
            for j, col in enumerate(_SCOLS):
                tbl.rows[0].cells[j].text = col
            for i, rec in enumerate(records, start=1):
                row = tbl.rows[i]
                if i % 2 == 0:
                    for cell in row.cells:
                        _shd(cell, _ALT_BG)
                for j, col in enumerate(_SCOLS):
                    row.cells[j].text = str(rec.get(col, ""))
                    _cell_font(row.cells[j], size_pt=8)
                sev_val = rec.get("Severity", "")
                if sev_val in _SEV_COLOURS_CMP:
                    _shd(row.cells[_sev_idx_cmp], _SEV_COLOURS_CMP[sev_val])
            _set_col_widths(tbl, _SW)
        doc.add_heading(f"{label_a} — Slug Segments", level=2)
        _slug_tbl(doc, _slug_a)
        doc.add_paragraph()
        doc.add_heading(f"{label_b} — Slug Segments", level=2)
        _slug_tbl(doc, _slug_b)
        doc.add_paragraph()

    # ── 8. Visualisations ────────────────────────────────────────────────────
    if fig_cmp is not None or fig_bar is not None:
        doc.add_page_break()
        doc.add_heading("7. Visualisations", level=1)
        if fig_cmp is not None:
            doc.add_heading(f"Pressure Profiles — {label_a} vs {label_b}", level=2)
            _pcmp = doc.add_paragraph(
                f"The overlay shows absolute pressure (bara) vs. cumulative pipe distance for "
                f"{label_a} (solid line) and {label_b} (dashed line). "
                f"Both lines start at their respective inlet pressures; the vertical separation "
                f"at the outlet end reflects the difference in total ΔP between the two branches. "
                f"A steeper slope on either curve identifies the segments with the highest "
                f"resistance per unit length — the prime candidates for bore or fitting optimisation."
            )
            _pcmp.paragraph_format.space_after = Pt(4)
            if _pcmp.runs:
                _pcmp.runs[0].font.size = Pt(9)
            img = _fig_to_png(fig_cmp, width=900, height=400, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out)")
            doc.add_paragraph()
        if fig_bar is not None:
            doc.add_heading(f"ΔP by Segment — {label_a} vs {label_b}", level=2)
            _pbar = doc.add_paragraph(
                f"The bar chart decomposes total ΔP into individual segment contributions for "
                f"{label_a} and {label_b} side by side. "
                f"Tall bars identify the dominant loss segments. "
                f"A significant difference in bar height between the two branches at the same "
                f"segment position means their pipe specifications diverge there — this is "
                f"often the root cause of the hydraulic imbalance between the two lines."
            )
            _pbar.paragraph_format.space_after = Pt(4)
            if _pbar.runs:
                _pbar.runs[0].font.size = Pt(9)
            img = _fig_to_png(fig_bar, width=900, height=340, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out)")

    # ── 8. Method Sensitivity Analysis ───────────────────────────────────────
    if sensitivity_data is not None:
        _sa = sensitivity_data.get("sa", [])
        _sb = sensitivity_data.get("sb", [])
        _fig_s = sensitivity_data.get("fig")

        doc.add_page_break()
        doc.add_heading("8. Method Sensitivity Analysis", level=1)

        _intro = doc.add_paragraph(
            "All 12 method combinations (6 ΔP correlations × 2 void-fraction models) "
            "were evaluated to quantify uncertainty in the total pressure drop due to "
            "method selection.  Each combination runs the full pipeline with pressure "
            "marching.  Combinations that failed to converge are excluded."
        )
        if _intro.runs:
            _intro.runs[0].font.size = Pt(9)
        _intro_s2 = doc.add_paragraph(
            "The spread between minimum and maximum ΔP values defines the correlation "
            "uncertainty band for this service. "
            "If the two cases' ΔP ranges do not overlap, the relative ordering is "
            "unambiguous — one branch always has higher ΔP regardless of method choice. "
            "If the ranges overlap, method selection can change which branch governs; "
            "in that situation treat the spread as an additional design margin. "
            "The Flow Regime Consistency table below confirms whether the predicted "
            "regime in each segment is stable across all 12 combinations — an unstable "
            "regime (✗ flag) at a particular segment warrants extra scrutiny."
        )
        _intro_s2.paragraph_format.space_after = Pt(4)
        if _intro_s2.runs:
            _intro_s2.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        # Build summary table
        _CORR_SHORT = {
            "Beggs-Brill": "BB", "Friedel": "Friedel",
            "Lockhart_Martinelli": "L-M", "Muller_Steinhagen_Heck": "MSH",
            "Chisholm": "Chisholm", "Kim_Mudawar": "Kim-M",
        }
        _VOID_SHORT = {
            "Homogeneous": "Homo",
            "Rouhani-1 (slip)": "Rouhani-1",
        }
        _va = [r["total_dp_kpa"] for r in _sa if r["ok"]]
        _vb = [r["total_dp_kpa"] for r in _sb if r["ok"]]
        if _va and _vb:
            _a_min, _a_max = min(_va), max(_va)
            _b_min, _b_max = min(_vb), max(_vb)
            _a_sel = results_a["total_dp_kpa"]
            _b_sel = results_b["total_dp_kpa"]
            _overlap = _a_min <= _b_max and _b_min <= _a_max
            _sum_rows = [
                (f"{label_a} — minimum ΔP (kPa)",        f"{_a_min:.3f}",   "—"),
                (f"{label_a} — selected method (kPa)",   f"{_a_sel:.3f}",   "—"),
                (f"{label_a} — maximum ΔP (kPa)",        f"{_a_max:.3f}",   "—"),
                (f"{label_a} — spread (kPa)",             f"{_a_max-_a_min:.3f}", "—"),
                (f"{label_b} — minimum ΔP (kPa)",        "—",               f"{_b_min:.3f}"),
                (f"{label_b} — selected method (kPa)",   "—",               f"{_b_sel:.3f}"),
                (f"{label_b} — maximum ΔP (kPa)",        "—",               f"{_b_max:.3f}"),
                (f"{label_b} — spread (kPa)",             "—",               f"{_b_max-_b_min:.3f}"),
                ("Ranges overlap?",
                 "Yes — ordering depends on method" if _overlap else "No — unambiguous",
                 ""),
            ]
            _kv3_table(doc, _sum_rows, label_a=label_a, label_b=label_b)
            doc.add_paragraph()

        # Per-method ΔP detail table
        _detail_rows = []
        for _r_a, _r_b in zip(_sa, _sb):
            _c = _CORR_SHORT.get(_r_a["correlation"], _r_a["correlation"])
            _v = _VOID_SHORT.get(_r_a["voidage"], _r_a["voidage"])
            _label = f"{_c} / {_v}"
            _va_str = f"{_r_a['total_dp_kpa']:.3f}" if _r_a["ok"] else f"FAIL: {_r_a['error']}"
            _vb_str = f"{_r_b['total_dp_kpa']:.3f}" if _r_b["ok"] else f"FAIL: {_r_b['error']}"
            _detail_rows.append((_label, _va_str, _vb_str))
        _kv3_table(doc, _detail_rows, label_a=label_a, label_b=label_b)
        doc.add_paragraph()

        # Flow regime consistency tables
        doc.add_heading("Flow Regime Consistency", level=2)
        _reg_note = doc.add_paragraph(
            "Regime classification uses fixed Vsg, Vsl, and pipe angle — "
            "it is independent of ΔP correlation. Only the void fraction model (α) "
            "can shift vertical-segment thresholds (bubble/slug/churn). "
            "✓ = all 12 combinations predict the same regime for that segment."
        )
        if _reg_note.runs:
            _reg_note.runs[0].font.size = Pt(8)
            _reg_note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

        def _regime_report_table(doc, sens_results, segments_list, case_label):
            ok_results = [r for r in sens_results if r["ok"] and r.get("segment_regimes")]
            if not ok_results or not segments_list:
                return
            void_keys = list(dict.fromkeys(
                _VOID_SHORT.get(r["voidage"], r["voidage"]) for r in ok_results))
            cols = ["Seg", "Pipe", "Orientation"] + void_keys + ["Unanimous"]
            tbl = doc.add_table(rows=len(segments_list) + 1, cols=len(cols))
            tbl.style = "Table Grid"
            _style_header(tbl.rows[0], font_size=8)
            for j, col in enumerate(cols):
                tbl.rows[0].cells[j].text = col
            for i, seg in enumerate(segments_list):
                row = tbl.rows[i + 1]
                if i % 2 == 0:
                    for cell in row.cells:
                        _shd(cell, _ALT_BG)
                by_void = {}
                all_regimes = set()
                for r in ok_results:
                    v = _VOID_SHORT.get(r["voidage"], r["voidage"])
                    rg = r["segment_regimes"][i] if i < len(r["segment_regimes"]) else "—"
                    by_void.setdefault(v, set()).add(rg)
                    all_regimes.add(rg)
                row.cells[0].text = f"#{i+1}"
                dn_str = f"{seg.get('dn','?')}/{seg.get('pn','?')}"
                lined_str = (f" + {seg.get('liner_material','?')} {seg.get('liner_thickness_mm',1.0):.1f}mm"
                             if seg.get("lined") else "")
                row.cells[1].text = dn_str + lined_str
                row.cells[2].text = seg["type"]
                for j, v in enumerate(void_keys):
                    row.cells[3 + j].text = " | ".join(sorted(by_void.get(v, {"—"})))
                row.cells[3 + len(void_keys)].text = (
                    "✓" if len(all_regimes) == 1 else f"✗ ({len(all_regimes)})")
                for cell in row.cells:
                    _cell_font(cell, size_pt=8)
            _set_col_widths(tbl, [0.30, 0.70, 0.90] + [1.60] * len(void_keys) + [0.50])

        doc.add_paragraph(f"{label_a} — Flow Regime by Method")
        _regime_report_table(doc, _sa, results_a.get("segments", []), label_a)
        doc.add_paragraph()
        doc.add_paragraph(f"{label_b} — Flow Regime by Method")
        _regime_report_table(doc, _sb, results_b.get("segments", []), label_b)
        doc.add_paragraph()

        # Chart
        if _fig_s is not None:
            _img_s = _fig_to_png(_fig_s, width=900, height=480, scale=2)
            if _img_s:
                doc.add_picture(BytesIO(_img_s), width=Inches(6.2))
            else:
                doc.add_paragraph("(sensitivity chart rendering timed out)")

    # ── Stack ΔP ──────────────────────────────────────────────────────────────
    if stack_dp is not None:
        _gsh  = stack_dp.get("gsr_h2") or {}
        _gso  = stack_dp.get("gsr_o2") or {}
        _ph   = stack_dp.get("P_sep_h2")
        _po   = stack_dp.get("P_sep_o2")
        _la_s = stack_dp.get("label_a", label_a)
        _lb_s = stack_dp.get("label_b", label_b)
        _p_in_a  = _gsh.get("P_line_in", 0.0)
        _p_in_b  = _gso.get("P_line_in", 0.0)
        _dp_s    = _p_in_a - _p_in_b
        _dp_kpa  = _dp_s * 100.0
        _dp_mbar = _dp_kpa * 10.0

        doc.add_page_break()
        doc.add_heading("Generator Differential Pressure", level=1)

        _pgen_cmp = doc.add_paragraph(
            "The Generator (Electrolyser) Differential Pressure is the difference between "
            f"the {_la_s}-side and {_lb_s}-side branch inlet pressures, calculated at equal "
            "separator target pressures. "
            "Both gas streams share a common electrolyte membrane; any hydraulic imbalance "
            "between the two pipe systems appears directly as a mechanical pressure differential "
            "across that membrane. "
            "The inlet pressures are found by goal-seeking each system independently: both "
            "separators are set to their target operating pressure and the required inlet "
            "pressure is back-calculated through the full branch and header pressure drop. "
            "The resulting difference is a purely hydraulic quantity — it excludes "
            "electrochemical overpotentials — and represents the piping-induced membrane load."
        )
        _pgen_cmp.paragraph_format.space_after = Pt(4)
        if _pgen_cmp.runs:
            _pgen_cmp.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        doc.add_heading("Target Conditions", level=2)
        _kv_table(doc, [
            (f"{_la_s} system — separator target pressure (bara)", f"{_ph:.3f}" if _ph is not None else "—"),
            (f"{_lb_s} system — separator target pressure (bara)", f"{_po:.3f}" if _po is not None else "—"),
        ])
        doc.add_paragraph()

        doc.add_heading(f"{_la_s} System  (Branch → Header C → Separator)", level=2)
        _kv_table(doc, [
            (f"{_la_s} line inlet pressure (bara)",       f"{_gsh.get('P_line_in', 0):.4f}"),
            (f"{_la_s} line ΔP (kPa)",                    f"{_gsh.get('dp_line', 0):.3f}"),
            (f"{_la_s} outlet / Header C inlet (bara)",   f"{_gsh.get('P_line_out', 0):.4f}"),
            ("Header C + T-seg ΔP (kPa)",                 f"{_gsh.get('dp_hdr', 0):.3f}"),
            (f"{_la_s} system separator pressure (bara)", f"{_gsh.get('P_sep', 0):.4f}"),
        ])
        doc.add_paragraph()

        doc.add_heading(f"{_lb_s} System  (Branch → Header D → Separator)", level=2)
        _kv_table(doc, [
            (f"{_lb_s} line inlet pressure (bara)",       f"{_gso.get('P_line_in', 0):.4f}"),
            (f"{_lb_s} line ΔP (kPa)",                    f"{_gso.get('dp_line', 0):.3f}"),
            (f"{_lb_s} outlet / Header D inlet (bara)",   f"{_gso.get('P_line_out', 0):.4f}"),
            ("Header D + T-seg ΔP (kPa)",                 f"{_gso.get('dp_hdr', 0):.3f}"),
            (f"{_lb_s} system separator pressure (bara)", f"{_gso.get('P_sep', 0):.4f}"),
        ])
        doc.add_paragraph()

        doc.add_heading(f"Generator ΔP Result  (P_inlet_{_la_s} − P_inlet_{_lb_s})", level=2)
        _kv_table(doc, [
            (f"ΔP  {_la_s} − {_lb_s}  (bara)",  f"{_dp_s:.4f}"),
            (f"ΔP  {_la_s} − {_lb_s}  (kPa)",   f"{_dp_kpa:.2f}"),
            (f"ΔP  {_la_s} − {_lb_s}  (mbar)",  f"{_dp_mbar:.1f}"),
        ])
        doc.add_paragraph()

    # ── Engineering note ──────────────────────────────────────────────────────
    doc.add_paragraph()
    _sp_a = " / ".join((results_a.get("gas_flows_kgh") or {}).keys()) or (results_a.get("liquid_type") or "—")
    _sp_b = " / ".join((results_b.get("gas_flows_kgh") or {}).keys()) or (results_b.get("liquid_type") or "—")
    _sp_str = _sp_a if _sp_a == _sp_b else f"{_sp_a}  |  {_sp_b}"
    note = doc.add_paragraph(
        f"Engineering Note: The correlations used here carry an estimated uncertainty of "
        f"±20–30 % for two-phase flow and ±10–15 % for single-phase flow. "
        f"Their application to this service ({_sp_str}) should be validated against "
        f"commissioning data before use in safety-critical design. "
        f"Use the sensitivity analysis (if present) to bracket the ΔP range."
    )
    if note.runs:
        note.runs[0].font.size      = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================================
# COMBINED REPORT  (Case A + B + optional C, comparison, sensitivity)
# ============================================================================

def generate_combined_report(
    cases,
    case_labels=None,
    fig_cmp=None,
    fig_bar=None,
    sensitivity_data=None,
    stack_dp=None,
    dn_study_data=None,
):
    """
    Generate a single Word report combining all cases, comparison, and sensitivity.

    cases       : list of result dicts from run_case() – 2 or 3 items.
    case_labels : e.g. ["Case A", "Case B", "Case C"].
    fig_cmp     : overlaid pressure-profile figure (A vs B).
    fig_bar     : per-segment ΔP bar chart (A vs B).
    sensitivity_data : dict {"sa": [...], "sb": [...], "fig": Figure} or None.
    """
    n = len(cases)
    if case_labels is None:
        case_labels = [f"Case {chr(65 + i)}" for i in range(n)]

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27);  sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9);   sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9);   sec.bottom_margin = Inches(0.9)

    # ── Title ────────────────────────────────────────────────────────────────
    h = doc.add_heading("Gas–Liquid Piping Hydraulic Study — Combined Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "  ·  ".join(case_labels)
        + f"  ·  {datetime.now().strftime('%d %B %Y  %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _add_footer_page_numbers(doc)
    _add_toc(doc)

    # ── Helpers ───────────────────────────────────────────────────────────────
    _sec = [0]
    def _h1(title):
        _sec[0] += 1
        doc.add_heading(f"{_sec[0]}. {title}", level=1)

    # Column widths fit within 6.47" text area
    if n >= 4:
        _col_w = (1.6, 1.2, 1.2, 1.2, 1.2)
    elif n >= 3:
        _col_w = (1.9, 1.5, 1.5, 1.5)
    else:
        _col_w = (2.2, 2.1, 2.1)
    _headers = ["Parameter"] + case_labels[:n]

    def _nt(rows):
        return _kv_n_table(doc, _headers, rows, col_widths=_col_w)

    def _body_para(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        if p.runs:
            p.runs[0].font.size = Pt(9)
        return p

    # Identify header vs branch cases once up front
    _is_hdr = [not c.get("segments") for c in cases]

    # ════════════════════════════════════════════════════════════════════════
    # 1. PURPOSE
    # ════════════════════════════════════════════════════════════════════════
    _h1("Purpose")
    _lbl_branch = "  ·  ".join(case_labels[:2]) if n >= 2 else case_labels[0]
    _lbl_header = "  ·  ".join(case_labels[2:]) if n >= 3 else ""
    _p_lines = [
        f"This study sizes the individual branch pipelines ({_lbl_branch}) and "
        f"{'collecting headers (' + _lbl_header + ') ' if _lbl_header else ''}"
        f"for a gas–liquid piping system. Each branch carries a two-phase mixture "
        f"from a process unit to a collecting header, which conveys the combined "
        f"flow to the gas–liquid separator."
    ]
    if n >= 4:
        _p_lines.append(
            f"Two independent collecting systems are evaluated: "
            f"{case_labels[0]} branches feed {case_labels[2]}, and "
            f"{case_labels[1]} branches feed {case_labels[3]}. "
            f"The goal-seek function finds the required branch inlet pressures for each "
            f"system such that each separator arrives at its target operating pressure. "
            f"The difference between the two branch inlet pressures is the differential "
            f"pressure across the upstream process unit."
        )
    elif n >= 3:
        _p_lines.append(
            f"The goal-seek function finds the required branch-line inlet pressure "
            f"such that the separator arrives at the target operating pressure. "
            f"For a two-line system, the difference between the two branch inlet "
            f"pressures gives the differential pressure across the upstream process unit."
        )
    for _txt in _p_lines:
        _body_para(_txt)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 2. KEY RESULTS
    # ════════════════════════════════════════════════════════════════════════
    _h1("Key Results")

    # Per-case summary row (inlet P, total ΔP, outlet/separator P)
    _kr_rows = [
        ("Inlet / tap inlet pressure (bara)",)
            + tuple(f"{c['P_bara']:.4f}" for c in cases),
        ("Total ΔP (kPa)",)
            + tuple(f"{c['total_dp_kpa']:.3f}" for c in cases),
        ("Outlet / separator pressure (bara)",)
            + tuple(f"{c['outlet_pressure_bara']:.4f}" for c in cases),
    ]
    _nt(_kr_rows)
    doc.add_paragraph()

    # Combined system totals (branch + header)
    if n >= 3:
        _systems = []
        if n >= 3 and not _is_hdr[2]:
            pass  # case 2 is not a header — skip
        elif n >= 3 and _is_hdr[2]:
            _dp_br_1  = cases[0]["total_dp_kpa"]
            _dp_hd_1  = cases[2]["total_dp_kpa"]
            _p_sep_1  = cases[2].get("P_separator_bara", cases[2]["outlet_pressure_bara"])
            _systems.append((f"{case_labels[0]}+{case_labels[2]}",
                              _dp_br_1, _dp_hd_1, _dp_br_1 + _dp_hd_1, _p_sep_1))
        if n >= 4 and _is_hdr[3]:
            _dp_br_2  = cases[1]["total_dp_kpa"]
            _dp_hd_2  = cases[3]["total_dp_kpa"]
            _p_sep_2  = cases[3].get("P_separator_bara", cases[3]["outlet_pressure_bara"])
            _systems.append((f"{case_labels[1]}+{case_labels[3]}",
                              _dp_br_2, _dp_hd_2, _dp_br_2 + _dp_hd_2, _p_sep_2))
        if _systems:
            doc.add_heading("Combined System ΔP", level=2)
            if len(_systems) == 1:
                _sc_hdrs = ["Parameter", _systems[0][0]]
                _sc_cw   = (3.5, 3.0)
                _sc_data = [
                    ("Branch ΔP (kPa)",                              f"{_systems[0][1]:.3f}"),
                    ("Header ΔP — worst arm + T-seg (kPa)",          f"{_systems[0][2]:.3f}"),
                    ("System total ΔP (kPa)",                        f"{_systems[0][3]:.3f}"),
                    ("Separator pressure (bara)",                     f"{_systems[0][4]:.4f}"),
                ]
            else:
                _sc_hdrs = ["Parameter"] + [s[0] for s in _systems]
                _sc_cw   = (2.47, 2.0, 2.0)
                _sc_data = [
                    ("Branch ΔP (kPa)",)
                        + tuple(f"{s[1]:.3f}" for s in _systems),
                    ("Header ΔP — worst arm + T-seg (kPa)",)
                        + tuple(f"{s[2]:.3f}" for s in _systems),
                    ("System total ΔP (kPa)",)
                        + tuple(f"{s[3]:.3f}" for s in _systems),
                    ("Separator pressure (bara)",)
                        + tuple(f"{s[4]:.4f}" for s in _systems),
                ]
            _kv_n_table(doc, _sc_hdrs, _sc_data, col_widths=_sc_cw)
            doc.add_paragraph()

    # Generator ΔP summary (branch A inlet − branch B inlet)
    if n >= 2 and not _is_hdr[0] and not _is_hdr[1]:
        _p_in_a   = cases[0]["P_bara"]
        _p_in_b   = cases[1]["P_bara"]
        _gen_dp   = _p_in_a - _p_in_b
        _gen_kpa  = _gen_dp * 100.0
        _gen_mbar = _gen_kpa * 10.0
        doc.add_heading("Generator Differential Pressure", level=2)
        _pgen_kr = doc.add_paragraph(
            f"The Generator ΔP is the pressure difference between the {case_labels[0]} and "
            f"{case_labels[1]} branch inlet ports, evaluated at equal separator target pressures. "
            f"Both systems share a common electrolyte membrane; this hydraulic imbalance "
            f"appears directly as a differential pressure across that membrane. "
            f"The design objective is to minimise the absolute value of this differential."
        )
        _pgen_kr.paragraph_format.space_after = Pt(4)
        if _pgen_kr.runs:
            _pgen_kr.runs[0].font.size = Pt(9)
        _kv_table(doc, [
            (f"{case_labels[0]} branch inlet pressure (bara)",              f"{_p_in_a:.4f}"),
            (f"{case_labels[1]} branch inlet pressure (bara)",              f"{_p_in_b:.4f}"),
            (f"Generator ΔP  [{case_labels[0]} − {case_labels[1]}]  (bara)", f"{_gen_dp:.4f}"),
            (f"Generator ΔP  [{case_labels[0]} − {case_labels[1]}]  (kPa)",  f"{_gen_kpa:.2f}"),
            (f"Generator ΔP  [{case_labels[0]} − {case_labels[1]}]  (mbar)", f"{_gen_mbar:.1f}"),
        ], col_widths=(4.0, 2.47))
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 3. METHOD
    # ════════════════════════════════════════════════════════════════════════
    _h1("Method")
    for _txt in [
        ("Six two-phase ΔP correlations are available: Beggs & Brill (1973, default), "
         "Friedel, Lockhart-Martinelli, Müller-Steinhagen & Heck, Chisholm, and "
         "Kim-Mudawar. Cases may use different correlations and void-fraction models. "
         "The pipeline is divided into user-defined segments with pressure marching — "
         "gas density is re-evaluated at each segment inlet. Each segment's ΔP is "
         "decomposed into frictional, gravitational, and accelerational components."),
        ("Void fraction: homogeneous model or Rouhani-1 slip-flow model. "
         "Flow regime classified automatically — Taitel-Dukler + Mandhane-Gregory-Aziz "
         "for horizontal, Wallis/Taitel (1980) for vertical segments. "
         "Gas properties: ideal-gas law, CoolProp viscosities, Dalton's Law for water "
         "vapour. Minor losses: Crane TP-410. Erosion: API RP 14E, C = 100. "
         "Packages: fluids · CoolProp · python-docx."),
    ]:
        _body_para(_txt)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 4. PROCESS CONDITIONS
    # ════════════════════════════════════════════════════════════════════════
    _h1("Process Conditions")
    _all_sp = list(dict.fromkeys(sp for c in cases for sp in (c.get("gas_flows_kgh") or {})))
    _cond = [
        ("Flow mode",)               + tuple(
            (c.get("flow_mode") or "gas_liquid").replace("_", " ").title() for c in cases),
        ("Inlet pressure (bara)",)   + tuple(f"{c['P_bara']:.2f}"  for c in cases),
        ("Temperature (°C)",)        + tuple(
            f"{c['T_C']:.1f}" if c.get("T_C") is not None
            else f"{c.get('props',{}).get('T_sat_C', 0.0):.1f}  (T_sat)"
            for c in cases),
    ]
    for _sp in _all_sp:
        _cond.append(
            (f"{_sp} mass flow (kg/h)",) +
            tuple(f"{(c.get('gas_flows_kgh') or {}).get(_sp, 0.0):.3f}" for c in cases))
    _cond += [
        ("Fluid / Liquid type",)     + tuple(
            c.get("liquid_type") or c.get("props", {}).get("vle_fluid") or "—"
            for c in cases),
        ("Liquid vol. flow (m³/h)",) + tuple(f"{c.get('q_lye', 0.0):.3f}" for c in cases),
        ("ΔP correlation",)          + tuple(c.get("correlation", "—")     for c in cases),
        ("Void fraction model",)     + tuple(c.get("voidage_method", "—")  for c in cases),
        ("Segments / taps",)         + tuple(
            str(len(c["segments"])) if c.get("segments")
            else f"{c.get('n_left',0)}L + {c.get('n_right',0)}R taps"
            for c in cases),
    ]
    _nt(_cond)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 5. BRANCH LINE RESULTS
    # ════════════════════════════════════════════════════════════════════════
    _branch_cases  = [c   for c, h in zip(cases, _is_hdr) if not h]
    _branch_labels = [lbl for lbl, h in zip(case_labels, _is_hdr) if not h]
    if _branch_cases:
        _h1("Branch Line Results")
        nb = len(_branch_cases)
        if nb >= 3:
            _bcw = (1.9, 1.5, 1.5, 1.5)[:nb + 1]
        elif nb == 2:
            _bcw = (2.5, 2.0, 2.0)
        else:
            _bcw = (3.0, 3.47)
        _bhdrs = ["Parameter"] + _branch_labels
        _max_ve = [
            max((r.get("V_m/V_e", 0) for r in c["grid_records"]), default=0.0)
            for c in _branch_cases
        ]
        _br_rows = [
            ("Inlet pressure (bara)",)
                + tuple(f"{c['P_bara']:.4f}"           for c in _branch_cases),
            ("Total ΔP (kPa)",)
                + tuple(f"{c['total_dp_kpa']:.4f}"     for c in _branch_cases),
            ("  ↳ Frictional ΔP (kPa)",)
                + tuple(f"{c['total_dp_fric_kpa']:.4f}" for c in _branch_cases),
            ("  ↳ Gravitational ΔP (kPa)",)
                + tuple(f"{c['total_dp_grav_kpa']:.4f}" for c in _branch_cases),
            ("Outlet pressure (bara)",)
                + tuple(f"{c['outlet_pressure_bara']:.4f}" for c in _branch_cases),
            ("Pipe length (m)",)
                + tuple(f"{c['pipe_length_m']:.2f}"    for c in _branch_cases),
            ("Worst V_m/V_e (–)",)
                + tuple(f"{v:.3f}"                      for v in _max_ve),
        ]
        if nb >= 2:
            _dp_delta = _branch_cases[1]["total_dp_kpa"] - _branch_cases[0]["total_dp_kpa"]
            _br_rows.append(
                (f"ΔP  {_branch_labels[1]} − {_branch_labels[0]}  (kPa)",
                 "—", f"{_dp_delta:+.4f}") + ("—",) * max(0, nb - 2)
            )
        _kv_n_table(doc, _bhdrs, _br_rows, col_widths=_bcw)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 6. COMBINED SYSTEM ΔP — one sub-section per header case
    # ════════════════════════════════════════════════════════════════════════
    _hdr_triples = [
        (idx, c, lbl)
        for idx, (c, lbl, h) in enumerate(zip(cases, case_labels, _is_hdr)) if h
    ]
    if _hdr_triples:
        _h1("Combined System ΔP — Header + Branch")
        for _hidx, _hc, _hlbl in _hdr_triples:
            # Header at index 2 pairs with branch at 0; index 3 pairs with branch at 1
            _br_idx = _hidx - 2
            if 0 <= _br_idx < n and not _is_hdr[_br_idx]:
                _bc, _blbl = cases[_br_idx], case_labels[_br_idx]
            else:
                _bc, _blbl = cases[0], case_labels[0]

            _dp_hdr   = _hc["total_dp_kpa"]
            _dp_br    = _bc["total_dp_kpa"]
            _p_in_br  = _bc["P_bara"]
            _p_out_br = _bc["outlet_pressure_bara"]
            _p_sep    = _hc.get("P_separator_bara", _hc["outlet_pressure_bara"])

            doc.add_heading(f"{_blbl} → {_hlbl} → Separator", level=2)
            _sys_p = doc.add_paragraph(
                f"The {_hlbl} collecting header receives flow from {_blbl} branch lines "
                f"and delivers the combined flow to the separator via the T-segment. "
                f"Equal flow per tap is assumed; the governing (highest-ΔP) arm sets the "
                f"required tap inlet pressure. This is conservative when header ΔP ≪ branch ΔP."
            )
            if _sys_p.runs:
                _sys_p.runs[0].font.size = Pt(9)
            doc.add_paragraph()
            _kv_table(doc, [
                (f"Branch inlet pressure — {_blbl} (bara)",          f"{_p_in_br:.4f}"),
                (f"Branch ΔP — {_blbl} (kPa)",                       f"{_dp_br:.4f}"),
                ("Branch outlet / header tap inlet pressure (bara)",  f"{_p_out_br:.4f}"),
                (f"Header ΔP (worst arm + T-seg) — {_hlbl} (kPa)",   f"{_dp_hdr:.4f}"),
                (f"Total system ΔP — {_blbl} + {_hlbl} (kPa)",       f"{_dp_br + _dp_hdr:.4f}"),
                ("Separator connection pressure (bara)",               f"{_p_sep:.4f}"),
            ])
            doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # 7. VISUALISATIONS  (page break)
    # ════════════════════════════════════════════════════════════════════════
    _has_figs = (fig_cmp is not None or fig_bar is not None
                 or any(c.get("fig_sch") or c.get("fig_prof") for c in cases))
    if _has_figs:
        doc.add_page_break()
        _h1("Visualisations")

        if fig_cmp is not None:
            _la_v = case_labels[0] if case_labels else "A"
            _lb_v = case_labels[1] if len(case_labels) > 1 else "B"
            doc.add_heading(f"Pressure Profiles — {_la_v} vs {_lb_v}", level=2)
            img = _fig_to_png(fig_cmp, width=900, height=400, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out)")
            _fig_caption(doc,
                f"Figure: Absolute pressure (bara) along the pipeline for {_la_v} (solid) "
                f"and {_lb_v} (dashed). A steeper slope indicates higher local resistance.")

        if fig_bar is not None:
            _la_v = case_labels[0] if case_labels else "A"
            _lb_v = case_labels[1] if len(case_labels) > 1 else "B"
            doc.add_heading(f"ΔP by Segment — {_la_v} vs {_lb_v}", level=2)
            img = _fig_to_png(fig_bar, width=900, height=340, scale=2)
            if img:
                doc.add_picture(BytesIO(img), width=Inches(6.2))
            else:
                doc.add_paragraph("(chart rendering timed out)")
            _fig_caption(doc,
                f"Figure: Pressure drop (kPa) per segment for {_la_v} and {_lb_v}. "
                f"Tallest bars are the dominant loss elements.")

        for c, lbl, is_hdr in zip(cases, case_labels, _is_hdr):
            _fs = c.get("fig_sch")
            _fp = c.get("fig_prof")
            if _fs is not None or _fp is not None:
                doc.add_heading(f"{lbl} — {'Header Layout' if is_hdr else 'Pipeline'}", level=2)
            if _fs is not None:
                img = _fig_to_png(_fs, width=900, height=340 if is_hdr else 520, scale=2)
                if img:
                    doc.add_picture(BytesIO(img), width=Inches(6.2))
                else:
                    doc.add_paragraph("(chart rendering timed out)")
                if is_hdr:
                    _fig_caption(doc,
                        f"Figure: Header piping layout for {lbl}. "
                        f"Blue = left arm, orange = right arm. Triangular markers = tap risers "
                        f"with distance from T. Flow arrows point toward T-junction. "
                        f"Thicker dark pipe = T-segment to separator. "
                        f"Governing arm (⚠) sets the required tap inlet pressure.")
                else:
                    _fig_caption(doc,
                        f"Figure: Pipeline schematic for {lbl}, colour-coded by flow regime. "
                        f"V_m/V_e > 1.0 flags erosion risk (API RP 14E, C = 100).")
            if _fp is not None:
                img = _fig_to_png(_fp, width=900, height=320 if is_hdr else 400, scale=2)
                if img:
                    doc.add_picture(BytesIO(img), width=Inches(6.2))
                else:
                    doc.add_paragraph("(chart rendering timed out)")
                if is_hdr:
                    _fig_caption(doc,
                        f"Figure: Header pressure profile for {lbl}. Left arm (blue) and "
                        f"right arm (orange) pressure vs. distance from T-junction. "
                        f"X-axis reversed so flow runs left to right toward T at x = 0.")
                else:
                    _fig_caption(doc,
                        f"Figure: Pressure profile for {lbl}. Pressure (bara) vs. cumulative "
                        f"distance. Coloured bands show predicted flow regime per segment.")

    # ════════════════════════════════════════════════════════════════════════
    # 8. SENSITIVITY ANALYSIS  (page break, if provided)
    # ════════════════════════════════════════════════════════════════════════
    if sensitivity_data is not None:
        _sa = sensitivity_data.get("sa", [])
        _sb = sensitivity_data.get("sb", [])
        _fig_s = sensitivity_data.get("fig")

        doc.add_page_break()
        _h1("Method Sensitivity Analysis")

        _intro = doc.add_paragraph(
            "All 12 method combinations (6 ΔP correlations × 2 void-fraction models) "
            "were evaluated to quantify uncertainty due to correlation choice. "
            "Combinations that failed to converge are excluded."
        )
        if _intro.runs:
            _intro.runs[0].font.size = Pt(9)
        _intro_c2 = doc.add_paragraph(
            "The spread between minimum and maximum ΔP defines the correlation uncertainty "
            "band for this specific service. "
            "A narrow spread (below ~20 % of the mean) indicates robust agreement across "
            "methods; a wide spread signals that the result is method-sensitive and that "
            "further validation — e.g. against vendor data or commissioning measurements — "
            "is advisable. "
            "If the two cases' ΔP ranges do not overlap, the ordering (which line has the "
            "higher ΔP) is unambiguous regardless of method choice; if they overlap, "
            "treat the full sensitivity range as the design uncertainty band. "
            "The chart at the end of this section shows all 12 results visually."
        )
        _intro_c2.paragraph_format.space_after = Pt(4)
        if _intro_c2.runs:
            _intro_c2.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        _CORR_S = {"Beggs-Brill": "BB", "Friedel": "Friedel",
                   "Lockhart_Martinelli": "L-M", "Muller_Steinhagen_Heck": "MSH",
                   "Chisholm": "Chisholm", "Kim_Mudawar": "Kim-M"}
        _VOID_S = {"Homogeneous": "Homo", "Rouhani-1 (slip)": "Rouhani-1"}

        _lbl_sa = case_labels[0] if len(case_labels) > 0 else "Case A"
        _lbl_sb = case_labels[1] if len(case_labels) > 1 else "Case B"

        _va = [r["total_dp_kpa"] for r in _sa if r["ok"]]
        _vb = [r["total_dp_kpa"] for r in _sb if r["ok"]]
        if _va and _vb:
            _a_sel  = cases[0]["total_dp_kpa"]
            _b_sel  = cases[1]["total_dp_kpa"]
            _overlap = min(_va) <= max(_vb) and min(_vb) <= max(_va)
            _kv3_table(doc, [
                (f"{_lbl_sa} — min ΔP (kPa)",    f"{min(_va):.3f}",  "—"),
                (f"{_lbl_sa} — selected (kPa)",  f"{_a_sel:.3f}",    "—"),
                (f"{_lbl_sa} — max ΔP (kPa)",    f"{max(_va):.3f}",  "—"),
                (f"{_lbl_sa} — spread (kPa)",     f"{max(_va)-min(_va):.3f}", "—"),
                (f"{_lbl_sb} — min ΔP (kPa)",    "—", f"{min(_vb):.3f}"),
                (f"{_lbl_sb} — selected (kPa)",  "—", f"{_b_sel:.3f}"),
                (f"{_lbl_sb} — max ΔP (kPa)",    "—", f"{max(_vb):.3f}"),
                (f"{_lbl_sb} — spread (kPa)",     "—", f"{max(_vb)-min(_vb):.3f}"),
                ("Ranges overlap?",
                 "Yes — ordering method-dependent" if _overlap else "No — unambiguous", ""),
            ], label_a=_lbl_sa, label_b=_lbl_sb)
            doc.add_paragraph()

        _det = []
        for _ra, _rb in zip(_sa, _sb):
            _c = _CORR_S.get(_ra["correlation"], _ra["correlation"])
            _v = _VOID_S.get(_ra["voidage"], _ra["voidage"])
            _det.append((
                f"{_c} / {_v}",
                f"{_ra['total_dp_kpa']:.3f}" if _ra["ok"] else f"FAIL: {_ra['error']}",
                f"{_rb['total_dp_kpa']:.3f}" if _rb["ok"] else f"FAIL: {_rb['error']}",
            ))
        _kv3_table(doc, _det, label_a=_lbl_sa, label_b=_lbl_sb)
        doc.add_paragraph()

        if _fig_s is not None:
            _img_s = _fig_to_png(_fig_s, width=900, height=480, scale=2)
            if _img_s:
                doc.add_picture(BytesIO(_img_s), width=Inches(6.2))
            else:
                doc.add_paragraph("(sensitivity chart rendering timed out)")
            _fig_caption(doc,
                f"Figure: Total ΔP for all 12 method combinations. "
                f"Spread quantifies the correlation-choice uncertainty band. "
                f"Non-overlapping clusters give an unambiguous result.")

    # ════════════════════════════════════════════════════════════════════════
    # 9. GENERATOR ΔP DETAIL  (page break, if goal-seek data available)
    # ════════════════════════════════════════════════════════════════════════
    if stack_dp is not None:
        _gsh  = stack_dp.get("gsr_h2") or {}
        _gso  = stack_dp.get("gsr_o2") or {}
        _ph   = stack_dp.get("P_sep_h2")
        _po   = stack_dp.get("P_sep_o2")
        _la_s = stack_dp.get("label_a", case_labels[0] if case_labels else "Case A")
        _lb_s = stack_dp.get("label_b", case_labels[1] if len(case_labels) > 1 else "Case B")
        _p_in_a  = _gsh.get("P_line_in", 0.0)
        _p_in_b  = _gso.get("P_line_in", 0.0)
        _dp_s    = _p_in_a - _p_in_b
        _dp_kpa  = _dp_s * 100.0
        _dp_mbar = _dp_kpa * 10.0

        doc.add_page_break()
        _h1("Generator Differential Pressure — Detail")

        _pgen_det = doc.add_paragraph(
            "This section traces the full pressure path from each branch inlet, through the "
            f"respective header, to the separator for both the {_la_s} and {_lb_s} systems. "
            "The goal-seek algorithm fixes both separator pressures at their targets and "
            "back-calculates the required branch inlet pressure for each system independently. "
            "The difference between the two inlet pressures is the Generator ΔP — the net "
            "hydraulic load on the electrolyser membrane arising solely from piping asymmetry. "
            "It is reported in bara, kPa, and mbar to align with typical membrane "
            "pressure-rating and instrumentation documentation."
        )
        _pgen_det.paragraph_format.space_after = Pt(4)
        if _pgen_det.runs:
            _pgen_det.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        doc.add_heading("Target Conditions", level=2)
        _kv_table(doc, [
            (f"{_la_s} separator target pressure (bara)", f"{_ph:.3f}" if _ph is not None else "—"),
            (f"{_lb_s} separator target pressure (bara)", f"{_po:.3f}" if _po is not None else "—"),
        ])
        doc.add_paragraph()

        for _la_x, _gs_x in [(_la_s, _gsh), (_lb_s, _gso)]:
            doc.add_heading(f"{_la_x}  (Branch → Header → Separator)", level=2)
            _kv_table(doc, [
                (f"{_la_x} branch inlet pressure (bara)",    f"{_gs_x.get('P_line_in', 0):.4f}"),
                (f"{_la_x} branch ΔP (kPa)",                 f"{_gs_x.get('dp_line', 0):.3f}"),
                ("Branch outlet / header tap inlet (bara)",  f"{_gs_x.get('P_line_out', 0):.4f}"),
                ("Header + T-seg ΔP (kPa)",                  f"{_gs_x.get('dp_hdr', 0):.3f}"),
                ("Separator pressure (bara)",                 f"{_gs_x.get('P_sep', 0):.4f}"),
            ])
            doc.add_paragraph()

        doc.add_heading(f"Generator ΔP  ({_la_s} − {_lb_s})", level=2)
        _kv_table(doc, [
            (f"ΔP  {_la_s} − {_lb_s}  (bara)",  f"{_dp_s:.4f}"),
            (f"ΔP  {_la_s} − {_lb_s}  (kPa)",   f"{_dp_kpa:.2f}"),
            (f"ΔP  {_la_s} − {_lb_s}  (mbar)",  f"{_dp_mbar:.1f}"),
        ])
        doc.add_paragraph()

    # ── Engineering note (end of main body) ──────────────────────────────────
    _all_gas = sorted({sp for c in cases for sp in (c.get("gas_flows_kgh") or {})})
    _all_liq = sorted({c.get("liquid_type") or c.get("props",{}).get("vle_fluid") or "—" for c in cases})
    _srv = " / ".join(_all_gas + _all_liq) if _all_gas else " / ".join(_all_liq)
    note = doc.add_paragraph(
        f"Engineering Note: The correlations used here carry an estimated uncertainty of "
        f"±20–30 % for two-phase flow and ±10–15 % for single-phase flow. "
        f"Their application to this service ({_srv}) should be validated against "
        f"commissioning data before use in safety-critical design. "
        f"Use the sensitivity analysis (if present) to bracket the ΔP range."
    )
    if note.runs:
        note.runs[0].font.size      = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ════════════════════════════════════════════════════════════════════════
    # APPENDIX
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _app_h = doc.add_heading("Appendix", level=0)
    _app_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── A. Phase Thermodynamics ───────────────────────────────────────────────
    doc.add_heading("A.  Phase Thermodynamics  (inlet conditions)", level=1)
    _ps = [c["props"] for c in cases]

    def _pfmt(p, key, scale=1.0, fmt=".4f"):
        v = p.get(key)
        return f"{v * scale:{fmt}}" if v is not None else "—"

    _thm = [
        ("Gas density ρ_g (kg/m³)",)     + tuple(_pfmt(p, "rho_g")                  for p in _ps),
        ("Gas mixture MW (g/mol)",)       + tuple(_pfmt(p, "MW_mix_gmol", fmt=".3f") for p in _ps),
        ("Liquid density ρ_l (kg/m³)",)   + tuple(_pfmt(p, "rho_l", fmt=".2f")       for p in _ps),
        ("Liquid viscosity μ_l (mPa·s)",) + tuple(_pfmt(p, "mu_l", 1e3)             for p in _ps),
        ("Gas viscosity μ_g (µPa·s)",)   + tuple(_pfmt(p, "mu_g", 1e6, ".2f")       for p in _ps),
        ("Surface tension σ (mN/m)",)    + tuple(_pfmt(p, "sigma", 1e3, ".3f")       for p in _ps),
        ("Mass quality x (%)",)          + tuple(_pfmt(p, "x_gas", 100)             for p in _ps),
        ("Void fraction α (%)",)         + tuple(_pfmt(p, "alpha", 100, ".2f")       for p in _ps),
    ]
    if any(p.get("P_sat_H2O_pa", 0) > 0 for p in _ps):
        _thm += [
            ("H₂O sat. pressure (bara)",) + tuple(
                f"{p.get('P_sat_H2O_pa', 0) / 1e5:.4f}"
                if p.get("P_sat_H2O_pa", 0) > 0 else "—" for p in _ps),
            ("H₂O vapour flow (kg/h)",) + tuple(
                f"{p.get('m_vapor_h2o_kgh', 0):.4f}"
                if p.get("P_sat_H2O_pa", 0) > 0 else "—" for p in _ps),
        ]
    _nt(_thm)
    doc.add_paragraph()

    # ── B. Branch Segment Analysis ────────────────────────────────────────────
    _branch_pairs = [(c, lbl) for c, lbl, h in zip(cases, case_labels, _is_hdr) if not h]
    if _branch_pairs:
        doc.add_heading("B.  Segment Analysis — Branch Lines", level=1)
        _SC = ["Seg", "Pipe", "ID (mm)", "Type", "L (m)", "L_eq (m)", "Fittings",
               "Regime", "V_m (m/s)", "V_m/V_e", "ΔP (kPa)", "P_out (bara)"]
        _SW = [0.25, 0.50, 0.42, 0.80, 0.38, 0.42, 0.70, 0.85, 0.48, 0.44, 0.50, 0.53]

        def _seg_tbl(records):
            if not records:
                return
            tbl = doc.add_table(rows=len(records) + 1, cols=len(_SC))
            tbl.style = "Table Grid"
            _style_header(tbl.rows[0], font_size=8)
            for j, col in enumerate(_SC):
                tbl.rows[0].cells[j].text = col
            for i, rec in enumerate(records, start=1):
                row = tbl.rows[i]
                if i % 2 == 0:
                    for cell in row.cells:
                        _shd(cell, _ALT_BG)
                for j, col in enumerate(_SC):
                    row.cells[j].text = str(rec.get(col, ""))
                    _cell_font(row.cells[j], size_pt=8)
            _set_col_widths(tbl, _SW)

        for _bc, _blbl in _branch_pairs:
            doc.add_heading(_blbl, level=2)
            _seg_tbl(_bc["grid_records"])
            doc.add_paragraph()

    # ── B2. Slug Flow Dynamics (branch lines with slug segments) ─────────────
    _slug_branch_pairs = [
        (c, lbl) for c, lbl, h in zip(cases, case_labels, _is_hdr)
        if not h and (c.get("slug_records") or [])
    ]
    if _slug_branch_pairs:
        doc.add_heading("B2.  Slug Flow Dynamics — Branch Lines", level=1)
        _pb2 = doc.add_paragraph(
            "One or more branch segments were classified as slug or intermittent flow. "
            "Slug frequency (Gregory-Scott 1969), translational velocity (Bendiksen 1984), "
            "liquid holdup (Gregory et al. 1978), and 90° elbow pressure pulse / force "
            "(momentum balance, ASME B31.3 DLF = 2.0) are tabulated below per branch."
        )
        _pb2.paragraph_format.space_after = Pt(4)
        if _pb2.runs:
            _pb2.runs[0].font.size = Pt(9)
        doc.add_paragraph()
        _SEV_COLOURS2 = {"Low": "D1FAE5", "Moderate": "FEF3C7",
                         "Severe": "FEE2E2", "High": "FEE2E2"}
        _SCOLS2 = ["Seg","DN","Regime","Severity","f_slug (Hz)","f_slug (slugs/min)",
                   "V_slug (m/s)","H_Ls","L_slug (m)",
                   "ΔP_pulse (kPa)","ΔP_design (kPa)","F_elbow (N)","F_design (N)"]
        _SW2 = [0.28,0.38,0.90,0.68,0.52,0.72,0.58,0.38,0.55,0.68,0.72,0.58,0.62]
        _sev_idx2 = _SCOLS2.index("Severity")
        def _slug_tbl2(doc, records):
            tbl = doc.add_table(rows=len(records)+1, cols=len(_SCOLS2))
            tbl.style = "Table Grid"
            _style_header(tbl.rows[0], font_size=8)
            for j, col in enumerate(_SCOLS2):
                tbl.rows[0].cells[j].text = col
            for i, rec in enumerate(records, start=1):
                row = tbl.rows[i]
                if i % 2 == 0:
                    for cell in row.cells:
                        _shd(cell, _ALT_BG)
                for j, col in enumerate(_SCOLS2):
                    row.cells[j].text = str(rec.get(col, ""))
                    _cell_font(row.cells[j], size_pt=8)
                sev_val = rec.get("Severity", "")
                if sev_val in _SEV_COLOURS2:
                    _shd(row.cells[_sev_idx2], _SEV_COLOURS2[sev_val])
            _set_col_widths(tbl, _SW2)
        for _sc, _slbl in _slug_branch_pairs:
            doc.add_heading(_slbl, level=2)
            _slug_tbl2(doc, _sc["slug_records"])
            doc.add_paragraph()

    # ── C. Header Configuration ───────────────────────────────────────────────
    _hdr_pairs = [(c, lbl) for c, lbl, h in zip(cases, case_labels, _is_hdr) if h]
    if _hdr_pairs:
        doc.add_heading("C.  Header Configuration", level=1)

        # Column headers for arm detail table
        _HC = ["Arm", "Seg", "From T (m)", "To T (m)", "L (m)", "Pipe", "ID (mm)",
               "Regime", "Q_gas (kg/h)", "Q_liq (m³/h)", "ΔP (kPa)",
               "P_in (bara)", "P_out (bara)"]
        _HW = [0.30, 0.28, 0.50, 0.50, 0.36, 0.48, 0.42, 0.72,
               0.58, 0.56, 0.48, 0.50, 0.50]

        def _hdr_detail_tbl(records):
            if not records:
                return
            tbl = doc.add_table(rows=len(records) + 1, cols=len(_HC))
            tbl.style = "Table Grid"
            _style_header(tbl.rows[0], font_size=8)
            for j, col in enumerate(_HC):
                tbl.rows[0].cells[j].text = col
            _col_map = {"Q_gas (kg/h)": "Q_gas_kgh", "Q_liq (m³/h)": "Q_liq_m3h"}
            for i, rec in enumerate(records, start=1):
                row = tbl.rows[i]
                if i % 2 == 0:
                    for cell in row.cells:
                        _shd(cell, _ALT_BG)
                seg_id = str(rec.get("Seg", ""))
                row.cells[0].text = (
                    "Left" if seg_id.startswith("L")
                    else "Right" if seg_id.startswith("R") else "T-seg")
                for j, col in enumerate(_HC[1:], start=1):
                    row.cells[j].text = str(rec.get(_col_map.get(col, col), ""))
                for cell in row.cells:
                    _cell_font(cell, size_pt=8)
            _set_col_widths(tbl, _HW)

        for _hc, _hlbl in _hdr_pairs:
            doc.add_heading(_hlbl, level=2)
            _hp  = _hc.get("header_pipe", {})
            _ts  = _hc.get("t_seg", {})
            _ltp = sorted(_hc.get("left_taps",  []), reverse=True)
            _rtp = sorted(_hc.get("right_taps", []))
            _cfg_rows = [
                ("Header pipe DN / PN",
                 f"{_hp.get('dn', '—')} / {_hp.get('pn', '—')}"),
                ("Header pipe material",
                 _hp.get("material", "—")),
            ]
            if _hp.get("lined"):
                _cfg_rows += [
                    ("Liner material",       _hp.get("liner_material", "—")),
                    ("Liner thickness (mm)", f"{_hp.get('liner_thickness_mm', 0):.1f}"),
                ]
            _cfg_rows += [
                ("Taps — left arm",
                 str(_hc.get("n_left", len(_ltp)))),
                ("Taps — right arm",
                 str(_hc.get("n_right", len(_rtp)))),
                ("Left tap distances from T (m)",
                 "  |  ".join(f"{p:.2f}" for p in _ltp) or "—"),
                ("Right tap distances from T (m)",
                 "  |  ".join(f"{p:.2f}" for p in _rtp) or "—"),
                ("Governing arm",
                 _hc.get("worst_arm", "—")),
                ("T-segment DN / PN",
                 f"{_ts.get('dn', '—')} / {_ts.get('pn', '—')}"),
                ("T-segment material",
                 _ts.get("material", "—")),
                ("T-segment length (m)",
                 f"{_ts.get('length', 0):.2f}"),
            ]
            _kv_table(doc, _cfg_rows, col_widths=(3.0, 3.47))
            doc.add_paragraph()
            if _hc.get("grid_records"):
                _p = doc.add_paragraph("Arm segment detail:")
                if _p.runs:
                    _p.runs[0].font.size = Pt(9)
                _hdr_detail_tbl(_hc["grid_records"])
                doc.add_paragraph()

    # ── D. DN Study ──────────────────────────────────────────────────────────
    if dn_study_data:
        _dns = dn_study_data
        _dn_p  = _dns["dn_primary"]
        _dn_a  = _dns["dn_alt"]
        _la_dn = _dns.get("label_a", case_labels[0] if case_labels else "A")
        _lb_dn = _dns.get("label_b", case_labels[1] if len(case_labels) > 1 else "B")
        _gp_h2 = _dns["gsr_h2_primary"]
        _gp_o2 = _dns["gsr_o2_primary"]
        _ga_h2 = _dns["gsr_h2_alt"]
        _ga_o2 = _dns["gsr_o2_alt"]
        _dp_p  = _dns["dp_gen_primary_mbar"]
        _dp_a  = _dns["dp_gen_alt_mbar"]
        _vd    = _dns["vel_data"]

        doc.add_heading("D.  DN Study — Branch Line Size Comparison", level=1)
        _pdn_intro = doc.add_paragraph(
            "The DN Study evaluates whether changing the branch pipe nominal diameter "
            f"from {_dn_p} (primary design) to {_dn_a} (alternative) improves system performance. "
            "A larger bore reduces mixture velocity and therefore frictional pressure drop "
            "in the branch, but it does so differently on each side because the two fluids "
            "(H₂ and O₂) have different physical properties and flow rates. "
            "This asymmetric response means that switching DN shifts the hydraulic balance "
            "between the two branches, changing the Generator ΔP. "
            "The study uses the same separator target pressures and goal-seek logic as the "
            "main case — only the branch DN changes; header sizes are kept constant."
        )
        _pdn_intro.paragraph_format.space_after = Pt(4)
        if _pdn_intro.runs:
            _pdn_intro.runs[0].font.size = Pt(9)
        doc.add_paragraph()
        _kv_table(doc, [
            ("Primary branch DN",         _dn_p),
            ("Alternative branch DN",     _dn_a),
            ("Header size",               "Unchanged for both cases"),
            ("H₂ separator target (bara)", f"{_dns.get('p_sep_h2', 0):.3f}"),
            ("O₂ separator target (bara)", f"{_dns.get('p_sep_o2', 0):.3f}"),
        ])
        doc.add_paragraph()

        doc.add_heading("Generator ΔP", level=2)
        _pdn_gen = doc.add_paragraph(
            "Each column represents a complete goal-seek run with all branches at the "
            "specified DN and headers unchanged. "
            "The preferred DN is the one with the smaller absolute Generator ΔP, "
            "subject to the erosion velocity constraint (V_m/V_e ≤ 1.0)."
        )
        _pdn_gen.paragraph_format.space_after = Pt(4)
        if _pdn_gen.runs:
            _pdn_gen.runs[0].font.size = Pt(9)
        _delta_mbar = _dp_a - _dp_p
        _winner = _dn_p if abs(_dp_p) <= abs(_dp_a) else _dn_a
        _kv3_table(doc, [
            (f"{_la_dn} branch inlet pressure (bara)",
                f"{_gp_h2['P_line_in']:.4f}", f"{_ga_h2['P_line_in']:.4f}"),
            (f"{_lb_dn} branch inlet pressure (bara)",
                f"{_gp_o2['P_line_in']:.4f}", f"{_ga_o2['P_line_in']:.4f}"),
            ("Generator ΔP (mbar)", f"{_dp_p:.1f}", f"{_dp_a:.1f}"),
            ("Change vs primary (mbar)", "—", f"{_delta_mbar:+.1f}"),
            ("Lower |Generator ΔP|", _winner, _winner),
        ], label_a=_dn_p, label_b=_dn_a)
        doc.add_paragraph()

        doc.add_heading("Pressure Drop by Case", level=2)
        _cases_dp = [
            (f"{_la_dn} branch", _gp_h2["dp_line"], _ga_h2["dp_line"]),
            (f"{_lb_dn} branch", _gp_o2["dp_line"], _ga_o2["dp_line"]),
            (f"{_la_dn} header", _gp_h2["dp_hdr"],  _ga_h2["dp_hdr"]),
            (f"{_lb_dn} header", _gp_o2["dp_hdr"],  _ga_o2["dp_hdr"]),
        ]
        _dp_rows = []
        for _lbl, _dp_pv, _dp_av in _cases_dp:
            _pct = (_dp_av - _dp_pv) / _dp_pv * 100 if abs(_dp_pv) > 1e-9 else 0.0
            _dp_rows.append((_lbl, f"{_dp_pv:.3f}", f"{_dp_av:.3f}  ({_pct:+.1f} %)"))
        _kv3_table(doc, _dp_rows,
                   label_a=f"{_dn_p} ΔP (kPa)", label_b=f"{_dn_a} ΔP (kPa)")
        doc.add_paragraph()

        doc.add_heading("Inlet Velocity — First Segment (Estimated)", level=2)
        _ratio_a = _vd["vm_a_alt"] / _vd["ve_a"] if _vd["ve_a"] > 0 else 0.0
        _ratio_b = _vd["vm_b_alt"] / _vd["ve_b"] if _vd["ve_b"] > 0 else 0.0
        _kv3_table(doc, [
            ("Effective ID (mm)",
                f"{_vd['D_p_mm']:.1f}", f"{_vd['D_a_mm']:.1f}"),
            ("Velocity scale factor (ID ratio²)", "1.00×", f"{_vd['vel_scale']:.2f}×"),
            (f"{_la_dn} V_m inlet (m/s)",
                f"{_vd['vm_a_primary']:.3f}", f"{_vd['vm_a_alt']:.3f}"),
            (f"{_la_dn} V_m / V_e",
                f"{_vd['vm_a_primary']/_vd['ve_a']:.2f}" if _vd["ve_a"] > 0 else "—",
                f"{_ratio_a:.2f}"),
            (f"{_lb_dn} V_m inlet (m/s)",
                f"{_vd['vm_b_primary']:.3f}", f"{_vd['vm_b_alt']:.3f}"),
            (f"{_lb_dn} V_m / V_e",
                f"{_vd['vm_b_primary']/_vd['ve_b']:.2f}" if _vd["ve_b"] > 0 else "—",
                f"{_ratio_b:.2f}"),
        ], label_a=_dn_p, label_b=_dn_a)
        _fig_caption(doc,
            "Velocity estimated via ID-ratio scaling. "
            "Erosion velocity V_e from API RP 14E (C = 100), primary case values.")
        doc.add_paragraph()

        doc.add_heading("Recommendation", level=2)
        _vel_ok   = _ratio_a <= 1.0 and _ratio_b <= 1.0
        _alt_wins = abs(_dp_a) < abs(_dp_p)
        if _alt_wins and _vel_ok:
            _rec = (f"{_dn_a} gives lower Generator |ΔP| ({_dp_a:.1f} vs {_dp_p:.1f} mbar) "
                    f"with V_m/V_e within the erosion limit. {_dn_a} is preferred.")
        elif _alt_wins:
            _rec = (f"{_dn_a} gives lower Generator |ΔP| ({_dp_a:.1f} vs {_dp_p:.1f} mbar) "
                    f"but estimated V_m/V_e may exceed 1.0 — verify erosion before selecting {_dn_a}.")
        else:
            _rec = (f"{_dn_p} (primary) gives lower Generator |ΔP| "
                    f"({_dp_p:.1f} vs {_dp_a:.1f} mbar). {_dn_a} appears undersized.")
        _p_rec = doc.add_paragraph(_rec)
        if _p_rec.runs:
            _p_rec.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_dn_study_report(
    dn_primary, dn_alt,
    label_a, label_b,
    gsr_h2_primary, gsr_o2_primary,
    gsr_h2_alt, gsr_o2_alt,
    dp_gen_primary_mbar, dp_gen_alt_mbar,
    vel_data,
    p_sep_h2, p_sep_o2,
):
    """Word report comparing two branch DN sizes across the full system."""
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    h = doc.add_heading(f"Pipe Size Study — {dn_primary} vs {dn_alt}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Branch Line DN Comparison  ·  {datetime.now().strftime('%d %B %Y  %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 0. Study objective ────────────────────────────────────────────────────
    doc.add_heading("Study Objective", level=1)
    _pobj = doc.add_paragraph(
        "This study quantifies the effect of changing the branch pipe nominal diameter "
        f"from {dn_primary} (primary design) to {dn_alt} on overall system hydraulics. "
        "A different-bore branch changes mixture velocity and therefore frictional pressure "
        "drop, but the H₂ and O₂ sides respond differently because their flow rates, "
        "gas densities, and fluid properties differ. "
        "This asymmetric response shifts the hydraulic balance between the two systems "
        "and therefore changes the Generator Differential Pressure — the net pressure "
        "across the electrolyser membrane arising from piping asymmetry. "
        "The study holds separator target pressures fixed and uses the same goal-seek "
        "logic as the main case; only the branch DN changes. "
        "The preferred DN is the one that gives the lower absolute Generator ΔP while "
        "keeping mixture velocities below the API RP 14E erosion limit (V_m/V_e ≤ 1.0)."
    )
    _pobj.paragraph_format.space_after = Pt(4)
    if _pobj.runs:
        _pobj.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 1. Study basis ────────────────────────────────────────────────────────
    doc.add_heading("Study Basis", level=1)
    _kv_table(doc, [
        ("Primary branch DN",           dn_primary),
        ("Alternative branch DN",        dn_alt),
        ("Header size",                  "Unchanged for both cases"),
        ("H₂ separator target (bara)",   f"{p_sep_h2:.3f}"),
        ("O₂ separator target (bara)",   f"{p_sep_o2:.3f}"),
        ("Correlation / voidage",        "As set in primary case"),
    ])
    doc.add_paragraph()

    # ── 2. Generator ΔP comparison ───────────────────────────────────────────
    doc.add_heading("Generator Differential Pressure", level=1)
    _pgen_dn = doc.add_paragraph(
        "The Generator ΔP is back-calculated by goal-seeking both the H₂ and O₂ systems "
        "to their respective separator target pressures and then taking the difference "
        "between the two branch inlet pressures. "
        "It represents the mechanical pressure differential across the electrolyser membrane "
        "resulting purely from hydraulic imbalance in the piping — not from "
        "electrochemical effects. "
        "A smaller absolute value indicates a more balanced system and lower membrane stress. "
        "The Change vs primary row shows how the differential shifts when the DN changes; "
        "a negative change means the imbalance decreased."
    )
    _pgen_dn.paragraph_format.space_after = Pt(4)
    if _pgen_dn.runs:
        _pgen_dn.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _delta_mbar = dp_gen_alt_mbar - dp_gen_primary_mbar
    _winner = dn_primary if abs(dp_gen_primary_mbar) <= abs(dp_gen_alt_mbar) else dn_alt
    _kv3_table(doc, [
        (f"{label_a} branch inlet pressure (bara)",
            f"{gsr_h2_primary['P_line_in']:.4f}",
            f"{gsr_h2_alt['P_line_in']:.4f}"),
        (f"{label_b} branch inlet pressure (bara)",
            f"{gsr_o2_primary['P_line_in']:.4f}",
            f"{gsr_o2_alt['P_line_in']:.4f}"),
        ("Generator ΔP (mbar)",
            f"{dp_gen_primary_mbar:.1f}",
            f"{dp_gen_alt_mbar:.1f}"),
        ("Change vs primary (mbar)", "—", f"{_delta_mbar:+.1f}"),
        ("Lower |Generator ΔP|", _winner, _winner),
    ], label_a=dn_primary, label_b=dn_alt)
    doc.add_paragraph()

    # ── 3. Pressure drop by case ──────────────────────────────────────────────
    doc.add_heading("Pressure Drop Summary by Case", level=1)
    _pdp_dn = doc.add_paragraph(
        "The table breaks down how pressure drop changes in each pipe section when the "
        f"branch DN is changed from {dn_primary} to {dn_alt}. "
        "Branch ΔP typically decreases with a larger bore (lower velocity, lower friction) "
        "and increases with a smaller bore. "
        "Header ΔP is unchanged because the header DN is kept constant. "
        "The percentage change quantifies the sensitivity of each section to the DN choice. "
        "A large change on one branch and a small change on the other explains why the "
        "Generator ΔP shifts when the DN is varied."
    )
    _pdp_dn.paragraph_format.space_after = Pt(4)
    if _pdp_dn.runs:
        _pdp_dn.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _cases_dp = [
        (f"{label_a} branch",  gsr_h2_primary["dp_line"], gsr_h2_alt["dp_line"]),
        (f"{label_b} branch",  gsr_o2_primary["dp_line"], gsr_o2_alt["dp_line"]),
        (f"{label_a} header",  gsr_h2_primary["dp_hdr"],  gsr_h2_alt["dp_hdr"]),
        (f"{label_b} header",  gsr_o2_primary["dp_hdr"],  gsr_o2_alt["dp_hdr"]),
    ]
    _rows_dp = []
    for _lbl, _dp_p, _dp_a in _cases_dp:
        _pct = (_dp_a - _dp_p) / _dp_p * 100 if abs(_dp_p) > 1e-9 else 0.0
        _rows_dp.append((
            _lbl,
            f"{_dp_p:.3f}",
            f"{_dp_a:.3f}  ({_pct:+.1f} %)",
        ))
    _kv3_table(doc, _rows_dp, label_a=f"{dn_primary} ΔP (kPa)", label_b=f"{dn_alt} ΔP (kPa)")
    doc.add_paragraph()

    # ── 4. Velocity estimate ──────────────────────────────────────────────────
    doc.add_heading("Inlet Velocity — First Segment (Estimated)", level=1)
    _pvel_dn = doc.add_paragraph(
        "Inlet velocity is the critical erosion check — particularly important when "
        "moving to a smaller bore where velocity increases. "
        "For a larger bore (typical in this study) velocity decreases, but the check "
        "is included for completeness. "
        "Velocity is estimated by scaling the primary-case value by the square of the "
        "ID ratio (pipe area ratio), which is exact for incompressible flow and "
        "conservative for two-phase compressible flow. "
        "V_m/V_e must remain ≤ 1.0 for continuous-service acceptance under API RP 14E "
        "(C = 100); values above this threshold require either a higher C-factor "
        "justification based on material and fluid corrosivity, or a bore increase."
    )
    _pvel_dn.paragraph_format.space_after = Pt(4)
    if _pvel_dn.runs:
        _pvel_dn.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _vd = vel_data
    _ratio_a = _vd["vm_a_alt"] / _vd["ve_a"] if _vd["ve_a"] > 0 else 0.0
    _ratio_b = _vd["vm_b_alt"] / _vd["ve_b"] if _vd["ve_b"] > 0 else 0.0
    _kv3_table(doc, [
        ("Effective ID (mm)",
            f"{_vd['D_p_mm']:.1f}", f"{_vd['D_a_mm']:.1f}"),
        ("Velocity scale factor (ID ratio²)", "1.00×", f"{_vd['vel_scale']:.2f}×"),
        (f"{label_a} V_m inlet (m/s)",
            f"{_vd['vm_a_primary']:.3f}", f"{_vd['vm_a_alt']:.3f}"),
        (f"{label_a} V_m / V_e",
            f"{_vd['vm_a_primary'] / _vd['ve_a']:.2f}" if _vd['ve_a'] > 0 else "—",
            f"{_ratio_a:.2f}"),
        (f"{label_b} V_m inlet (m/s)",
            f"{_vd['vm_b_primary']:.3f}", f"{_vd['vm_b_alt']:.3f}"),
        (f"{label_b} V_m / V_e",
            f"{_vd['vm_b_primary'] / _vd['ve_b']:.2f}" if _vd['ve_b'] > 0 else "—",
            f"{_ratio_b:.2f}"),
    ], label_a=dn_primary, label_b=dn_alt)
    doc.add_paragraph()
    _fig_caption(doc,
        "Velocity estimated via ID-ratio scaling: V_m(alt) = V_m(primary) × (ID_primary/ID_alt)². "
        "Erosion velocity V_e from primary case (API RP 14E, C = 100). Not recalculated for alt DN.")

    # ── 5. Recommendation ─────────────────────────────────────────────────────
    doc.add_heading("Recommendation", level=1)
    _vel_ok = _ratio_a <= 1.0 and _ratio_b <= 1.0
    _alt_better = abs(dp_gen_alt_mbar) < abs(dp_gen_primary_mbar)
    if _alt_better and _vel_ok:
        _rec = (
            f"{dn_alt} gives a lower Generator |ΔP| ({dp_gen_alt_mbar:.1f} mbar vs "
            f"{dp_gen_primary_mbar:.1f} mbar) with V_m/V_e within the erosion limit "
            f"for both branch lines. {dn_alt} is preferred for this duty."
        )
    elif _alt_better and not _vel_ok:
        _rec = (
            f"{dn_alt} gives a lower Generator |ΔP| ({dp_gen_alt_mbar:.1f} mbar vs "
            f"{dp_gen_primary_mbar:.1f} mbar) but the estimated inlet velocity ratio "
            f"V_m/V_e exceeds 1.0 for one or both branch lines. Verify erosion "
            f"against API RP 14E before selecting {dn_alt}."
        )
    else:
        _rec = (
            f"{dn_primary} (primary) gives a lower Generator |ΔP| "
            f"({dp_gen_primary_mbar:.1f} mbar vs {dp_gen_alt_mbar:.1f} mbar). "
            f"{dn_alt} appears undersized for this duty; the reduced pipe cross-section "
            f"does not improve the system differential pressure."
        )
    _p_rec = doc.add_paragraph(_rec)
    if _p_rec.runs:
        _p_rec.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    note = doc.add_paragraph(
        "Engineering Note: Velocity is estimated via ID-ratio scaling and is not exact. "
        "For a rigorous erosion check on the alternative DN, run it as the primary case. "
        "Correlations carry ±20–30 % uncertainty for H₂/O₂ over KOH systems."
    )
    if note.runs:
        note.runs[0].font.size      = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================================
# ENGINEERING CALCULATOR REPORT  (Fanno / RO / PSV / CV / DG / Pump / Line Size)
# ============================================================================

def generate_calculator_report(
    tool_name: str,
    subtitle: str,
    method_text,
    inputs_rows: list,
    results_rows: list,
    extra_tables=None,
    notes: str = None,
    fig=None,
    fig_caption_text: str = "",
    fig_width: int = 900,
    fig_height: int = 380,
) -> "BytesIO":
    """Generic Word report for any engineering calculator.

    tool_name    : e.g. "Fanno Flow"
    subtitle     : one-line description of what was calculated
    method_text  : str or list[str] — method / assumption paragraphs
    inputs_rows  : list of (label, value) tuples — inputs table
    results_rows : list of (label, value) tuples — results table
    extra_tables : list of dicts:
        KV table: {"title": str, "rows": [(label, value), ...], "widths": (w1, w2)}
        Grid table: {"title": str, "headers": [str, ...], "data": [[str, ...], ...],
                     "col_widths": [float, ...]}
    notes        : optional engineering note appended as footer
    fig          : optional Plotly figure
    fig_caption_text : caption shown above the figure image
    fig_width, fig_height : kaleido render dimensions in pixels
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)

    _add_footer_page_numbers(doc)

    def _body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        if p.runs:
            p.runs[0].font.size = Pt(9)
        return p

    _sec_n = [0]
    def _h1(title):
        _sec_n[0] += 1
        doc.add_heading(f"{_sec_n[0]}. {title}", level=1)

    # ── Title block ──────────────────────────────────────────────────────────
    h = doc.add_heading(f"{tool_name} — Calculation Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"{subtitle}  ·  {datetime.now().strftime('%d %B %Y  %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 1. Method ────────────────────────────────────────────────────────────
    _h1("Method")
    if isinstance(method_text, str):
        _body(method_text)
    else:
        for t in method_text:
            _body(t)
    doc.add_paragraph()

    # ── 2. Inputs ────────────────────────────────────────────────────────────
    _h1("Inputs")
    _kv_table(doc, inputs_rows)
    doc.add_paragraph()

    # ── 3. Results ───────────────────────────────────────────────────────────
    _h1("Results")
    _kv_table(doc, results_rows)
    doc.add_paragraph()

    # ── Extra tables ─────────────────────────────────────────────────────────
    if extra_tables:
        for tbl_spec in extra_tables:
            _h1(tbl_spec["title"])
            if "headers" in tbl_spec:
                headers = tbl_spec["headers"]
                data    = tbl_spec["data"]
                widths  = tbl_spec.get("col_widths")
                tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
                tbl.style = "Table Grid"
                _style_header(tbl.rows[0], font_size=8)
                for j, hdr in enumerate(headers):
                    tbl.rows[0].cells[j].text = hdr
                for i, row_vals in enumerate(data, start=1):
                    row = tbl.rows[i]
                    if i % 2 == 0:
                        for cell in row.cells:
                            _shd(cell, _ALT_BG)
                    for j, v in enumerate(row_vals):
                        row.cells[j].text = str(v)
                        _cell_font(row.cells[j], size_pt=8)
                if widths:
                    _set_col_widths(tbl, widths)
            else:
                _kv_table(doc, tbl_spec["rows"],
                          col_widths=tbl_spec.get("widths", (2.5, 3.5)))
            doc.add_paragraph()

    # ── Figure ───────────────────────────────────────────────────────────────
    if fig is not None:
        _h1("Chart")
        if fig_caption_text:
            _body(fig_caption_text)
        img = _fig_to_png(fig, width=fig_width, height=fig_height, scale=2)
        if img:
            doc.add_picture(BytesIO(img), width=Inches(6.2))
        else:
            doc.add_paragraph(
                "(Chart rendering timed out — install kaleido for embedded charts.)"
            )
        doc.add_paragraph()

    # ── Disclaimer ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    _note_text = (notes if notes else (
        f"FlowBench {tool_name} calculation. "
        "Provided for general engineering reference only. "
        "No warranty is given for accuracy, completeness, or fitness for any "
        "particular purpose. Validate all results independently before use in "
        "any design, procurement, or safety-critical application."
    ))
    _note_p = doc.add_paragraph(_note_text)
    if _note_p.runs:
        _note_p.runs[0].font.size = Pt(8)
        _note_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
